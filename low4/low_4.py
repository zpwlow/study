from PyQt5.QtWidgets import QWidget, QLabel, QPushButton, QGridLayout
from PyQt5.QtWidgets import QApplication
from PyQt5.QtCore import Qt
from captcha.image import ImageCaptcha
from PyQt5.QtGui import QFont
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtCore import *
from PyQt5 import QtCore, QAxContainer, QtWidgets
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
import os,sys, time
import glob
import random
import base64
#from datetime import   timedelta
import  datetime
from bs4 import BeautifulSoup
from docx import Document
from win32com import client
import fitz
import sqlite3
import requests
import zipfile
import shutil
from PIL import Image
from docx.shared import Pt



class QUnFrameWindow(QMainWindow):
    """
    无边框窗口类
    """
    def __init__(self):   #设置界面布局，界面大小，声名控件
        super(QUnFrameWindow, self).__init__(None) # 设置为顶级窗口
        self.setWindowTitle("low(**_**)")
        self.setWindowIcon(QIcon("logo.ico"))
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        self.setMinimumWidth(670)
        self.setMinimumHeight(560)
        self.resize(self.x, self.y)
        self.sign = 0  #标记
        self.grade = ''    #保存用户年级
        self.numble = ''  #保存用户密码
        self.data = []   #保存用户我的信息
        self.data1=[]   #保存用户学习报告信息
        self.time1 = [] #保存用户登录时间
        self.time2 = [] 
        self.htmls = []
        self.filename = []
        self.usr_window1_child1 = Usr_window1_child1()
        self.usr_window1_child2=Usr_window1_child2()
        self.usr_window1_child3 = Usr_window1_child3()
        self.usr_window2_child1 = Usr_window2_child1()
        self.usr_window2_child2 = Usr_window2_child2()
        self.usr_window2_child3 = Usr_window2_child3()
        self.centralwidget = QtWidgets.QWidget(self)
        self.centralwidget.setObjectName("centralwidget")
        self.horizontalLayout = QtWidgets.QHBoxLayout(self.centralwidget)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.splitter = QtWidgets.QSplitter(self.centralwidget)
        self.bar = self.menuBar()
        self.bar.setStyleSheet("QMenuBar{background-color:rgb(240,240,240);font-family:'宋体';font-size:14px;font-weight:Bold;};\
                         QMenuBar:hover{background-color:rgb(50, 170, 200)};font-family:'宋体';font-size:14px;font-weight:Bold}; ")
#        self.bar.setMaximumSize(50, 40)
        file =  self.bar.addMenu("文件")
        logonquit = QAction("退出登录", self)
        file.addAction(logonquit)
        quit = QAction("退出", self)
        file.addAction(quit)
        logonquit.triggered.connect(self.logonquit_fun)
        quit.triggered.connect(self.close_win)
        self.splitter.setOrientation(QtCore.Qt.Horizontal)
        self.splitter.setObjectName("splitter")
        self.horizontalLayout.addWidget(self.splitter)
        self.setCentralWidget(self.centralwidget)
        self.choice_status = Choice_status()   #选择身份界面
#        self.choice_status = Edit_usr()
        self.splitter.addWidget(self.choice_status) 
        
    def close_win(self):
        rely = QMessageBox.question(self, "提示!", "是否退出程序？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        if self.sign == 1:
            a = datetime.datetime.now()
            b = a - self.time2
            self.data1[3] = self.data1[3] + b.seconds/1.0
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],self.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        elif self.sign == 2:
            a = datetime.datetime.now()
            b = a - self.time2
            self.data1[4] = self.data1[4] + b.seconds/1.0
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude2_day =(?) where numble=(?)",(win.data1[3],self.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "练习", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        self.close()
        sys.exit()
    
    def logonquit_fun(self):
        rely = QMessageBox.question(self, "提示!", "是否退出登录？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        if self.sign == 1:
            a = datetime.datetime.now()
            b = a - self.time2
            self.data1[3] = self.data1[3] + b.seconds/1.0
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],self.numble))
            conn.commit()
            conn.close()
        elif self.sign == 2:
            a = datetime.datetime.now()
            b = a - self.time2
            self.data1[4] = self.data1[4] + b.seconds/1.0
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude2_day =(?) where numble=(?)",(win.data1[3],self.numble))
            conn.commit()
            conn.close()
        self.choice_status = Choice_status()   #选择身份界面 
        self.splitter.widget(0).setParent(None)  
        self.splitter.addWidget(self.choice_status) 
    
    
#选择身份界面 
class Choice_status(QFrame):
    def __init__(self):
        super(Choice_status, self).__init__()
        #self.setStyleSheet("background-color:white;")
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.prompt = QLabel("\t   请  选  择  身  份  登  录")
        self.Controller = QPushButton("管理员")
        self.user  = QPushButton("用户")
        self.devise_Ui()
        self.built_file()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.prompt.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:28px;font-weight:Bold;font-family:Arial;}")
        self.prompt.setMaximumSize(700, 70)
        self.Controller.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        self.Controller.setMaximumSize(700, 70)
        self.user.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        self.user.setMaximumSize(700, 70)
        self.Controller.clicked.connect(self.change_choice_status1)  #连接管理员登录界面
        self.user.clicked.connect(self.change_choice_status2)      #连接用户登录界面
        self.layout.addWidget(self.prompt, 0, 0)
        self.layout.addWidget(self.Controller, 1, 0)
        self.layout.addWidget(self.user, 2, 0)
    
    
    def built_file(self):  #建立文件夹跟数据库
        filepath = 'D:/项目数据库'
        if(not(os.path.exists(filepath))):   #创建文件夹。
            os.mkdir(filepath)
        filepath = 'D:/项目数据库/数据库'
        if(not(os.path.exists(filepath))):   #创建文件夹。
            os.mkdir(filepath)
        filepath = 'D:/项目数据库/tupian'
        if(not(os.path.exists(filepath))):   #创建文件夹。
            os.mkdir(filepath)
        filepath = 'D:/项目数据库/wen'
        if(not(os.path.exists(filepath))):   #创建文件夹。
            os.mkdir(filepath)
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        try:                                       #序号，      课件，     小学,     年级，      科目，      文件名   
            c.execute('''CREATE TABLE Education(no text,level1 text,level2 text,level3 text,level4 text,name text,
                         filename text)''')   #文件格式
        except:
            pass
        try:                                     #序号，文件内容
            c.execute('''CREATE TABLE filedata(no text,total LONGBLOB)''')
        except:
            pass
        try:                                    #序号，     课件/练习，      小学 ,    年级，    科目，      文件名   
            c.execute('''CREATE TABLE Education1(no text,level1 text,level2 text,level3 text,level4 text,name text,
                         filename text)''') #文件格式
        except:
            pass
        try:                                        #网址，  网址内容字节
            c.execute('''CREATE TABLE successfulurl(url text,howbyte integer)''')
        except:
            pass
        c.close()
        conn.close()
    
    def change_choice_status1(self):   #连接管理员登录界面
        win.splitter.widget(0).setParent(None)
        Controller_record().renovate_code()
        win.splitter.insertWidget(0, Controller_record())
        
    def change_choice_status2(self):    #连接用户登录界面
        win.splitter.widget(0).setParent(None)
        Usr_record().renovate_code()
        win.splitter.insertWidget(0, Usr_record())
    



#管理员注册界面
class Controller_logon(QFrame):
    def __init__(self):
        super(Controller_logon, self).__init__()
        #a = Usr_record().usrLineEdit.text()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        
        self.usr2 =  QLabel("用户:")
        self.usrname = QLabel("用户名：")
        self.pwd2 = QLabel("密码:")
        self.pwd3 = QLabel("确认密码:")
        self.usrLineEdit2 = QLineEdit()
        self.usrLineEdit = QLineEdit()
        self.pwdLineEdit2 = QLineEdit()
        self.pwdLineEdit3 = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn1 = QPushButton("注册")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        self.usr2.setMaximumSize(50, 40)
        self.usrname.setMaximumSize(60, 40)
        self.pwd2.setMaximumSize(50, 40)
        self.pwd3.setMaximumSize(80, 40)
        #设置QLabel 的字体颜色，大小，
        self.usr2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.usrname.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.pwd2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.pwd3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.usrLineEdit2.setMaximumSize(420, 40)
        self.usrLineEdit.setMaximumSize(420, 40)
        self.pwdLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit3.setMaximumSize(420, 40)
        self.codeLineEdit1.setMaximumSize(310, 40)
        #self.usrLineEdit2.setText(a)
        self.usrLineEdit2.setPlaceholderText("请输入手机号码")
        self.usrLineEdit.setPlaceholderText("请输入您的昵称")
        self.pwdLineEdit2.setPlaceholderText("请输入密码")
        self.pwdLineEdit3.setPlaceholderText("请重新输入密码")
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        self.usrLineEdit2.setFont(QFont("宋体" , 12))  #设置QLineEditn 的字体及大小
        self.usrLineEdit.setFont(QFont("宋体" , 12))
        self.pwdLineEdit2.setFont(QFont("宋体" , 12))
        self.pwdLineEdit3.setFont(QFont("宋体" , 12))
        self.codeLineEdit1.setFont(QFont("宋体" , 12))
        self.pwdLineEdit2.setEchoMode(QLineEdit.Password)
        self.pwdLineEdit3.setEchoMode(QLineEdit.Password)
        self.okBtn1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn1.setMaximumSize(420, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.codebel.setMaximumSize(100, 40)
        self.usrLineEdit2.returnPressed.connect(self.enterPress1)  #输入结束后按回车键跳到下一个控件
        self.usrLineEdit.returnPressed.connect(self.enterPress4)
        self.pwdLineEdit2.returnPressed.connect(self.enterPress2)
        self.pwdLineEdit3.returnPressed.connect(self.enterPress3)
        self.returnBtn.clicked.connect(self.change_choice_status1)  #点击返回键连接管理员登录界面
        self.codeLineEdit1.returnPressed.connect(self.accept)   #管理员忘记密码登录
        self.okBtn1.clicked.connect(self.accept)
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0, 1, 1, 1)
        self.layout.addWidget(self.usr2, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit2, 1, 5, 1, 14)
        self.layout.addWidget(self.usrname, 2, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit, 2, 5, 1, 14)
        self.layout.addWidget(self.pwd2, 3, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit2, 3, 5, 1, 14)
        self.layout.addWidget(self.pwd3, 4, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit3, 4, 5,1, 14 )
        self.layout.addWidget(self.codeLineEdit1, 5, 5, 1, 5)
        self.layout.addWidget(self.codebel, 5, 10, 1, 6)
        self.layout.addWidget(self.change_code, 5, 12, 1, 1)
        self.layout.addWidget(self.okBtn1, 6, 5, 1, 14)
        self.renovate_code()
     
    def renovate_code(self):
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小
    
    def checking1(self):  #注册时输入的号码检验是否已经注册过的
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from Controller")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit2.text() :
                return True
        c.close()
        conn.close()
        return False
    
    def checking2(self):  #登录时密码在数据库中保存过来
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        a = self.usrLineEdit2.text()
        b = self.usrLineEdit.text()
        c = self.pwdLineEdit2.text()
        conn.execute("INSERT INTO Controller VALUES(?,?,?)",(a, b, c))
        conn.commit()	
        conn.close()
        


    def enterPress1(self):  #注册-》用户框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit2.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit2.setFocus()
        elif len(self.usrLineEdit2.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit2.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码已注册！\n请您登录！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Controller_record())
        else:
            self.usrLineEdit.setFocus()
    
    def enterPress4(self):  #注册-》用户名框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit.text())==0:
            QMessageBox.about(self, "提示!", "用户名不能为空！" )
            self.usrLineEdit.setFocus()
        else:
            self.pwdLineEdit2.setFocus()
    
    def enterPress2(self):  #注册-》密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        else:
            self.pwdLineEdit3.setFocus()
     
    def enterPress3(self):   #注册-》确认密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        else:
            self.codeLineEdit1.setFocus()
    

    
    def accept(self):#注册时将账号密码保存并登录。
        if len(self.usrLineEdit2.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit2.setFocus()
        elif len(self.usrLineEdit2.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit2.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码已注册！\n请您登录！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Controller_record())
        elif len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        elif len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        elif self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            self.checking2()
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Controller_function())
            #连接主窗口界面。

    def change_choice_status1(self):   #连接管理员登录界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Controller_record())


#管理员登录界面
class Controller_record(QFrame):
    def __init__(self):
        super(Controller_record, self).__init__()
        #self.setStyleSheet("background-color:white;")
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usr1 = QLabel("用户：")
        self.pwd1 = QLabel("密码：")
        self.usrLineEdit = QLineEdit()
        self.pwdLineEdit = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn = QPushButton("登录")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.forgetbtn = QLabel( )
        self.logonbtn = QLabel()
        self.devise_Ui()
        self.found_sql()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        self.usr1.setMaximumSize(60, 60)
        #设置QLabel 的字体颜色，大小，
        self.usr1.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.pwd1.setMaximumSize(60, 60)
        #设置QLabel 的字体颜色，大小，
        self.pwd1.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.usrLineEdit.setPlaceholderText("请输入手机号码")
        self.usrLineEdit.setMaximumSize(420, 40)
        self.usrLineEdit.setFont(QFont("宋体" , 16))  #设置QLineEditn 的字体及大小
        self.pwdLineEdit.setMaximumSize(420, 40)
        self.pwdLineEdit.setPlaceholderText("请输入密码") 
        self.pwdLineEdit.setFont(QFont("宋体" , 16))
        self.pwdLineEdit.setEchoMode(QLineEdit.Password)
        self.okBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn.setMaximumSize(420, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        self.codeLineEdit1.setFont(QFont("宋体" , 16))
        self.codeLineEdit1.setMaximumSize(310, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.codebel.setMaximumSize(100, 40)
        self.forgetbtn.setText("<A href='www.baidu.com'>忘记密码</a>")
        self.logonbtn.setText("<A href='www.baidu.com'>注册</a>")
        self.forgetbtn.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:20px;font-weight:normal;font-family:Arial;}")
        self.logonbtn.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:20px;font-weight:normal;font-family:Arial;}")
        self.forgetbtn.setMaximumSize(90, 50)
        self.logonbtn.setMaximumSize(50, 50)
        
        self.okBtn.clicked.connect(self.accept)  
        self.forgetbtn.linkActivated.connect(self.controller_forgetbtn1)   #连接管理员忘记密码界面
        self.logonbtn.linkActivated.connect(self.controller_logonbtn1)   #连接管理员注册界面
        self.returnBtn.clicked.connect(self.change_controller_record1)  #点击返回键连接选择身份界面
        self.usrLineEdit.returnPressed.connect(self.enterPress1)  #输入结束后按回车键跳到下一个控件
        self.pwdLineEdit.returnPressed.connect(self.enterPress2)    
        self.codeLineEdit1.returnPressed.connect(self.accept)   #管理员忘记密码登录
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0, 1, 1, 1)
        self.layout.addWidget(self.usr1, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit, 1, 4, 1, 14)
        self.layout.addWidget(self.pwd1, 2, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit, 2, 4, 1, 14)
        self.layout.addWidget(self.codeLineEdit1, 3, 4, 1, 5)
        self.layout.addWidget(self.codebel, 3, 9, 1, 6)
        self.layout.addWidget(self.change_code, 3, 11, 1, 1)
        self.layout.addWidget(self.okBtn, 4, 4, 1, 14)
        self.layout.addWidget(self.forgetbtn, 5, 4, 1, 2)
        self.layout.addWidget(self.logonbtn, 5, 10, 1, 2)
        self.renovate_code()
     
    def renovate_code(self):
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小
        

    def change_controller_record1(self):   #点击返回键连接选择身份界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Choice_status())            
    
    def controller_forgetbtn1(self):   #连接管理员忘记密码界面
        win.splitter.widget(0).setParent(None)
        Controller_forget().renovate_code()
        win.splitter.insertWidget(0, Controller_forget())
    
    def controller_logonbtn1(self):    #连接管理员注册界面
        win.splitter.widget(0).setParent(None)
        Controller_logon().renovate_code()
        win.splitter.insertWidget(0, Controller_logon())
    

    
    
    def found_sql(self):   #创建保存用户信息的数据库
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        try:                                                      #号码            用户名                  密码
            c.execute('''CREATE TABLE Controller(numble text,usrname text,password text)''')	
        except:
            pass
        c.close()
        conn.close()
    
    def checking1(self):  #登录时检验号码是否没有注册
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from Controller")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit.text() :
                return False
        c.close()
        conn.close()
        return True
        
    def enterPress1(self):  #登录回车确定时判断文字框是否有输入
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text())!=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "该账号还未注册！\n请先注册！" )
        else:
            self.pwdLineEdit.setFocus()
    
    def enterPress2(self):  #登录回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit.setFocus()
        else:
            self.codeLineEdit1.setFocus()
        
    def accept(self):         #登录时判断密码是否正确
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text())!=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "该账号还未注册！\n请先注册！" )
        elif len(self.pwdLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit.setFocus()
        elif self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("select * from Controller")
            d =0
            for variate in c.fetchall():
                if variate[0]==self.usrLineEdit.text() and variate[2]== self.pwdLineEdit.text():
                    d = 1
                    break
            c.close()
            conn.close()
            if d == 1:
                win.splitter.widget(0).setParent(None)
                win.splitter.insertWidget(0, Controller_function())
            else:
                QMessageBox.about(self, "提示!", "账号或密码输入错误" )
            #连接主界面函数

#管理员忘记密码
class Controller_forget(QFrame):
    def __init__(self):
        super(Controller_forget, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usr2 =  QLabel("用户:")
        self.pwd2 = QLabel("密码:")
        self.pwd3 = QLabel("确认密码:")
        self.usrLineEdit2 = QLineEdit()
        self.pwdLineEdit2 = QLineEdit()
        self.pwdLineEdit3 = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn1 = QPushButton("确认")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        self.usr2.setMaximumSize(50, 40)
        self.pwd2.setMaximumSize(50, 40)
        self.pwd3.setMaximumSize(80, 40)
        #设置QLabel 的字体颜色，大小，
        self.usr2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.pwd2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.pwd3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.usrLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit3.setMaximumSize(420, 40)
        self.codeLineEdit1.setMaximumSize(310, 40)
        #self.usrLineEdit2.setText(a)
        self.usrLineEdit2.setPlaceholderText("请输入手机号码")
        self.pwdLineEdit2.setPlaceholderText("请输入新的密码")
        self.pwdLineEdit3.setPlaceholderText("请重新输入新的密码")
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        self.usrLineEdit2.setFont(QFont("宋体" , 12))  #设置QLineEditn 的字体及大小
        self.pwdLineEdit2.setFont(QFont("宋体" , 12))
        self.pwdLineEdit3.setFont(QFont("宋体" , 12))
        self.codeLineEdit1.setFont(QFont("宋体" , 12))
        self.pwdLineEdit2.setEchoMode(QLineEdit.Password)
        self.pwdLineEdit3.setEchoMode(QLineEdit.Password)
        self.okBtn1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn1.setMaximumSize(420, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.codebel.setMaximumSize(100, 40)
        self.usrLineEdit2.returnPressed.connect(self.enterPress1)  #输入结束后按回车键跳到下一个控件
        self.pwdLineEdit2.returnPressed.connect(self.enterPress2)
        self.pwdLineEdit3.returnPressed.connect(self.enterPress3)
        self.returnBtn.clicked.connect(self.change_choice_status1)  #点击返回键连接管理员登录界面
        self.codeLineEdit1.returnPressed.connect(self.accept)   #管理员忘记密码登录
        self.okBtn1.clicked.connect(self.accept)
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0, 1, 1, 1)
        self.layout.addWidget(self.usr2, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit2, 1, 5, 1, 14)
        self.layout.addWidget(self.pwd2, 2, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit2, 2, 5, 1, 14)
        self.layout.addWidget(self.pwd3, 3, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit3, 3, 5,1, 14)
        self.layout.addWidget(self.codeLineEdit1, 4, 5, 1, 5)
        self.layout.addWidget(self.codebel, 4, 10, 1, 6)
        self.layout.addWidget(self.change_code, 4, 12, 1, 1)
        self.layout.addWidget(self.okBtn1, 5, 5, 1, 14)
        self.renovate_code()
      
    def renovate_code(self):  
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小


    def checking1(self):  #输入的号码检验是否已经注册过的
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from Controller")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit2.text() :
                return False
        c.close()
        conn.close()
        return True
    
    def checking2(self):  #忘记密码时密码在数据库中修改过来
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from Controller")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit2.text() :
                conn.execute("update Controller set password=(?) where numble=(?)",(self.pwdLineEdit2.text(),variate[0],))
                break
        conn.commit()	
        conn.close()
        


    def enterPress1(self):  #忘记密码-》用户框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit2.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit2.setFocus()
        elif len(self.usrLineEdit2.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit2.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码未注册！\n请您先注册！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Controller_logon())
        else:
            self.pwdLineEdit2.setFocus()
    
    def enterPress2(self):  #忘记密码-》密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        else:
            self.pwdLineEdit3.setFocus()
     
    def enterPress3(self):   #注册-》确认密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        else:
            self.codeLineEdit1.setFocus()
    

    
    def accept(self):#忘记密码时将账号密码修改保存并登录。
        if len(self.usrLineEdit2.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit2.setFocus()
        elif len(self.usrLineEdit2.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit2.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码未注册！\n请您先注册！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Controller_logon())
        elif len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        elif len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        elif self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            self.checking2()
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Controller_function())
            #连接主窗口界面。
   
    def change_choice_status1(self):   #连接管理员登录界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Controller_record())
        
    
#管理员功能界面
class Controller_function(QFrame):
    def __init__(self):
        super(Controller_function, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.mainbutton1 = QPushButton("用户信息")  #用户功能界面的控件
        self.mainbutton2 = QPushButton("爬虫")
        self.mainbutton3 = QPushButton("添加资料")
        self.mainbutton4 = QPushButton("统计")
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        b= self.screenRect.height()  *1.0/5  
        a = self.screenRect.width() *1.0/5 
       
        self.mainbutton1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.mainbutton2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.mainbutton3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.mainbutton4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                QPushButton:hover{background-color:rgb(50, 170, 200)}")  
        self.mainbutton1.clicked.connect(self.select_fun1)
        self.mainbutton2.clicked.connect(self.select_fun2)
        self.mainbutton3.clicked.connect(self.select_fun3)
        self.mainbutton4.clicked.connect(self.select_fun4)
        self.layout.addWidget(self.mainbutton1, 0, 0)#往网格的不同坐标添加不同的组件
        self.layout.addWidget(self.mainbutton2, 0, 1)
        self.layout.addWidget(self.mainbutton3, 1, 0)
        self.layout.addWidget(self.mainbutton4, 1, 1)
        self.mainbutton1.setMaximumSize(a, b)
        self.mainbutton2.setMaximumSize(a, b)
        self.mainbutton3.setMaximumSize(a, b)
        self.mainbutton4.setMaximumSize(a, b)
    
    def select_fun1(self):  #连接用户信息
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_news())

    def select_fun2(self):   #连接爬虫
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_window1())
        
    def select_fun3(self):  #连接添加资料
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date())
    
    def select_fun4(self):  #连接统计信息
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_census())
        
        

#管理员用户信息界面
class Controller_news(QFrame):
    def __init__(self):
        super(Controller_news, self).__init__()
#        self.setStyleSheet("background-color:white;")
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.select_query = QComboBox()
        self.query = QLineEdit()
        self.search = QPushButton("搜索")
        self.returnBtn = QPushButton("返回")
        self.addusrbut = QPushButton("添加用户")
        self.editbut = QPushButton("编辑")
        self.delete1 = QPushButton("删除")
        self.usrs = QLabel("总用户数:")
        self.table = QTableWidget()
        self.data = []
        self.data1=[]
        self.devise_Ui()
        
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.select_query.addItems(['姓名', '号码', '学校', '年级'])
        self.usrs.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.select_query.setStyleSheet("QComboBox{font-family:'宋体';font-size: 18px;}")
        self.search.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.delete1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.editbut.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.addusrbut.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.query.setFont(QFont("宋体" , 12))
        self.select_query.setMaximumSize(70, 30)
        self.search.setMaximumSize(50, 30)
        self.returnBtn.setMaximumSize(60, 30)
        self.delete1.setMaximumSize(60, 30)
        self.editbut.setMaximumSize(60, 30)
        self.addusrbut.setMaximumSize(120, 30)
        self.usrs.setMaximumSize(120, 30)
        self.query.setMaximumSize(150, 30)
        self.returnBtn.clicked.connect(self.return_fun)
        self.search.clicked.connect(self.finddata1)
        self.delete1.clicked.connect(self.deletedata)
        self.addusrbut.clicked.connect(self.add_usr)
        self.editbut.clicked.connect(self.edit_usr)
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.usrs, 0, 4, 1, 1)
        self.layout.addWidget(self.addusrbut, 0, 8, 1, 1)
        self.layout.addWidget(self.editbut, 0, 11, 1, 1)
        self.layout.addWidget(self.delete1, 0, 14, 1, 1)
        self.layout.addWidget(self.select_query, 0, 16, 1, 1)
        self.layout.addWidget(self.query, 0, 17, 1, 1)
        self.layout.addWidget(self.search, 0, 18, 1, 1)
        self.finddata()
        
    
    def edit_usr(self):
        n =0
        for data in self.data1:
            if data[0].isChecked():
                n+=1
        if n != 1:
            QMessageBox.about(self, "提示!", "抱歉，该功能只能选择一个用户进行编辑！！" )
            return
        for data in self.data1:
            if data[0].isChecked():
                print(data[1])
                print(1)
                win.numble = data[1]
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Edit_usr())

    
    def add_usr(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Cont_Usr_logon())

    
    
    def finddata(self):
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User_data")
        self.data = c.fetchall()
        b = len(self.data)
        self.usrs1 = QLabel(str(b))  #读取用户人数
        self.table.setRowCount(b)
        self.table.setColumnCount(7)
        self.table.setHorizontalHeaderLabels(['选择','用户', '姓名','出生年月','性别', '学校', '年级'])
        i=0
        for variate in self.data:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.data1.append([ck, variate[0]])
            for j  in range(6):
                itemContent = variate[j]
                self.table.setItem(i, j+1, QTableWidgetItem(itemContent))
            i= i+1
        c.close()
        conn.close()
        self.usrs1.setStyleSheet("QLabel{color:rgb(0,0,240,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.usrs1.setMaximumSize(50, 30)
        self.layout.addWidget(self.usrs1, 0, 5, 1, 1)
        self.table.setStyleSheet("QTableWidget{background-color:rgb(235,235,235);font:13pt '宋体';font-weight:Bold;};")
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)  #不能编辑table
        self.layout.addWidget(self.table, 1, 0, 15, 19)
       
        
    def finddata1(self):
        if (self.select_query.currentText() == '姓名'):
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            d = self.query.text()
            c.execute("select * from User_data where name = (?)", (d, ))
            self.data = c.fetchall()
            b = len(self.data)
            if b:
                self.table.setRowCount(b)
                i=0
                for variate in data1:
                    ck = QCheckBox()
                    h = QHBoxLayout()
                    h.setAlignment(Qt.AlignCenter)
                    h.addWidget(ck)
                    w = QWidget()
                    w.setLayout(h)
                    self.table.setCellWidget(i, 0, w)
                    self.data1.append([ck, variate[0]])
                    for j  in range(6):
                        itemContent = variate[j]
                        self.table.setItem(i, j+1, QTableWidgetItem(itemContent))
                    self.table.item(i ,1).setForeground(QBrush(QColor(255,0,0)))
                    i= i+1
                c.close()
                conn.close()
            else:
                QMessageBox.about(self, "提示!","没有该用户的任何信息" )
        elif (self.select_query.currentText() == '号码'):
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            d = self.query.text()
            c.execute("select * from User_data where numble = (?)", (d, ))
            self.data = c.fetchall()
            b = len(self.data)
            if b:
                self.table.setRowCount(b)
                i=0
                for variate in self.data:
                    ck = QCheckBox()
                    h = QHBoxLayout()
                    h.setAlignment(Qt.AlignCenter)
                    h.addWidget(ck)
                    w = QWidget()
                    w.setLayout(h)
                    self.table.setCellWidget(i, 0, w)
                    self.data1.append([ck, variate[0]])
                    for j  in range(6):
                        itemContent = variate[j]
                        self.table.setItem(i, j+1, QTableWidgetItem(itemContent))
                    self.table.item(i ,1).setForeground(QBrush(QColor(255,0,0)))
                    i= i+1
                c.close()
                conn.close()
            else:
                QMessageBox.about(self, "提示!","没有该用户的任何信息" )
        elif (self.select_query.currentText() == '学校'):
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            d = self.query.text()
            c.execute("select * from User_data where school = (?)", (d, ))
            self.data = c.fetchall()
            b = len(self.data)
            if b:
                self.table.setRowCount(b)
                i=0
                for variate in self.data:
                    ck = QCheckBox()
                    h = QHBoxLayout()
                    h.setAlignment(Qt.AlignCenter)
                    h.addWidget(ck)
                    w = QWidget()
                    w.setLayout(h)
                    self.table.setCellWidget(i, 0, w)
                    self.data1.append([ck, variate[0]])
                    for j  in range(6):
                        itemContent = variate[j]
                        self.table.setItem(i, j+1, QTableWidgetItem(itemContent))
                    self.table.item(i ,5).setForeground(QBrush(QColor(255,0,0)))
                    i= i+1
                c.close()
                conn.close()
            else:
                QMessageBox.about(self, "提示!","没有该用户的任何信息" )
        elif (self.select_query.currentText() == '年级'):
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            d = self.query.text()
            c.execute("select * from User_data where grade = (?)", (d, ))
            self.data = c.fetchall()
            b = len(self.data)
            if b:
                self.table.setRowCount(b)
                i=0
                for variate in self.data:
                    ck = QCheckBox()
                    h = QHBoxLayout()
                    h.setAlignment(Qt.AlignCenter)
                    h.addWidget(ck)
                    w = QWidget()
                    w.setLayout(h)
                    self.table.setCellWidget(i, 0, w)
                    self.data1.append([ck, variate[0]])
                    for j  in range(6):
                        itemContent = variate[j]
                        self.table.setItem(i, j+1, QTableWidgetItem(itemContent))
                    self.table.item(i ,6).setForeground(QBrush(QColor(255,0,0)))
                    i= i+1
                c.close()
                conn.close()
            else:
                QMessageBox.about(self, "提示!","没有该用户的任何信息" )
    
    def deletedata(self):
        rely = QMessageBox.question(self, "提示!", "删除会造成数据无法恢复！！！\n确定删除？？？" , QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        removeline =[]
        n =0
        for data in self.data1:
            if data[0].isChecked():
                n+=1
        if n == 0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件，请您重新选择！！" )
            return
        for data in self.data1:
            if data[0].isChecked():
                row = self.table.rowCount()
                for x in range(row, 0, -1):
                    if data[1] == self.table.item(x ,1).text():
                        self.table.removeRow(x)
                        removeline.append(data)
        if len(removeline)>0:
            for line in removeline:
                self.data1.remove(line)
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        for line in removeline:
            c.execute("delete from User_data where numble = (?)", (line[1], ))
            c.execute("delete from User_data1 where numble = (?)", (line[1], ))
        conn.commit()
        c.close()
        conn.close()
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_function())


#编辑用户信息
class Edit_usr(QFrame):
    def __init__(self):
        super(Edit_usr, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.sure = QPushButton("确认")
        self.amend = QPushButton("修改用户密码")
        self.returnBtn = QPushButton("返回")
        self.name =  QLabel("姓名:")
        self.year  = QLabel("出生年月")
        self.yearcb = QComboBox()
        self.monthcb = QComboBox()
        self.sex =  QLabel("性别:")
        self.sexcb = QComboBox()
        self.school = QLabel("学校:")
        self.grade = QLabel("选择年级")
        self.gradecb = QComboBox()
        self.nameEdit = QLineEdit()
        self.schoolEiit = QLineEdit()
        self.devise_Ui()
     
   
    def devise_Ui(self):
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User_data where numble=(?)", (win.numble, ))
        for data in c.fetchall():
            self.data = data
        c.close()
        conn.close()
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        yearnb = []
        for  i in range(1980, 2020):
            yearnb.append(str(i))
        monthmb = []
        for i in range(1, 13):
            monthmb.append(str(i))
        grade = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级', 
                     '初一', '初二', '初三', 
                     '高一', '高二', '高三']
        self.sex.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        self.amend.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(255,0,0,255);}")
        self.school.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.grade.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.name.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.year.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.sure.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        
        self.nameEdit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.schoolEiit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.name.setMaximumSize(50, 40)
        self.school.setMaximumSize(50, 40)
        self.returnBtn.setMaximumSize(60, 40)
        self.amend.setMaximumSize(190, 40)
        self.year.setMaximumSize(95, 40)
        self.sex.setMaximumSize(50, 40)
        self.grade.setMaximumSize(95, 40)
        self.nameEdit.setMaximumSize(420, 40)
        self.schoolEiit.setMaximumSize(420, 40)
        self.sure.setMaximumSize(420, 40)
        self.sexcb.setMaximumSize(420, 40)
        self.gradecb.setMaximumSize(420, 40)
        self.yearcb.setMaximumSize(220, 40)
        self.monthcb.setMaximumSize(175, 40)
        self.sexcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.yearcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.monthcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.gradecb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.sexcb.addItems(['男', '女'])
        print(self.data[3])
        self.sexcb.setCurrentText(self.data[3])  #设置文本的默认选项
        self.yearcb.addItems(yearnb)
        self.yearcb.setCurrentText(self.data[2][0:4])  #设置文本的默认选项
        self.monthcb.addItems(monthmb)
        self.monthcb.setCurrentText(self.data[2][5:7])  #设置文本的默认选项
        self.gradecb.addItems(grade)
        self.gradecb.setCurrentText(self.data[5])  #设置文本的默认选项
        self.nameEdit.setText(self.data[1])
        self.schoolEiit.setText(self.data[4])
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.amend, 0, 10, 1, 1)
        self.layout.addWidget(self.name, 1, 3, 1, 1)
        self.layout.addWidget(self.nameEdit, 1, 4, 1, 18)
        self.layout.addWidget(self.sex, 2, 3, 1, 1)
        self.layout.addWidget(self.sexcb, 2, 4, 1, 18)
        self.layout.addWidget(self.year, 3, 3, 1, 1)
        self.layout.addWidget(self.yearcb,3, 4, 1, 8 )
        self.layout.addWidget(self.monthcb, 3, 10, 1, 7)
        self.layout.addWidget(self.school, 4, 3, 1, 1)
        self.layout.addWidget(self.schoolEiit, 4, 4, 1, 18)
        self.layout.addWidget(self.grade, 5, 3, 1, 1)
        self.layout.addWidget(self.gradecb, 5, 4, 1, 18)
        self.layout.addWidget(self.sure, 6, 4, 1, 18)
        self.sure.clicked.connect(self.connect_fun)
        self.returnBtn.clicked.connect(self.return_fun)
        self.amend.clicked.connect(self.connect_fun1)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Controller_news())
    
    def save_data(self):
        a =  self.nameEdit.text()
        b = self.yearcb.currentText() +'-' +self.monthcb.currentText()
        c = self.sexcb.currentText()
        d = self.schoolEiit.text()
        e = self.gradecb.currentText()
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        conn.execute("update User_data set name =(?),birthday=(?),sex=(?),school=(?),grade=(?) where numble=(?)",(a, b, c, d, e, win.numble))
        conn.commit()	
        conn.close()
    
    def connect_fun(self):
        win.splitter.widget(0).setParent(None)
        self.save_data()
        Controller_news().devise_Ui()
        win.splitter.insertWidget(0, Controller_news())
        
    def connect_fun1(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, chang_Usr_amend())

#修改用户密码
class chang_Usr_amend(QFrame):
    def __init__(self):
        super(chang_Usr_amend, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usrlab = QLabel("账号:")
        self.amendlab1 = QLabel("原密码:")
        self.amendlab2 = QLabel("新密码:")
        self.amendlab3 = QLabel("确认密码:")
        self.amendedit1 = QLineEdit()
        self.amendedit2 = QLineEdit()
        self.amendedit3 = QLineEdit()
        self.sure = QPushButton("确认修改")
        self.returnBtn = QPushButton("返回")
        self.devise_Ui()
        
    def devise_Ui(self):
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User where numble=(?)", (win.numble, ))
        for data in c.fetchall():
            self.data = data
        c.close()
        conn.close()
        self.usrlab1 = QLabel(win.numble)
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (350, 0, 0, 0)
        self.usrlab.setMaximumSize(80, 40)
        self.amendlab1.setMaximumSize(80, 40)
        self.amendlab2.setMaximumSize(80, 40)
        self.amendlab3.setMaximumSize(100, 40)
        #设置QLabel 的字体颜色，大小，
        self.usrlab.setStyleSheet("QLabel{color:rgb(100,100,100,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.usrlab1.setStyleSheet("QLabel{color:rgb(100,100,100,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.amendlab1.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")      
        self.amendlab2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")      
        self.amendlab3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")      
        self.usrlab1.setMaximumSize(420, 40)
        self.amendedit1.setMaximumSize(420, 40)
        self.amendedit2.setMaximumSize(420, 40)
        self.amendedit3.setMaximumSize(420, 40)
        self.sure.setMaximumSize(420, 40)
        self.amendedit1.setText(self.data[2])
        self.amendedit2.setPlaceholderText("请输入新密码")
        self.amendedit3.setPlaceholderText("请重新输入密码")
        self.amendedit1.setFont(QFont("宋体" , 16))  #设置QLineEditn 的字体及大小
        self.amendedit2.setFont(QFont("宋体" , 16))
        self.amendedit3.setFont(QFont("宋体" , 16))
        self.amendedit1.setEchoMode(QLineEdit.Password)
        self.amendedit2.setEchoMode(QLineEdit.Password)
        self.amendedit3.setEchoMode(QLineEdit.Password)
        self.sure.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.returnBtn.clicked.connect(self.return_fun)
        self.amendedit2.returnPressed.connect(self.enterPress2)
        self.amendedit3.returnPressed.connect(self.accept)
        self.sure.clicked.connect(self.accept)
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.usrlab, 1, 3, 1, 1)
        self.layout.addWidget(self.usrlab1, 1, 5, 1, 5)
        self.layout.addWidget(self.amendlab1, 2, 3, 1, 1)
        self.layout.addWidget(self.amendedit1, 2, 5, 1, 5)
        self.layout.addWidget(self.amendlab2, 3, 3, 1, 1)
        self.layout.addWidget(self.amendedit2, 3, 5, 1, 5)
        self.layout.addWidget(self.amendlab3, 4, 3, 1, 1)
        self.layout.addWidget(self.amendedit3, 4, 5, 1, 5)
        self.layout.addWidget(self.sure, 5, 5, 1, 5)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Edit_usr())
        
    
    def enterPress2(self):
        if len(self.amendedit2.text()) == 0:
            QMessageBox.about(self, "提示!", "新密码框不能为空！" )
            self.amendedit2.setFocus()
        else:
            self.amendedit3.setFocus()
            
    def accept(self):
        if len(self.amendedit2.text()) == 0:
            QMessageBox.about(self, "提示!", "新密码框不能为空！" )
            self.amendedit2.setFocus()
        elif len(self.amendedit3.text()) == 0:
            QMessageBox.about(self, "提示!", "确认密码框不能为空！" )
            self.amendedit3.setFocus()
        elif self.amendedit3.text() != self.amendedit2.text():
            QMessageBox.about(self, "提示!", "前后密码输入不一样！" )
            self.amendedit3.setFocus()
        else:
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User set password=(?) where numble=(?)",(self.amendedit2.text(),win.numble,))
            conn.commit()
            c.close()
            conn.close()
            QMessageBox.about(self, "提示!", "修改密码成功！！！" )
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Edit_usr())
            
   

#管理员添加用户注册界面
class Cont_Usr_logon(QFrame):
    def __init__(self):
        super(Cont_Usr_logon, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usr =  QLabel("号码:")
        self.usrname = QLabel("用户名：")
        self.pwd2 = QLabel("密码:")
        self.pwd3 = QLabel("确认密码:")
        self.usrLineEdit = QLineEdit()
        self.usrLineEdit2 = QLineEdit()
        self.pwdLineEdit2 = QLineEdit()
        self.pwdLineEdit3 = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn1 = QPushButton("注册")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        self.usr.setMaximumSize(50, 40)
        self.pwd2.setMaximumSize(50, 40)
        self.pwd3.setMaximumSize(80, 40)
        #设置QLabel 的字体颜色，大小，
        self.usr.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.usrname.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.pwd2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.pwd3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.usrLineEdit.setMaximumSize(420, 40)
        self.usrLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit3.setMaximumSize(420, 40)
        self.codeLineEdit1.setMaximumSize(310, 40)
        #self.usrLineEdit2.setText(a)
        self.usrLineEdit.setPlaceholderText("请输入手机号码")
        self.usrLineEdit2.setPlaceholderText("请输入您的昵称")
        self.pwdLineEdit2.setPlaceholderText("请输入密码")
        self.pwdLineEdit3.setPlaceholderText("请重新输入密码")
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        self.usrLineEdit2.setFont(QFont("宋体" , 12))  #设置QLineEditn 的字体及大小
        self.usrLineEdit.setFont(QFont("宋体" , 12)) 
        self.pwdLineEdit2.setFont(QFont("宋体" , 12))
        self.pwdLineEdit3.setFont(QFont("宋体" , 12))
        self.codeLineEdit1.setFont(QFont("宋体" , 12))
        self.pwdLineEdit2.setEchoMode(QLineEdit.Password)
        self.pwdLineEdit3.setEchoMode(QLineEdit.Password)
        self.okBtn1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn1.setMaximumSize(420, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.codebel.setMaximumSize(100, 40)
        self.usrLineEdit.returnPressed.connect(self.enterPress1)  #输入结束后按回车键跳到下一个控件
        self.usrLineEdit2.returnPressed.connect(self.enterPress4)
        self.pwdLineEdit2.returnPressed.connect(self.enterPress2)
        self.pwdLineEdit3.returnPressed.connect(self.enterPress3)
        self.codeLineEdit1.returnPressed.connect(self.accept) #验证码输入后回车直接验证是否可以登录
        self.okBtn1.clicked.connect(self.accept)
        self.returnBtn.clicked.connect(self.return_record)   #点击返回键返回登录界面
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0,1 , 1, 1)
        self.layout.addWidget(self.usr, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit, 1, 5, 1, 14)
        self.layout.addWidget(self.usrname, 2, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit2, 2, 5, 1, 14)
        self.layout.addWidget(self.pwd2, 3, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit2, 3, 5, 1, 14)
        self.layout.addWidget(self.pwd3, 4, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit3, 4, 5,1, 14 )
        self.layout.addWidget(self.codeLineEdit1, 5, 5, 1, 5)
        self.layout.addWidget(self.codebel, 5, 10, 1, 6)
        self.layout.addWidget(self.change_code, 5, 12, 1, 1)
        self.layout.addWidget(self.okBtn1, 6, 5, 1, 14)
        self.renovate_code()
    
    def renovate_code(self):
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小

    def return_record(self):
        win.splitter.widget(0).setParent(None)
        Controller_news().finddata()
        win.splitter.insertWidget(0, Controller_news())
    
    def checking1(self):  #注册时输入的号码检验是否已经注册过的
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit.text() :
                return True
        c.close()
        conn.close()
        return False
    
    def checking2(self):  #注册时密码在数据库中保存过来
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        a = self.usrLineEdit.text()
        b = self.usrLineEdit2.text()
        c = self.pwdLineEdit2.text()
        conn.execute("INSERT INTO User VALUES(?,?,?)",(a, b, c))
        conn.commit()	
        conn.close()
        
    def enterPress1(self):  #注册-》用户框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码已注册！\n请您重新确定一下！" )
            
        else:
            self.usrLineEdit2.setFocus()
    
    def enterPress4(self):  #注册-》用户名框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "用户名不能为空！" )
            self.usrLineEdit2.setFocus()
        else:
            self.pwdLineEdit2.setFocus()
    
    def enterPress2(self):  #注册-》密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        else:
            self.pwdLineEdit3.setFocus()
     
    def enterPress3(self):   #注册-》确认密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        else:
            self.codeLineEdit1.setFocus()
    

    def accept(self):#注册时将账号密码保存并登录。
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码已注册！\n请您重新确认一下！" )
            
        elif len(self.usrLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "用户名不能为空！" )
            self.usrLineEdit2.setFocus()
        elif len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        elif len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        elif self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            win.numble = self.usrLineEdit.text()
            self.checking2()
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Cont_informent())
      


#管理员填写用户信息填写
class Cont_informent(QFrame):
    def __init__(self):
        super(Cont_informent, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.sure = QPushButton("确认")
        self.chang_image  = QPushButton("换头像")
        self.name =  QLabel("姓名:")
        self.year  = QLabel("出生年月")
        self.yearcb = QComboBox()
        self.monthcb = QComboBox()
        self.sex =  QLabel("性别:")
        self.sexcb = QComboBox()
        self.school = QLabel("学校:")
        self.grade = QLabel("选择年级")
        self.gradecb = QComboBox()
        self.nameEdit = QLineEdit()
        self.tupian = QLabel()
        self.schoolEiit = QLineEdit()
        self.devise_Ui()
     
   
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (100, 0, 0, 0)
        yearnb = []
        for  i in range(1980, 2020):
            yearnb.append(str(i))
        monthmb = []
        for i in range(1, 13):
            monthmb.append(str(i))
        grade = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级', 
                     '初一', '初二', '初三', 
                     '高一', '高二', '高三']
        self.sex.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.chang_image.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.school.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.grade.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.name.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.year.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.sure.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        self.nameEdit.setPlaceholderText("请输入姓名")
        self.schoolEiit.setPlaceholderText("请输入学校名称")
        self.nameEdit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.schoolEiit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.name.setMaximumSize(50, 40)
        self.chang_image.setMaximumSize(90, 40)
        self.school.setMaximumSize(50, 40)
        self.year.setMaximumSize(95, 40)
        self.sex.setMaximumSize(50, 40)
        self.grade.setMaximumSize(95, 40)
        self.nameEdit.setMaximumSize(420, 40)
        self.schoolEiit.setMaximumSize(420, 40)
        self.sure.setMaximumSize(420, 40)
        self.sexcb.setMaximumSize(420, 40)
        self.gradecb.setMaximumSize(420, 40)
        self.yearcb.setMaximumSize(220, 40)
        self.monthcb.setMaximumSize(175, 40)
        self.tupian.setMaximumSize(250, 250)
        self.sexcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.yearcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.monthcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.gradecb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.sexcb.addItems(['男', '女'])
        self.yearcb.addItems(yearnb)
        self.monthcb.addItems(monthmb)
        self.gradecb.addItems(grade)
        self.layout.addWidget(self.tupian, 1, 1, 4, 4)
        self.layout.addWidget(self.chang_image, 4, 2, 1, 2)
        self.layout.addWidget(self.name, 1, 6, 1, 1)
        self.layout.addWidget(self.nameEdit, 1, 8, 1, 8)
        self.layout.addWidget(self.sex, 2, 6, 1, 1)
        self.layout.addWidget(self.sexcb, 2, 8, 1, 8)
        self.layout.addWidget(self.year, 3, 6, 1, 1)
        self.layout.addWidget(self.yearcb,3, 8, 1, 4 )
        self.layout.addWidget(self.monthcb, 3, 11, 1, 7)
        self.layout.addWidget(self.school, 4, 6, 1, 1)
        self.layout.addWidget(self.schoolEiit, 4, 8, 1, 8)
        self.layout.addWidget(self.grade, 5, 6, 1, 1)
        self.layout.addWidget(self.gradecb, 5, 8, 1, 8)
        self.layout.addWidget(self.sure, 6, 8, 1, 8)
        self.image()
        self.sure.clicked.connect(self.connect_fun)
        self.chang_image.clicked.connect(self.chang_fun)
    
    def image(self):
        self.image_path = "D:/项目数据库/头像/a7.jpeg"
        self.file = os.path.splitext(self.image_path)[1]
        self.tupian.setPixmap(QPixmap(self.image_path))
        self.tupian.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
    
    def chang_fun(self):
        path, _ = QFileDialog.getOpenFileName(self, '请选择文件', 
        'D:\\', 'image(*.jpg)')
        if path:
            self.image_path = path
            self.file = os.path.splitext(self.image_path)[1]
            self.tupian.setPixmap(QPixmap(self.image_path))
            self.tupian.setScaledContents (True) # 让图片自适应label大小
        else:
            self.image()
            
        
    
    def save_data(self):
        a =  self.nameEdit.text()
        b = self.yearcb.currentText() +'-' +self.monthcb.currentText()
        c = self.sexcb.currentText()
        d = self.schoolEiit.text()
        e = self.gradecb.currentText()
        with open(self.image_path, "rb") as f:
            total=base64.b64encode(f.read())   #将文件转换为字节。
        f.close()
        win.grade = e
        win.data = [win.numble, a, b, c, d, e, total, self.file]
        ab = '%Y-%m-%d %H:%M:%S'
        theTime = datetime.datetime.now().strftime(ab) 
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        conn.execute("INSERT INTO User_data VALUES(?,?,?,?,?,?,?,?)",(win.numble, a, b, c, d, e,total, self.file))
        conn.execute("INSERT INTO User_data1 VALUES(?,?,?,?,?,?)",(win.numble, theTime , 1,0.0,0.0, theTime))
        win.data1 = [win.numble, theTime , 1, 0.0,0.0, theTime ]
        conn.commit()	
        conn.close()
    
    def connect_fun(self):
        if len(self.nameEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "姓名框不能为空！！" )
            self.nameEdit.setFocus()
        if len(self.schoolEiit.text()) == 0:
            QMessageBox.about(self, "提示!", "学校框不能为空！！" )
            self.schoolEiit.setFocus()
        else:
            self.save_data()
            win.splitter.widget(0).setParent(None)
            Controller_news().devise_Ui()
            win.splitter.insertWidget(0, Controller_news())
            
        



    




#爬虫选择爬取课件或练习
class Controller_window1(QFrame):  
    def __init__(self):
        super(Controller_window1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.win1button1 = QPushButton("返回")  #爬虫中的控件
        self.win1button2 = QPushButton("爬取课件")
        self.win1button3 = QPushButton("爬取练习")
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        a = self.x *1.0/3
        b = self.y*1.0/5
        self.win1button1.setMaximumSize(a, b)
        self.win1button2.setMaximumSize(a, b)
        self.win1button3.setMaximumSize(a, b)
        self.win1button1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.win1button2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.win1button3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.win1button1.clicked.connect(self.return_fun)
        self.win1button2.clicked.connect(self.select_fun1)
        self.win1button3.clicked.connect(self.select_fun2)
        self.layout.addWidget(self.win1button1, 0, 0)#往网格的不同坐标添加不同的组件
        self.layout.addWidget(self.win1button2, 1, 0)
        self.layout.addWidget(self.win1button3, 2, 0) 
    
    def return_fun(self):   #返回功能界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_function())
       
    def select_fun1(self):  #爬取课件界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Courseware())
        
    def select_fun2(self):  #爬取练习界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Practice())
    

#爬取课件的界面
class Controller_win_Courseware(QFrame):
    def __init__(self):
        super(Controller_win_Courseware, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Courseware_but1 = QPushButton("返回")  #爬取课件中的控件
        self.Courseware_but2 = QPushButton("小学")
        self.Courseware_but3 = QPushButton("初中")
        self.Courseware_but4 = QPushButton("高中")
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        a = self.x *1.0/8
        b = self.y*1.0/3
        self.Courseware_but1.setMaximumSize(a, b)
        self.Courseware_but2.setMaximumSize(a, b)
        self.Courseware_but3.setMaximumSize(a, b)
        self.Courseware_but4.setMaximumSize(a, b)
        self.Courseware_but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but1.clicked.connect(self.return_fun)
        self.Courseware_but2.clicked.connect(self.select_fun1)
        self.Courseware_but3.clicked.connect(self.select_fun2)
        self.Courseware_but4.clicked.connect(self.select_fun3)
        self.layout.addWidget(self.Courseware_but1, 0,1)
        self.layout.addWidget(self.Courseware_but2,  0,2)
        self.layout.addWidget(self.Courseware_but3,  0,3)
        self.layout.addWidget(self.Courseware_but4,  0,4)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_window1())
        
    def select_fun1(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_Courseware_child1())
    
    def select_fun2(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_Courseware_child2())
    
    def select_fun3(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_Courseware_child3())
    

#爬虫课件小学
class Controller_Courseware_child1(QFrame):
    def __init__(self):
        super(Controller_Courseware_child1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Courseware_child1but1 = QPushButton("返回")
        self.Courseware_child1but2 = QPushButton("语文")
        self.Courseware_child1but3 = QPushButton("数学")
        self.Courseware_child1but4 = QPushButton("英语")
        self.window1tree = QTextEdit()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.Courseware_child1but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Courseware_child1but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child1but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child1but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child1but1.clicked.connect(self.return_fun)
        self.Courseware_child1but2.clicked.connect(self.select_fun1)
        self.Courseware_child1but3.clicked.connect(self.select_fun2)
        self.Courseware_child1but4.clicked.connect(self.select_fun3)
        self.win_layout1.addWidget(self.Courseware_child1but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child1but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child1but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child1but4, 4, 0,1,2)
        self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Courseware())

    
    def select_fun1(self):
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("小学语文 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj1/", "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj2/",
                    "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj3/", "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj4/",
                    "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj5/", "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj6/",
                    "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj7/", "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj8/",
                    "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj9/", "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj10/",
                    "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj11/", "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/sheji_1/kj12/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    n= 64
                    if html[n]=='/':
                        n=65
                    data = Reptile().crawling_data('课件', '小学', '语文',5, n, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)
        self.open_but()
        
    def select_fun2(self):
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("小学数学 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/xxsx/jszx/tbjxzy/xsjxkj/"
        data = Reptile().crawling_url(url, 1)
        self.window1tree.append(data)
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '小学', '数学',4, 64, html )
                    self.window1tree.append(data)
                    x=x+1
                    time.sleep(1)        
        self.open_but()
        
    def select_fun3(self):
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("小学英语 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/xe/jszx/tbjxzy/kjsc/PEPkjsc/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data)
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '小学', '英语',6, 67, html )
                    self.window1tree.append(data)
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)     
        self.open_but()
        
    def close_but(self):
        self.Courseware_child1but1.setEnabled(False)
        self.Courseware_child1but2.setEnabled(False)
        self.Courseware_child1but3.setEnabled(False)
        self.Courseware_child1but4.setEnabled(False)
        
    
    def open_but(self):
        self.Courseware_child1but1.setEnabled(True)
        self.Courseware_child1but2.setEnabled(True)
        self.Courseware_child1but3.setEnabled(True)
        self.Courseware_child1but4.setEnabled(True)
        
    
    
#爬虫课件初中
class Controller_Courseware_child2(QFrame):
    def __init__(self):
        super(Controller_Courseware_child2, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Courseware_child2but1 = QPushButton("返回")
        self.Courseware_child2but2 = QPushButton("语文")
        self.Courseware_child2but3 = QPushButton("数学")
        self.Courseware_child2but4 = QPushButton("英语")
        self.Courseware_child2but5 = QPushButton("物理")
        self.Courseware_child2but6 = QPushButton("化学")
        self.Courseware_child2but7 = QPushButton("生物")
        self.Courseware_child2but8 = QPushButton("政治")
        self.Courseware_child2but9 = QPushButton("历史")
        self.Courseware_child2but10 = QPushButton("地理")
        self.window1tree = QTextEdit()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.Courseware_child2but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child2but1.clicked.connect(self.return_fun)
        self.Courseware_child2but2.clicked.connect(self.select_fun1)
        self.Courseware_child2but3.clicked.connect(self.select_fun2)
        self.Courseware_child2but4.clicked.connect(self.select_fun3)
        self.Courseware_child2but5.clicked.connect(self.select_fun4)
        self.Courseware_child2but6.clicked.connect(self.select_fun5)
        self.Courseware_child2but7.clicked.connect(self.select_fun6)
        self.Courseware_child2but8.clicked.connect(self.select_fun7)
        self.Courseware_child2but9.clicked.connect(self.select_fun8)
        self.Courseware_child2but10.clicked.connect(self.select_fun9)
        self.win_layout1.addWidget(self.Courseware_child2but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child2but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
        
        
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Courseware())
    def select_fun1(self):   #语文
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中语文 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/czyw/jszx/tbjxzy/kjzy/qs/", "http://old.pep.com.cn/czyw/jszx/tbjxzy/kjzy/qx_1/", 
                    "http://old.pep.com.cn/czyw/jszx/tbjxzy/kjzy/bs_1/", "http://old.pep.com.cn/czyw/jszx/tbjxzy/kjzy/bx_1/", 
                    "http://old.pep.com.cn/czyw/jszx/tbjxzy/kjzy/js_1/", "http://old.pep.com.cn/czyw/jszx/tbjxzy/kjzy/jx_1/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '语文',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)    
        self.open_but()
        
    def select_fun2(self):   #数学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中数学 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/czsx/jszx/czsxtbjxzy/czsxdzkb_1_1_2/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '数学',5, -20, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1 
                    time.sleep(1)     
        self.open_but()
        
    def select_fun3(self):   #英语
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中英语 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/ce/czyy/tbjxzy/kjsc/7s_1_1_1_1/", "http://old.pep.com.cn/ce/czyy/tbjxzy/kjsc/7s_1_1_1/", 
                    "http://old.pep.com.cn/ce/czyy/tbjxzy/kjsc/7s_1_1/", "http://old.pep.com.cn/ce/czyy/tbjxzy/kjsc/7s_1/",
                    "http://old.pep.com.cn/ce/czyy/tbjxzy/kjsc/7s/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '英语',5, -20, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)     
        self.open_but()
        
    def select_fun4(self):   #物理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中物理 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/czwl/jszx/tbjx/kj/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '物理',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)           
        self.open_but()
        
    def select_fun5(self):   #化学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中化学 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/czhx/jshzhx/tbxzy/tblxi_1/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '化学',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)        
        self.open_but()
        
    def select_fun6(self):   #生物
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中生物 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/czsw/jshzhx/tbjxzy/kjzy/7sh/", "http://old.pep.com.cn/czsw/jshzhx/tbjxzy/kjzy/7xxsh/", 
                    "http://old.pep.com.cn/czsw/jshzhx/tbjxzy/kjzy/8sh/", "http://old.pep.com.cn/czsw/jshzhx/tbjxzy/kjzy/8xxsh/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '生物',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)  
        self.open_but()
        
    def select_fun7(self):   #政治
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中政治 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/sxpd/js/tbjx/kj/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '化学',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1 
                    time.sleep(1)     
        self.open_but()
        
    def select_fun8(self):   #历史
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中历史 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/czls/js/tbjx/kj/7s/", "http://old.pep.com.cn/czls/js/tbjx/kj/7x/", 
                   "http://old.pep.com.cn/czls/js/tbjx/kj/8s/", 'http://old.pep.com.cn/czls/js/tbjx/kj/8x/'
                   "http://old.pep.com.cn/czls/js/tbjx/kj/9s/", 'http://old.pep.com.cn/czls/js/tbjx/kj/9x/']
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '历史',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)     
        self.open_but()
        
    def select_fun9(self):   #地理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("初中地理 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/czdl/jszx/tbjxzy/kj/qs_1/", "http://old.pep.com.cn/czdl/jszx/tbjxzy/kj/qx_1/", 
                   "http://old.pep.com.cn/czdl/jszx/tbjxzy/kj/bs_1/", "http://old.pep.com.cn/czdl/jszx/tbjxzy/kj/bx_1/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '初中', '地理',5, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)     
        self.open_but()
        
    def close_but(self):
        self.Courseware_child2but1.setEnabled(False)
        self.Courseware_child2but2.setEnabled(False)
        self.Courseware_child2but3.setEnabled(False)
        self.Courseware_child2but4.setEnabled(False)
        self.Courseware_child2but5.setEnabled(False)
        self.Courseware_child2but6.setEnabled(False)
        self.Courseware_child2but7.setEnabled(False)
        self.Courseware_child2but8.setEnabled(False)
        self.Courseware_child2but9.setEnabled(False)
        self.Courseware_child2but10.setEnabled(False)
    
    def open_but(self):
        self.Courseware_child2but1.setEnabled(True)
        self.Courseware_child2but2.setEnabled(True)
        self.Courseware_child2but3.setEnabled(True)
        self.Courseware_child2but4.setEnabled(True)
        self.Courseware_child2but5.setEnabled(True)
        self.Courseware_child2but6.setEnabled(True)
        self.Courseware_child2but7.setEnabled(True)
        self.Courseware_child2but8.setEnabled(True)
        self.Courseware_child2but9.setEnabled(True)
        self.Courseware_child2but10.setEnabled(True)
    
#爬虫课件高中
class Controller_Courseware_child3(QFrame):
    def __init__(self):
        super(Controller_Courseware_child3, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Courseware_child3but1 = QPushButton("返回")
        self.Courseware_child3but2 = QPushButton("语文")
        self.Courseware_child3but3 = QPushButton("数学")
        self.Courseware_child3but4 = QPushButton("英语")
        self.Courseware_child3but5 = QPushButton("物理")
        self.Courseware_child3but6 = QPushButton("化学")
        self.Courseware_child3but7 = QPushButton("生物")
        self.Courseware_child3but8 = QPushButton("政治")
        self.Courseware_child3but9 = QPushButton("历史")
        self.Courseware_child3but10 = QPushButton("地理")
        self.window1tree = QTextEdit()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.Courseware_child3but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Courseware_child3but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_child3but1.clicked.connect(self.return_fun)
        self.Courseware_child3but2.clicked.connect(self.select_fun1)
        self.Courseware_child3but3.clicked.connect(self.select_fun2)
        self.Courseware_child3but4.clicked.connect(self.select_fun3)
        self.Courseware_child3but5.clicked.connect(self.select_fun4)
        self.Courseware_child3but6.clicked.connect(self.select_fun5)
        self.Courseware_child3but7.clicked.connect(self.select_fun6)
        self.Courseware_child3but8.clicked.connect(self.select_fun7)
        self.Courseware_child3but9.clicked.connect(self.select_fun8)
        self.Courseware_child3but10.clicked.connect(self.select_fun9)
        self.win_layout1.addWidget(self.Courseware_child3but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.Courseware_child3but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
        

    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Courseware())
    def select_fun1(self):  #语文
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中语文 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/gzyw/jszx/tbjxzy/kbjc/kjzy/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '语文',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)    
        self.open_but()
        
    def select_fun2(self):  #数学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中数学 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/gzsx/jszx_1/czsxtbjxzy/xkbsyjc/jxkj/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '数学',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1 
                    time.sleep(1)   
        self.open_but()
    def select_fun3(self):  #英语
        QMessageBox.question(self, "提示!", "暂时没有英语的数据可供爬取\n请您爬取其他内容！" , QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
    
    def select_fun4(self):  #物理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中物理 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/gzwl/jszx/tbjx/kb/kj/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in self.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '物理',6, -20, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)      
        self.open_but()
        
    def select_fun5(self):  #化学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中化学 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/nmby/", "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/xcwaz/", 
                    "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/oplkm/", "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/iklpk/", 
                    "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/oplnm1/", "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/opujm/", 
                    "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/12wa/", "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/kjsc1234/xcz/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '化学',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1)     
        self.open_but()
        
    def select_fun6(self):  #生物
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中生物 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/gzsw/jshzhx/tbziy/kbshy/jxsj_1/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '生物',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1 
                    time.sleep(1)    
        self.open_but()
        
    def select_fun7(self):  #政治
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中政治 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/sxzz/js/tbjx/kb/kj/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '政治',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1) 
        self.open_but()
        
    def select_fun8(self):  #历史
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中历史 \n\n数据爬取如下:")
        QApplication.processEvents()
        url = "http://old.pep.com.cn/gzls/js/tbjx/kb/kj/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '历史',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1) 
        self.open_but()
        
    def select_fun9(self):  #地理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("高中地理 \n\n数据爬取如下:")
        QApplication.processEvents()
        urls = ["http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/kjzy/bx1/", 
                    "http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/kjzy/bx2/", 
                    "http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/kjzy/bx3/"]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 " + str(x) + " 次") 
                    QApplication.processEvents()
                    data = Reptile().crawling_data('课件', '高中', '地理',6, -21, html )
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(1) 
        self.open_but()
        
    def close_but(self):
        self.Courseware_child3but1.setEnabled(False)
        self.Courseware_child3but2.setEnabled(False)
        self.Courseware_child3but3.setEnabled(False)
        self.Courseware_child3but4.setEnabled(False)
        self.Courseware_child3but5.setEnabled(False)
        self.Courseware_child3but6.setEnabled(False)
        self.Courseware_child3but7.setEnabled(False)
        self.Courseware_child3but8.setEnabled(False)
        self.Courseware_child3but9.setEnabled(False)
        self.Courseware_child3but10.setEnabled(False)
    
    def open_but(self):
        self.Courseware_child3but1.setEnabled(True)
        self.Courseware_child3but2.setEnabled(True)
        self.Courseware_child3but3.setEnabled(True)
        self.Courseware_child3but4.setEnabled(True)
        self.Courseware_child3but5.setEnabled(True)
        self.Courseware_child3but6.setEnabled(True)
        self.Courseware_child3but7.setEnabled(True)
        self.Courseware_child3but8.setEnabled(True)
        self.Courseware_child3but9.setEnabled(True)
        self.Courseware_child3but10.setEnabled(True)
    
  
#爬虫练习
class Controller_win_Practice(QFrame):
    def __init__(self):
        super(Controller_win_Practice, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_but1 = QPushButton("返回")  #爬取课件中的控件
        self.Practice_but2 = QPushButton("小学")
        self.Practice_but3 = QPushButton("初中")
        self.Practice_but4 = QPushButton("高中")
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        a = self.x *1.0/8
        b = self.y*1.0/3
        self.Practice_but1.setMaximumSize(a, b)
        self.Practice_but2.setMaximumSize(a, b)
        self.Practice_but3.setMaximumSize(a, b)
        self.Practice_but4.setMaximumSize(a, b)
        self.Practice_but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_but1.clicked.connect(self.return_fun)
        self.Practice_but2.clicked.connect(self.select_fun1)
        self.Practice_but3.clicked.connect(self.select_fun2)
        self.Practice_but4.clicked.connect(self.select_fun3)
        self.layout.addWidget(self.Practice_but1, 0,1)
        self.layout.addWidget(self.Practice_but2,  0,2)
        self.layout.addWidget(self.Practice_but3,  0,3)
        self.layout.addWidget(self.Practice_but4,  0,4)
        



    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_window1())
        
    def select_fun1(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_Practice_child1())
    
    def select_fun2(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_Practice_child2())
    
    def select_fun3(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_Practice_child3())
    

#爬虫练习小学
class Controller_Practice_child1(QFrame):
    def __init__(self):
        super(Controller_Practice_child1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_child1but1 = QPushButton("返回")
        self.Practice_child1but2 = QPushButton("语文")
        self.Practice_child1but3 = QPushButton("数学")
        self.Practice_child1but4 = QPushButton("英语")
        self.window1tree = QTextEdit()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.Practice_child1but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Practice_child1but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child1but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child1but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child1but1.clicked.connect(self.return_fun)
        self.Practice_child1but2.clicked.connect(self.select_fun1)
        self.Practice_child1but3.clicked.connect(self.select_fun2)
        self.Practice_child1but4.clicked.connect(self.select_fun3)
        self.win_layout1.addWidget(self.Practice_child1but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child1but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child1but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child1but4, 4, 0,1,2)
        self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)

    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Practice())
    def select_fun1(self):
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  小学语文 \n数据爬取如下:")
        url = "http://old.pep.com.cn/xiaoyu/jiaoshi/tbjx/yuedu_1/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '小学', '语文',5, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)
        self.open_but()
        
    def select_fun2(self):
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  小学数学 \n数据爬取如下:")
        url = "http://old.pep.com.cn/xxsx/xxsxxs/xstblx/"
        data = Reptile().crawling_url2(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '小学', '数学',4, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)
        self.open_but()
    
    def select_fun3(self):
        QMessageBox.question(self, "提示!", "暂时没有该科目的数据可爬取", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        
    def close_but(self):
        self.Practice_child1but1.setEnabled(False)
        self.Practice_child1but2.setEnabled(False)
        self.Practice_child1but3.setEnabled(False)
        self.Practice_child1but4.setEnabled(False)
    
    def open_but(self):
        self.Practice_child1but1.setEnabled(True)
        self.Practice_child1but2.setEnabled(True)
        self.Practice_child1but3.setEnabled(True)
        self.Practice_child1but4.setEnabled(True)
    

#爬虫练习初中
class Controller_Practice_child2(QFrame):
    def __init__(self):
        super(Controller_Practice_child2, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_child2but1 = QPushButton("返回")
        self.Practice_child2but2 = QPushButton("语文")
        self.Practice_child2but3 = QPushButton("数学")
        self.Practice_child2but4 = QPushButton("英语")
        self.Practice_child2but5 = QPushButton("物理")
        self.Practice_child2but6 = QPushButton("化学")
        self.Practice_child2but7 = QPushButton("生物")
        self.Practice_child2but8 = QPushButton("政治")
        self.Practice_child2but9 = QPushButton("历史")
        self.Practice_child2but10 = QPushButton("地理")
        self.window1tree = QTextEdit()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.Practice_child2but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Practice_child2but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but1.clicked.connect(self.return_fun)
        self.Practice_child2but2.clicked.connect(self.select_fun1)
        self.Practice_child2but3.clicked.connect(self.select_fun2)
        self.Practice_child2but4.clicked.connect(self.select_fun3)
        self.Practice_child2but5.clicked.connect(self.select_fun4)
        self.Practice_child2but6.clicked.connect(self.select_fun5)
        self.Practice_child2but7.clicked.connect(self.select_fun6)
        self.Practice_child2but8.clicked.connect(self.select_fun7)
        self.Practice_child2but9.clicked.connect(self.select_fun8)
        self.Practice_child2but10.clicked.connect(self.select_fun9)
        self.win_layout1.addWidget(self.Practice_child2but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child2but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
        
        
        
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Practice())
    def select_fun1(self): #语文
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中语文 \n数据爬取如下:")
        urls = ["http://old.pep.com.cn/czyw/xszx/tbxx/ywlx/qs/", "http://old.pep.com.cn/czyw/xszx/tbxx/ywlx/qx/", 
                    "http://old.pep.com.cn/czyw/xszx/tbxx/ywlx/bs/", "http://old.pep.com.cn/czyw/xszx/tbxx/ywlx/bx/", 
                    "http://old.pep.com.cn/czyw/xszx/tbxx/ywlx/js/", "http://old.pep.com.cn/czyw/xszx/tbxx/ywlx/jx/"]
        for url in urls:        
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '语文',5, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)
        self.open_but()
    
    def select_fun2(self): #数学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中数学 \n数据爬取如下:")
        url = "http://old.pep.com.cn/czsx/jszx/czsxtbjxzy/czsxdzkb_1_1_2_1/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '数学',4, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)
        self.open_but()
        
    def select_fun3(self): #英语
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中英语 \n数据爬取如下:")
        urls = ["http://old.pep.com.cn/ce/cexsml/tbzy/sj_1/7s_1_1_1_1/", "http://old.pep.com.cn/ce/cexsml/tbzy/sj_1/7s_1_1_1/", 
                    "http://old.pep.com.cn/ce/cexsml/tbzy/sj_1/7s_1_1/", "http://old.pep.com.cn/ce/cexsml/tbzy/sj_1/7s_1/", 
                    "http://old.pep.com.cn/ce/cexsml/tbzy/sj_1/7s/"]
        for url in urls:        
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '英语',5, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)
        self.open_but()
        
    def select_fun4(self): #物理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中物理 \n数据爬取如下:")
        url = "http://old.pep.com.cn/czwl/xszx/tbxx/tb8s/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '物理',4, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)   
        self.open_but()
        
    def select_fun5(self): #化学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中化学 \n数据爬取如下:")
        url = "http://old.pep.com.cn/czhx/xshzx/xsjxt/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '化学',4, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2) 
        self.open_but()
    
    def select_fun6(self): #生物
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中生物 \n数据爬取如下:")
        url = "http://old.pep.com.cn/czhx/xshzx/xsjxt/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '生物',5, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)   
        self.open_but()
        
    def select_fun7(self): #政治
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中政治 \n数据爬取如下:")
        url = "http://old.pep.com.cn/czhx/xshzx/xsjxt/"
        data = Reptile().crawling_url2(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '政治',4, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)     
        self.open_but()
        
    def select_fun8(self): #历史
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中历史 \n数据爬取如下:")
        url = "http://old.pep.com.cn/czls/xs/tbxx/cs/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 10):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '历史',5, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents() 
                    x=x+1
                    time.sleep(2)    
        self.open_but()
        
    def select_fun9(self): #地理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  初中地理 \n数据爬取如下:")
        urls = ["http://old.pep.com.cn/czdl/jszx/tbjxzy/st/qs_1/", "http://old.pep.com.cn/czdl/jszx/tbjxzy/st/qx_1/", 
                    "http://old.pep.com.cn/czdl/jszx/tbjxzy/st/bs_1/", "http://old.pep.com.cn/czdl/jszx/tbjxzy/st/bx_1/"]
        for url in urls:        
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents() 
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '初中', '地理',5, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents() 
                    x=x+1
                    time.sleep(2)            
        self.open_but()
        
    def close_but(self):
        self.Practice_child2but1.setEnabled(False)
        self.Practice_child2but2.setEnabled(False)
        self.Practice_child2but3.setEnabled(False)
        self.Practice_child2but4.setEnabled(False)
        self.Practice_child2but5.setEnabled(False)
        self.Practice_child2but6.setEnabled(False)
        self.Practice_child2but7.setEnabled(False)
        self.Practice_child2but8.setEnabled(False)
        self.Practice_child2but9.setEnabled(False)
        self.Practice_child2but10.setEnabled(False)
    
    def open_but(self):
        self.Practice_child2but1.setEnabled(True)
        self.Practice_child2but2.setEnabled(True)
        self.Practice_child2but3.setEnabled(True)
        self.Practice_child2but4.setEnabled(True)
        self.Practice_child2but5.setEnabled(True)
        self.Practice_child2but6.setEnabled(True)
        self.Practice_child2but7.setEnabled(True)
        self.Practice_child2but8.setEnabled(True)
        self.Practice_child2but9.setEnabled(True)
        self.Practice_child2but10.setEnabled(True)
    
#爬虫练习高中
class Controller_Practice_child3(QFrame):
    def __init__(self):
        super(Controller_Practice_child3, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_child3but1 = QPushButton("返回")
        self.Practice_child3but2 = QPushButton("语文")
        self.Practice_child3but3 = QPushButton("数学")
        self.Practice_child3but4 = QPushButton("英语")
        self.Practice_child3but5 = QPushButton("物理")
        self.Practice_child3but6 = QPushButton("化学")
        self.Practice_child3but7 = QPushButton("生物")
        self.Practice_child3but8 = QPushButton("政治")
        self.Practice_child3but9 = QPushButton("历史")
        self.Practice_child3but10 = QPushButton("地理")
        self.window1tree = QTextEdit()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.Practice_child3but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Practice_child3but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but1.clicked.connect(self.return_fun)
        self.Practice_child3but2.clicked.connect(self.select_fun1)
        self.Practice_child3but3.clicked.connect(self.select_fun2)
        self.Practice_child3but4.clicked.connect(self.select_fun3)
        self.Practice_child3but5.clicked.connect(self.select_fun4)
        self.Practice_child3but6.clicked.connect(self.select_fun5)
        self.Practice_child3but7.clicked.connect(self.select_fun6)
        self.Practice_child3but8.clicked.connect(self.select_fun7)
        self.Practice_child3but9.clicked.connect(self.select_fun8)
        self.Practice_child3but10.clicked.connect(self.select_fun9)
        self.win_layout1.addWidget(self.Practice_child3but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.Practice_child3but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
        
        
        




    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_win_Practice())
    def select_fun1(self): #语文
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中语文 \n数据爬取如下:")
        url = "http://old.pep.com.cn/gzyw/jszx/tbjxzy/kbjc/lx/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents() 
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '语文',6, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents() 
                    x=x+1
                    time.sleep(2)           
        self.open_but()
        
    def select_fun2(self): #数学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中数学 \n数据爬取如下:")
        url = "http://old.pep.com.cn/gzsx/jszx_1/czsxtbjxzy/xkbsyjc/st/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents() 
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '数学',6, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents() 
                    x=x+1
                    time.sleep(2)       
        self.open_but()
        
    def select_fun3(self): #英语
        QMessageBox.question(self, "提示!", "暂时没有英语的数据可供爬取\n请您爬取其他内容！", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        
    def select_fun4(self): #物理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中物理 \n数据爬取如下:")
        url = "http://old.pep.com.cn/gzwl/xszx/tbxx/st/"
        data = Reptile().crawling_url(url, 0)
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '物理',5, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)  
        self.open_but()
    def select_fun5(self): #化学
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中化学 \n数据爬取如下:")
        urls = ["http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/tbxl/bxytbxl/",
                    "http://old.pep.com.cn/gzhx/gzhxjs/0pl/kb/tbxl/bxetbxl/" ]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '化学',6, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)   
        self.open_but()
        
    def select_fun6(self): #生物
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中生物 \n数据爬取如下:")
        url = "http://old.pep.com.cn/gzsw/xszx/tbxxi/jxsj_2_1/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '生物',5, -20, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)       
        self.open_but()
        
    def select_fun7(self): #政治
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中政治 \n数据爬取如下:")
        url = "http://old.pep.com.cn/sxzz/js/gkzl/index.htm"
        data = Reptile().crawling_url2(url, 1)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '政治',4, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)      
        self.open_but()
        
    def select_fun8(self): #历史
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中历史 \n数据爬取如下:")
        url = "http://old.pep.com.cn/gzls/js/tbjx/kb/tbdl/"
        data = Reptile().crawling_url(url, 0)
        self.window1tree.append(data) 
        QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '历史',6, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)        
        self.open_but()
        
    def select_fun9(self): #地理
        rely = QMessageBox.question(self, "提示!", "爬取过程需要时间比较久\n请问是否继续？", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        self.close_but()
        self.window1tree.setPlainText("爬取课件数据开始\n爬取过程需要时间比较久，请您耐心等待!!\n\n")                          
        x=1
        self.window1tree.append("  高中地理  \n数据爬取如下:")
        urls = ["http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/dlst/bx1/","http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/dlst/bx3/", 
                    "http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/dlst/bx2/", "http://old.pep.com.cn/gzdl/jszx/tbjxzy/kbjc/dlst/xxjc/" ]
        for url in urls:
            data = Reptile().crawling_url(url, 0)
            self.window1tree.append(data) 
            QApplication.processEvents()
        for i in range(0, 3):
            for html in win.htmls:
                if Reptile().check_url(html):
                    self.window1tree.append("第 "+ str(x) + " 次\n")
                    data = Reptile().data_word('练习', '高中', '地理',6, -21, html)
                    self.window1tree.append(data) 
                    QApplication.processEvents()
                    x=x+1
                    time.sleep(2)  
        self.open_but()
 
    def close_but(self):
        self.Practice_child3but1.setEnabled(False)
        self.Practice_child3but2.setEnabled(False)
        self.Practice_child3but3.setEnabled(False)
        self.Practice_child3but4.setEnabled(False)
        self.Practice_child3but5.setEnabled(False)
        self.Practice_child3but6.setEnabled(False)
        self.Practice_child3but7.setEnabled(False)
        self.Practice_child3but8.setEnabled(False)
        self.Practice_child3but9.setEnabled(False)
        self.Practice_child3but10.setEnabled(False)
    
    def open_but(self):
        self.Practice_child3but1.setEnabled(True)
        self.Practice_child3but2.setEnabled(True)
        self.Practice_child3but3.setEnabled(True)
        self.Practice_child3but4.setEnabled(True)
        self.Practice_child3but5.setEnabled(True)
        self.Practice_child3but6.setEnabled(True)
        self.Practice_child3but7.setEnabled(True)
        self.Practice_child3but8.setEnabled(True)
        self.Practice_child3but9.setEnabled(True)
        self.Practice_child3but10.setEnabled(True)
    
 
class Reptile():
    def __init__(self):
        super(Reptile, self).__init__()
        self.file_sort = ['一年级上册', '二年级上册', '三年级上册', '四年级上册', '五年级上册', '六年级上册', 
                                '一年级下册', '二年级下册', '三年级下册', '四年级下册', '五年级下册', '六年级下册', 
                                '初一上册', '初二上册', '初三上册', '初一下册', '初二下册', '初三下册']
        self.file_sort1 = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级', 
                                 '一年级', '二年级', '三年级', '四年级', '五年级', '六年级', 
                               '初一', '初二', '初三', '初一', '初二', '初三']
        
    
    def get_agent(self):  #模拟浏览器
        '''
        模拟header的user-agent字段，
        返回一个随机的user-agent字典类型的键值对
        '''
        agents = ['Mozilla/5.0 (compatible; MSIE 9.0; Windows NT 6.1; Trident/5.0;',
              'Mozilla/5.0 (Macintosh; Intel Mac OS X 10.6; rv,2.0.1) Gecko/20100101 Firefox/4.0.1',
              'Opera/9.80 (Macintosh; Intel Mac OS X 10.6.8; U; en) Presto/2.8.131 Version/11.11',
              'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_7_0) AppleWebKit/535.11 (KHTML, like Gecko) Chrome/17.0.963.56 Safari/535.11',
              'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 5.1; 360SE)']
        fakeheader = {}
        fakeheader['User-agent'] = agents[random.randint(0, len(agents)-1)]
        return fakeheader    
    
    def check_url(self, url):  #检查网页是否已经爬取过
        r = requests.get(url,timeout=10,headers=self.get_agent())
        r = r.text
        data = len(r)
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        d=conn.cursor()
        d.execute("select * from successfulurl")
        for variate in d.fetchall():
            if variate[0]==url and variate[1]== data:
                return False
        d.close()
        conn.close()
        return True
    
    def crawling_url(self, url, sign):     #爬取网页网址
        for j in range(0,10):        #使用循环，避免爬取网址时出错，无法进行下面的爬取。
            try:
                content = requests.get(url,timeout=10,headers=self.get_agent())
                content.encoding = content.apparent_encoding
                soup = BeautifulSoup(content.text, 'lxml')
                soups = soup.find_all('div',attrs={'class': 'clear'})
                for soup in soups:
                    datas = soup.find_all('a')
                    for data in datas: 
                        #对于一些网址可以这样处理，将多余的字符从网址中出去。
                        data = data['href'].replace('\n', '').replace('.../', '').replace('../', '').replace('./', '')
                        if sign:
                            data = url[:-7]+ data
                        else:
                            data = url + data    
                        win.htmls.append(data)
                win.htmls=list(set(win.htmls))  #去重复元素
                return("爬取成功\n\n")
            except:
                return("爬取错误")
    
    def crawling_url2(self, url, sign):  #爬取网页网址
        for j in range(0,10):        #使用循环，避免爬取网址时出错，无法进行下面的爬取。
            try:
                content = requests.get(url,timeout=10,headers=self.get_agent())
                content.encoding = content.apparent_encoding
                soup = BeautifulSoup(content.text, 'lxml')
                soups = soup.find_all('div',attrs={'class': 'ttlist'})
                for soup in soups:
                    datas = soup.find_all('a')
                    for data in datas: 
                        #对于一些网址可以这样处理，将多余的字符从网址中出去。
                        data = data['href'].replace('.../', '').replace('../', '').replace('./', '')
                        if sign:
                            data = url[:-9]+ data
                        else:
                            data = url + data    
                        win.htmls.append(data)
                win.htmls=list(set(win.htmls))  #去重复元素
                return("爬取成功\n\n")
            except:
                return("爬取错误\n\n")
    
    def file_to_zip(self, path):   #将文件夹压缩为压缩包。
        filepath = path +'.zip'
        if os.path.exists(filepath):
            os.remove(filepath)
        z = zipfile.ZipFile(filepath,'w',zipfile.ZIP_DEFLATED)
        for dirpath, dirnames, filenames in os.walk(path):
            fpath = dirpath.replace(path,'')
            fpath = fpath and fpath + os.sep or ''
            for filename in filenames:
                z.write(os.path.join(dirpath, filename),fpath+filename)
        z.close()
       
    
    def zip_to_files(self, zippath):  #将压缩包解压
        path = zippath[:-4]
        if (os.path.isdir(path)):  #判断文件夹是否存在
            fileNames = glob.glob(path + r'\*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(path)
        zf = zipfile.ZipFile(zippath)
        for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
            right_fn = fn.replace('\\\\', '_') .replace('\\', '_').replace('//', '_').replace('/', '_') # 将文件名正确编码   
            right_fn =  path +'/'+ right_fn 
            with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                with zf.open(fn, 'r') as origin_file:  # 打开原文件
                    shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
        zf.close()
        os.remove(zippath)
        
    
    def ppt_to_image(self, output_file, input_file):  #将ppt 转换为图片
        fileNames = glob.glob(output_file + r'\*')   
        if fileNames:
            for fileName in fileNames:     #将pa 文件夹中的文件删除。
                os.remove( fileName)
        #powerpoint = comtypes.client.CreateObject("kwpp.Application") #使用wps的接口
        powerpoint = client.DispatchEx("kwpp.Application") #使用wps的接口
        powerpoint.Visible = 1 #0 不显示wps程序,1 显示程序  ->不显示程序无法将ppt转化为图片。
        powerpoint.DisplayAlerts = 0  #不警告
        ppt = powerpoint.Presentations.Open(input_file)
        ppt.SaveAs(output_file + '.jpg', 17)    # 另存为
        ppt.Close()   # 退出
        os.remove(input_file)
        powerpoint.Quit()
        
        

    #爬取课件数据
    def crawling_data(self, file1, file2, file4, a, b,  url): #a 代表爬取文件名的位置，b代表组成网址时需要减少的字符串长度
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c = conn.cursor()
        c.execute("select * from Education")
        no = len(c.fetchall())
        try: 
            r = requests.get(url,timeout=10,headers=self.get_agent())
            data = len(r.text)
            r.encoding = r.apparent_encoding
            soup = BeautifulSoup(r.text, 'lxml')
            da = soup.find_all('a', target="_self")
            file3 = da[a].text
            for i in range(0, len(self.file_sort)):
                if file3 == self.file_sort[i]:
                    file3 = self.file_sort1[i]
                    break
            title = soup.find('title').text.replace('\n', '').replace('/', '_')
            try:
                soup1 = soup.find('div',id="downloadcontent")
                url2 = soup1.find('a').get('href').replace('.../', '').replace('../', '').replace('./', '')
                print(url2)
                url2 = url[0:b]+url2
                print(url2)
            except:
                soup1 = soup.find('div',id="doccontent")
                url2 = soup1.find('A').get('href')
            filename = url2[-3:]
            file =  'D:/项目数据库/wen/xinwen.' +filename
            d = requests.get(url2,timeout=10,headers=self.get_agent())
            with open(file,'wb')as f:  #将网上的文件下载保存进电脑。
                for chunk in d.iter_content(chunk_size=100):
                    f.write(chunk) 
            f.close()
            if filename == "ppt":
                pa = "D:/项目数据库/tupian"  #保存图片的路径。
                self.ppt_to_image(pa, file)  #将ppt 转换为图片。
                self.file_to_zip(pa)  #将文件夹压缩为压缩包。
                file = pa + ".zip"
                with open(file, "rb") as f:
                    total=base64.b64encode(f.read())   #将文件转换为字节。
                f.close()
                filename1 = 'zip'
                conn.execute("INSERT INTO Education(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename1))
                conn.commit()
                conn.execute("insert into filedata(no,total) values(?,?)", ("S"+str(no+1), total))
                conn.commit()
            elif filename == 'zip':
                self.zip_to_files(file)  #将压缩包解压。
                pa1 = 'D:/项目数据库/wen/xinwen'  #解压后的文件名
                fileNames = glob.glob(pa1 + r'\*')  #读取解压文件夹里的文件。
                for fileName in fileNames:
                    da = fileName[-3:]
                    if da == 'ppt' or da =='ptx' :
                        pa = "D:/项目数据库/tupian"  #保存图片的路径
                        self.ppt_to_image(pa, fileName) #将ppt转换为图片。
                        self.file_to_zip(pa)  #将文件夹压缩为压缩包。
                        file = pa + ".zip"   #压缩包的路径
                        with open(file, "rb") as f:
                            total=base64.b64encode(f.read())   #将文件转换为字节。
                        f.close()
                        filename1 = "zip"
                        conn.execute("INSERT INTO Education(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename1))
                        conn.commit()
                        conn.execute("insert into filedata(no,total) values(?,?)", ("S"+str(no+1), total))
                        conn.commit()
                        break
                    elif da == 'swf':
                        file = fileName  #不是ppt时，文件的路径赋值给file,保证文件转为字节时正确。
                        with open(file, "rb") as f:
                            total=base64.b64encode(f.read())   #将文件转换为字节。
                        f.close()
                        filename1 = da
                        conn.execute("INSERT INTO Education(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename1))
                        conn.commit()
                        conn.execute("insert into filedata(no,total) values(?,?)", ("S"+str(no+1), total))
                        conn.commit()
                        break
                    
            elif filename == 'swf':
                with open(file, "rb") as f:
                    total=base64.b64encode(f.read())   #将文件转换为字节。
                f.close()
                filename1=filename
                conn.execute("INSERT INTO Education(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename1))
                conn.commit()
                conn.execute("insert into filedata(no,total) values(?,?)", ("S"+str(no+1), total))
                conn.commit()
                
            os.remove(file)
            conn.execute("insert into successfulurl(url,howbyte)values(?,?)", (url, data))
            conn.commit()
            conn.close()
            return(title + "\n爬取成功\n")
        except:
            try:
                return(title + "\n爬取失败\n")
            except:
                return("爬取错误\n")
            
    
    

    
    #爬取练习课件
    def data_word(self, file1,file2, file4, a, b, url):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c = conn.cursor()
        c.execute("select * from Education")
        no = len(c.fetchall())
        file = r'D:/项目数据库/tupian'
        fileNames = glob.glob(file + r'\*')   
        if fileNames:
            for fileName in fileNames:     #将pa 文件夹中的文件删除。
                os.remove( fileName)
        try:
            content = requests.get(url,headers = self.get_agent())
            data = content.text
            content.encoding = content.apparent_encoding
            soup = BeautifulSoup(content.text, 'lxml')
            title = soup.find('title').text.replace('\n','')
            if title =='404错误页面':
                return('爬取失败\n')
            da = soup.find_all('a', target="_self")
            if len(da)>a:
                file3 = da[a].text
                for i in range(0, len(self.file_sort)):
                    if file3 == self.file_sort[i]:
                        file3 = self.file_sort1[i]
                        break
            else:
                file3 = "附加文件"
            
            QApplication.processEvents()
            #创建空白的word文档
            currentDocument = Document()
            style = currentDocument.styles['Normal']
            font = style.font
            #写入文章标题
#            currentDocument.add_heading(title)
            #找到文章的内容
            childrens = soup.find_all('p')
            if not soup:
                return("爬取失败\n")
            for child in childrens:   
                child = BeautifulSoup(str(child), 'lxml')
                #包含<img>的子节点，在word文档中插入对应的图片
                if child.img:
                    de = child 
                    p = currentDocument.add_paragraph('')
                    for child in de.children:
                        child = BeautifulSoup(str(child), 'lxml')
                        if child.img:
                            pic = 'D:/项目数据库/tupian/temp'
                            pe = child.img['src'].replace('.../', '').replace('../', '').replace('./', '')
                            de = pe[-4:]
                            pic = pic+de
                            ur = url[:b]+ pe 
                            data1 = requests.get(ur,headers=self.get_agent())
                            with open(pic, 'wb') as fp:
                                for chunk in data1.iter_content(chunk_size=100):
                                    fp.write(chunk) 
                            fp.close()  
                            try:
                                run = p.add_run()
                                run.add_picture(pic)
                                os.remove(pic) 
                            except:
                                pass
                        else:
                            para = child.text.replace('\n', '').replace('_', '_') 
                            font.size = Pt(16)
                            font.bold = True     
                            p.add_run(para)
                        
                #包含<tr>的子节点，在word文档中插入表格
                elif child.tr:
                    rows = child.find_all('tr')
                    cols = rows[0].find_all('td')
                    #创建空白表格
                    table = currentDocument.add_table(len(rows), len(cols))
                    #往对应的单元格中写入内容
                    for rindex, row in enumerate(rows):
                        for cindex , col in enumerate(row.find_all('td')):
                            try:
                                cell = table.cell(rindex, cindex)
                                cell.text = col.text
                            except:
                                pass
                #纯文字，直接写入word文件
                elif child.p:
                    para = child.p.text.replace('\n', '').replace('_', '_')
                    p =currentDocument.add_paragraph()
                    font.size = Pt(16)
                    font.bold = True
                    p.add_run(para)
            #保存当前文章的word 文档
            filepath = 'D:/项目数据库/wen/xinwen.docx'
            currentDocument.save(filepath)
            word = client.DispatchEx("kwps.Application")
            word.Visible = 0 #0 不显示wps程序
            word.DisplayAlerts = 0  #不警告
            new_file = 'D:/项目数据库/wen/xinwen.pdf'
            doc = word.Documents.Open(filepath)
            doc.PageSetup.PageWidth = 26*28.35     # 纸张大小, A3=6, A4=7 
            doc.PageSetup.PageHeight = 22*28.35   #1cm = 28.35pt
#            doc.PageSetup.PaperSize = 7     # 纸张大小, A3=6, A4=7 
#            doc.PageSetup.Orientation = 1     # 页面方向, 竖直=0, 水平=1
            doc.SaveAs(filepath)  # 文档保存
            doc.Close(-1)     # doc.Close(-1)保存后关闭，doc.Close()或doc.Close(0)直接关闭不保存
            doc = word.Documents.Open(filepath)
            doc.SaveAs(new_file, FileFormat = 17)
            doc.Close()
            word.Quit()
            os.remove(filepath)
            try:
                pdf = fitz.open(new_file)
                for pg in range(pdf.pageCount):
                    page = pdf[pg]
                    rotate = int(0)
                    # 每个尺寸的缩放系数为2，这将为我们生成分辨率提高4倍的图像。
                    zoom_x = 3
                    zoom_y = 3
                    trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
                    pm = page.getPixmap(matrix=trans, alpha=False)
                    pm.writePNG(file +'/幻灯片' + '%s.jpg' % (pg+1))
                pdf.close()
            except:
                pdf.close()
                doc.close()
            os.remove(new_file)
            self.file_to_zip(file)     
            file = file + '.zip'
            filename = 'zip'
            with open(file, "rb") as f:
                total=base64.b64encode(f.read())   #将文件转换为字节。
            f.close()
            os.remove(file)
            conn.execute("INSERT INTO Education(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename))
            conn.commit()
            conn.execute("insert into filedata(no,total) values(?,?)", ("S"+str(no+1), total))
            conn.commit()
            conn.execute("insert into successfulurl(url,howbyte)values(?,?)", (url, data))
            conn.commit()
            conn.close()
            return(title +'\n爬取成功\n')
        except:
            try:
                return(title +'\n爬取失败\n')
            except:
                return('爬取失败\n')
  




#添加数据的界面
class  addition_Date(QFrame):
    def __init__(self):
        super(addition_Date, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Courseware_but1 = QPushButton("返回") 
        self.Courseware_but2 = QPushButton("一年级")
        self.Courseware_but3 = QPushButton("二年级")
        self.Courseware_but4 = QPushButton("三年级")
        self.Courseware_but5 = QPushButton("四年级")
        self.Courseware_but6 = QPushButton("五年级")
        self.Courseware_but7 = QPushButton("六年级")
        self.Courseware_but8 = QPushButton("初一")
        self.Courseware_but9 = QPushButton("初二")
        self.Courseware_but10 = QPushButton("初三")
        self.Courseware_but11 = QPushButton("高一")
        self.Courseware_but12 = QPushButton("高二")
        self.Courseware_but13 = QPushButton("高三")
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        a = self.x *1.0/8
        b = self.y*1.0/5
        self.Courseware_but1.setMaximumSize(a, b)
        self.Courseware_but2.setMaximumSize(a, b)
        self.Courseware_but3.setMaximumSize(a, b)
        self.Courseware_but4.setMaximumSize(a, b)
        self.Courseware_but5.setMaximumSize(a, b)
        self.Courseware_but6.setMaximumSize(a, b)
        self.Courseware_but7.setMaximumSize(a, b)
        self.Courseware_but8.setMaximumSize(a, b)
        self.Courseware_but9.setMaximumSize(a, b)
        self.Courseware_but10.setMaximumSize(a, b)
        self.Courseware_but11.setMaximumSize(a, b)
        self.Courseware_but12.setMaximumSize(a, b)
        self.Courseware_but13.setMaximumSize(a, b)
        self.Courseware_but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                         
        self.Courseware_but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but11.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but12.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but13.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Courseware_but1.clicked.connect(self.return_fun)
        self.Courseware_but2.clicked.connect(lambda:self.select_fun1(1))
        self.Courseware_but3.clicked.connect(lambda:self.select_fun1(2))
        self.Courseware_but4.clicked.connect(lambda:self.select_fun1(3))
        self.Courseware_but5.clicked.connect(lambda:self.select_fun1(4))
        self.Courseware_but6.clicked.connect(lambda:self.select_fun1(5))
        self.Courseware_but7.clicked.connect(lambda:self.select_fun1(6))
        self.Courseware_but8.clicked.connect(lambda:self.select_fun2(1))
        self.Courseware_but9.clicked.connect(lambda:self.select_fun2(2))
        self.Courseware_but10.clicked.connect(lambda:self.select_fun2(3))
        self.Courseware_but11.clicked.connect(lambda:self.select_fun3(1))
        self.Courseware_but12.clicked.connect(lambda:self.select_fun3(2))
        self.Courseware_but13.clicked.connect(lambda:self.select_fun3(3))
        self.layout.addWidget(self.Courseware_but1, 0,1)
        self.layout.addWidget(self.Courseware_but2,  0,2)
        self.layout.addWidget(self.Courseware_but3,  1,2)
        self.layout.addWidget(self.Courseware_but4,  2,2)
        self.layout.addWidget(self.Courseware_but5,  0,3)
        self.layout.addWidget(self.Courseware_but6,  1,3)
        self.layout.addWidget(self.Courseware_but7 ,  2,3)
        self.layout.addWidget(self.Courseware_but8,  0,4)
        self.layout.addWidget(self.Courseware_but9,  1,4)
        self.layout.addWidget(self.Courseware_but10,  2,4)
        self.layout.addWidget(self.Courseware_but11,  0,5)
        self.layout.addWidget(self.Courseware_but12,  1,5)
        self.layout.addWidget(self.Courseware_but13,  2,5)
        

    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_function())
    
    def select_fun1(self, sign):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date_child1(sign))   
    
    def select_fun2(self, sign):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date_child2(sign))
        
    def select_fun3(self, sign):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date_child3(sign))
  
    
#添加数据小学
class addition_Date_child1(QFrame):
    def __init__(self, sign):
        super(addition_Date_child1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_child1but1 = QPushButton("返回")
        self.Practice_child1but2 = QPushButton("语文")
        self.Practice_child1but3 = QPushButton("数学")
        self.Practice_child1but4 = QPushButton("英语")
        self.date_but1 = QPushButton("返回")
        self.date_but2 = QPushButton("导入文件")
        self.date_but3 = QPushButton("添加文件")
        self.date_but4 = QPushButton("删除")
        self.add1 =  QPushButton("添加")
        self.table = QTableWidget()
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout1 = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout1)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        self.sign = 0
        self.sign1 = sign
        self.sign2= []
        self.data1 =[]
        self.button()
        self.devise_Ui()
    
    def button(self):
        self.Practice_child1but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Practice_child1but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child1but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child1but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child1but1.clicked.connect(self.return_fun)
        self.Practice_child1but2.clicked.connect(lambda:self.select(1))
        self.Practice_child1but3.clicked.connect(lambda:self.select(2))
        self.Practice_child1but4.clicked.connect(lambda:self.select(3))
        
        self.date_but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        
        self.add1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.add1.setMaximumSize(60, 30)
        self.add1.clicked.connect(self.save_file)
        self.table.setStyleSheet("QTableWidget{background-color:rgb(235,235,235);font:13pt '宋体';font-weight:Bold;};")
#        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        #self.table.setColumnWidth(0, 80) #设置第一列宽度为80
#        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        
        
      
    def devise_Ui(self):
        a = self.x *1.0/5
        b = self.y*1.0/5
        self.Practice_child1but1.setMaximumSize(a, b)
        self.Practice_child1but2.setMaximumSize(a, b)
        self.Practice_child1but3.setMaximumSize(a, b)
        self.Practice_child1but4.setMaximumSize(a, b)
        self.layout1.addWidget(self.Practice_child1but1, 0, 0)
        self.layout1.addWidget(self.Practice_child1but2, 0, 1)
        self.layout1.addWidget(self.Practice_child1but3, 1, 0)
        self.layout1.addWidget(self.Practice_child1but4, 1, 1)
        self.Practice_child1but1.show()
        self.Practice_child1but2.show()
        self.Practice_child1but3.show()
        self.Practice_child1but4.show()
        QApplication.processEvents()



    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date())

    
    
    def select(self, sign):
        self.sign2 =sign
        self.layout1.removeWidget(self.Practice_child1but1)
        self.layout1.removeWidget(self.Practice_child1but2)
        self.layout1.removeWidget(self.Practice_child1but3)
        self.layout1.removeWidget(self.Practice_child1but4)
        self.Practice_child1but1.close()
        self.Practice_child1but2.close()
        self.Practice_child1but3.close()
        self.Practice_child1but4.close()
        self.date_but1.clicked.connect(self.return_fun1)
        self.date_but2.clicked.connect(self.select2)
        self.date_but3.clicked.connect(self.select3)
        self.date_but4.clicked.connect(self.deletefile)
        self.select1()
    
    def read_usrdata(self):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        if self.sign1 == 1 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("一年级", "语文" ))
        elif self.sign1 == 1 and self.sign2==2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("一年级", "数学" ))
        elif self.sign1 == 1 and self.sign2==3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("一年级", "英语" ))
        elif self.sign1 == 2 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("二年级", "语文" ))
        elif self.sign1 == 2 and self.sign2==2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("二年级", "数学" ))
        elif self.sign1 == 2 and self.sign2==3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("二年级", "英语" ))
        elif self.sign1 == 3 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("三年级", "语文" ))
        elif self.sign1 == 3 and self.sign2==2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("三年级", "数学" ))
        elif self.sign1 == 3 and self.sign2==3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("三年级", "英语" ))
        elif self.sign1 == 4 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("四年级", "语文" ))
        elif self.sign1 == 4 and self.sign2==2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("四年级", "数学" ))
        elif self.sign1 == 4 and self.sign2==3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("四年级", "英语" ))
        elif self.sign1 == 5 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("五年级", "语文" ))
        elif self.sign1 == 5 and self.sign2==2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("五年级", "数学" ))
        elif self.sign1 == 5 and self.sign2==3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("五年级", "英语" ))
        elif self.sign1 == 6 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("六年级", "语文" ))
        elif self.sign1 == 6 and self.sign2==2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("六年级", "数学" ))
        elif self.sign1 == 6 and self.sign2==3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("六年级", "英语" ))
        self.filedata = c.fetchall()
        c.close()
        conn.close()

    
    
    def select1(self):
        self.read_usrdata()
        self.sign = 1
        self.lab = QLabel("已添加成功的文件如下:")
        self.lab.setStyleSheet("QLabel{color:rgb(0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.lab.setMaximumSize(300, 40)
        self.layout1.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout1.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.win_layout1.addWidget(self.date_but1, 1, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but2, 2, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but3, 3, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but4, 4, 0, 1, 2)
        self.win_layout2.addWidget(self.lab, 0, 0, 1, 1)
        self.win_layout2.addWidget(self.table, 1, 0, 19, 20)
        self.table.setRowCount(len(self.filedata))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['选择', "序号", "类别", "文件名"])
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 110)
        self.table.setColumnWidth(2, 110)
        self.table.setColumnWidth(3, 820)
        i=0
        self.filedata1= []
        for variate in self.filedata:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.filedata1.append([ck, variate[0], variate[1], variate[2], variate[3], variate[4], variate[5], variate[6]])
            itemContent = variate[0]
            self.table.setItem(i, 1, QTableWidgetItem(itemContent))
            itemContent = variate[1]
            self.table.setItem(i, 2, QTableWidgetItem(itemContent))
            itemContent = variate[5]+'.'+variate[6]
            self.table.setItem(i, 3, QTableWidgetItem(itemContent))
            i= i+1
        self.date_but1.show()
        self.date_but2.show()
        self.date_but3.show()
        self.date_but4.show()
        self.lab.show()
        self.table.show()
        self.Lchild_win1.show()
        self.Rchild_win1.show()
        QApplication.processEvents()
        
        

        
     
    def delete_but(self):
        self.win_layout2.removeWidget(self.lab)
        self.win_layout2.removeWidget(self.table)
        self.lab.deleteLater()
        self.table.close()
        
     
    def select2(self): #sign1 ->一年级到六年级， sign2 ->语文数学英语
        path, _ = QFileDialog.getOpenFileName(self, '请选择文件', 
        'C:\\', 'ppt(*.ppt *.pptx);;word(*.docx *.doc);;flash(*.swf)')
        if not path:
            return
        if self.sign1 == 1 and self.sign2 == 1:
            read_file().select_file("小学", "一年级", "语文", path)
        elif self.sign1 == 1 and self.sign2 == 2:
            read_file().select_file("小学", "一年级", "数学", path)
        elif self.sign1 == 1 and self.sign2 == 3:
            read_file().select_file("小学", "一年级", "英语", path)    
        elif self.sign1 == 2 and self.sign2 == 1:
            read_file().select_file("小学", "二年级", "语文", path)
        elif self.sign1 == 2 and self.sign2 == 2:
            read_file().select_file("小学", "二年级", "数学", path)
        elif self.sign1 == 2 and self.sign2 == 3:
            read_file().select_file("小学", "二年级", "英语", path)    
        elif self.sign1 == 3 and self.sign2 == 1:
            read_file().select_file("小学", "三年级", "语文", path)
        elif self.sign1 == 3 and self.sign2 == 2:
            read_file().select_file("小学", "三年级", "数学", path)
        elif self.sign1 == 3 and self.sign2 == 3:
            read_file().select_file("小学", "三年级", "英语", path)    
        elif self.sign1 == 4 and self.sign2 == 1:
            read_file().select_file("小学", "四年级", "语文", path)
        elif self.sign1 == 4 and self.sign2 == 2:
            read_file().select_file("小学", "四年级", "数学", path)
        elif self.sign1 == 4 and self.sign2 == 3:
            read_file().select_file("小学", "四年级", "英语", path)        
        elif self.sign1 == 5 and self.sign2 == 1:
            read_file().select_file("小学", "五年级", "语文", path)
        elif self.sign1 == 5 and self.sign2 == 2:
            read_file().select_file("小学", "五年级", "数学", path)
        elif self.sign1 == 5 and self.sign2 == 3:
            read_file().select_file("小学", "五年级", "英语", path)    
        elif self.sign1 == 6 and self.sign2 == 1:
            read_file().select_file("小学", "六年级", "语文", path)
        elif self.sign1 == 6 and self.sign2 == 2:
            read_file().select_file("小学", "六年级", "数学", path)
        elif self.sign1 == 6 and self.sign2 == 3:
            read_file().select_file("小学", "六年级", "英语", path)    
        QMessageBox.question(self, "提示!", "添加文件成功", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
   
    def select3(self):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        if self.sign1 == 1 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("一年级", "语文" ))
        elif self.sign1 == 1 and self.sign2==2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("一年级", "数学" ))
        elif self.sign1 == 1 and self.sign2==3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("一年级", "英语" ))
        elif self.sign1 == 2 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("二年级", "语文" ))
        elif self.sign1 == 2 and self.sign2==2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("二年级", "数学" ))
        elif self.sign1 == 2 and self.sign2==3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("二年级", "英语" ))
        elif self.sign1 == 3 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("三年级", "语文" ))
        elif self.sign1 == 3 and self.sign2==2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("三年级", "数学" ))
        elif self.sign1 == 3 and self.sign2==3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("三年级", "英语" ))
        elif self.sign1 == 4 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("四年级", "语文" ))
        elif self.sign1 == 4 and self.sign2==2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("四年级", "数学" ))
        elif self.sign1 == 4 and self.sign2==3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("四年级", "英语" ))
        elif self.sign1 == 5 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("五年级", "语文" ))
        elif self.sign1 == 5 and self.sign2==2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("五年级", "数学" ))
        elif self.sign1 == 5 and self.sign2==3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("五年级", "英语" ))
        elif self.sign1 == 6 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("六年级", "语文" ))
        elif self.sign1 == 6 and self.sign2==2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("六年级", "数学" ))
        elif self.sign1 == 6 and self.sign2==3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("六年级", "英语" ))
        self.data = c.fetchall()
        c.close()
        conn.close()
        self.add_data()

        

    def add_data(self):
        self.delete_but()
        self.sign = 2
        self.lab = QLabel("可添加的文件如下:")
        self.lab.setStyleSheet("QLabel{color:rgb(0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.lab.setMaximumSize(300, 40)

        self.table.setRowCount(len(self.data))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['选择', "序号", "类别", "文件名"])
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 110)
        self.table.setColumnWidth(2, 110)
        self.table.setColumnWidth(3, 820)
        QApplication.processEvents()
        i=0
        self.data1 = []
        for variate in self.data:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.data1.append([ck, variate[0], variate[1], variate[2], variate[3], variate[4], variate[5], variate[6]])
            itemContent = variate[0]
            self.table.setItem(i, 1, QTableWidgetItem(itemContent))
            itemContent = variate[1]
            self.table.setItem(i, 2, QTableWidgetItem(itemContent))
            itemContent = variate[5]+'.'+variate[6]
            self.table.setItem(i, 3, QTableWidgetItem(itemContent))
            i= i+1
        self.win_layout2.addWidget(self.add1, 0, 16, 1, 1)
        self.win_layout2.addWidget(self.lab, 0, 0, 1, 1)
        self.win_layout2.addWidget(self.table, 1, 0, 19, 20)
        self.add1.show()
        self.table.show()
        QApplication.processEvents()

    def return_fun1(self):
        if self.sign == 1:
            self.win_layout1.removeWidget(self.date_but1)
            self.win_layout1.removeWidget(self.date_but2)
            self.win_layout1.removeWidget(self.date_but3)
            self.win_layout1.removeWidget(self.date_but4)
            self.win_layout2.removeWidget(self.lab)
            self.win_layout2.removeWidget(self.table)
            self.layout1.removeWidget(self.Lchild_win1)
            self.layout1.removeWidget(self.Rchild_win1)
            self.lab.deleteLater()
            self.table.close()
            self.date_but1.close()
            self.date_but2.close()
            self.date_but3.close()
            self.date_but4.close()
            self.Lchild_win1.close()
            self.Rchild_win1.close()
            self.devise_Ui()
            
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.lab)
            self.win_layout2.removeWidget(self.add1)
            self.win_layout2.removeWidget(self.table)
            self.lab.deleteLater()
            self.table.close()
            self.add1.close()
            self.select1()
        

    
    def deletefile(self):
        removeline = []
        a= 0
        for data in self.filedata1:
            if data[0].isChecked():
                a+=1
        if a==0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件！！！" )
            return
        rely = QMessageBox.question(self, "提示!", "删除会造成数据无法恢复！！！\n确定删除？？？" , QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        for data in self.filedata1:
            if data[0].isChecked():
                file = 'D:/项目数据库'+ '/' +data[2]+'/' +data[3]+'/' +data[4]+ '/' + data[5]
                filepath= file +'/'+data[6]+'.'+data[7]
                os.remove(filepath)
                sqlpath = "D:/项目数据库/数据库/Data.db"
                conn=sqlite3.connect(sqlpath)
                c=conn.cursor()
                c.execute("delete from Education1 where no = (?)", (data[1], ))
                c.close()
                conn.close()
                row = self.table.rowCount()
                for x in range(row, 0, -1):
                    if data[1]==(self.table.item(x-1 ,1).text()):
                        self.table.removeRow(x-1)
                        removeline.append(data)
        for data in removeline:
            self.filedata1.remove(data)  
        QMessageBox.about(self, "提示!", "文件删除成功！！！" )
    
    def save_file(self):
        removeline = []
        a= 0
        for data in self.data1:
            if data[0].isChecked():
                a+=1
        if a==0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件！！！" )
            return
        for data in self.data1:
            if data[0].isChecked():
                file = 'D:/项目数据库'
                file = file + '/' +data[2]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                file = file +'/' +data[3]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                file = file +'/' +data[4]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                path = file + '/' + data[5]
                if(not(os.path.exists(path))):   #创建文件夹。
                    os.mkdir(path)
                filepath= path +'/'+data[6]+'.'+data[7]
                sqlpath = "D:/项目数据库/数据库/Data.db"
                conn=sqlite3.connect(sqlpath)
                c=conn.cursor()
                c.execute("select * from filedata where no = (?)", (data[1], ))
                for d in c.fetchall():
                    total=base64.b64decode(d[1])
                    f=open(filepath,'wb')
                    f.write(total)
                    f.close()
                c.close()
                conn.close()
                row = self.table.rowCount()
                print(row)
                for x in range(row, 0, -1):
                    if data[1]==(self.table.item(x-1 ,1).text()):
                        print(data[1])
                        self.table.removeRow(x-1)
                        removeline.append(data)
        for data in removeline:
            print(101)
            print(data[1])
            self.data1.remove(data)  
            sqlpath = "D:/项目数据库/数据库/Data.db"
            conn=sqlite3.connect(sqlpath)
            conn.execute("INSERT INTO Education1(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",(data[1], data[2], data[3], data[4], data[5], data[6], data[7]))
            conn.commit()
            conn.close()
        QMessageBox.about(self, "提示!", "文件已经保存成功了！！！" )
        

    
#添加数据中学
class addition_Date_child2(QFrame):
    def __init__(self, sign):
        super(addition_Date_child2, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_child2but1 = QPushButton("返回")
        self.Practice_child2but2 = QPushButton("语文")
        self.Practice_child2but3 = QPushButton("数学")
        self.Practice_child2but4 = QPushButton("英语")
        self.Practice_child2but5 = QPushButton("物理")
        self.Practice_child2but6 = QPushButton("化学")
        self.Practice_child2but7 = QPushButton("生物")
        self.Practice_child2but8 = QPushButton("政治")
        self.Practice_child2but9 = QPushButton("历史")
        self.Practice_child2but10 = QPushButton("地理")
        self.date_but1 = QPushButton("返回")
        self.date_but2 = QPushButton("导入文件")
        self.date_but3 = QPushButton("添加文件")
        self.date_but4 = QPushButton("删除")
        self.add1 =  QPushButton("添加")
        self.table = QTableWidget()
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout1 = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout1)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        self.sign = 0
        self.sign1 = sign
        self.sign2 =[]
        self.button()
        self.devise_Ui()
        
    
    def button(self):
        self.Practice_child2but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Practice_child2but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child2but1.clicked.connect(self.return_fun)
        self.Practice_child2but2.clicked.connect(lambda:self.select(1))
        self.Practice_child2but3.clicked.connect(lambda:self.select(2))
        self.Practice_child2but4.clicked.connect(lambda:self.select(3))
        self.Practice_child2but5.clicked.connect(lambda:self.select(4))
        self.Practice_child2but6.clicked.connect(lambda:self.select(5))
        self.Practice_child2but7.clicked.connect(lambda:self.select(6))
        self.Practice_child2but8.clicked.connect(lambda:self.select(7))
        self.Practice_child2but9.clicked.connect(lambda:self.select(8))
        self.Practice_child2but10.clicked.connect(lambda:self.select(9))
        
        self.date_but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        
        self.add1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.add1.setMaximumSize(60, 30)
        self.add1.clicked.connect(self.save_file)
        self.table.setStyleSheet("QTableWidget{background-color:rgb(235,235,235);font:13pt '宋体';font-weight:Bold;};")
        
    
    def devise_Ui(self):
        a = self.x *1.0/6
        b = self.y*1.0/5
        self.Practice_child2but1.setMaximumSize(a, b)
        self.Practice_child2but2.setMaximumSize(a, b)
        self.Practice_child2but3.setMaximumSize(a, b)
        self.Practice_child2but4.setMaximumSize(a, b)
        self.Practice_child2but5.setMaximumSize(a, b)
        self.Practice_child2but6.setMaximumSize(a, b)
        self.Practice_child2but7.setMaximumSize(a, b)
        self.Practice_child2but8.setMaximumSize(a, b)
        self.Practice_child2but9.setMaximumSize(a, b)
        self.Practice_child2but10.setMaximumSize(a, b)
        self.layout1.addWidget(self.Practice_child2but1, 0, 0)
        self.layout1.addWidget(self.Practice_child2but2, 0, 1)
        self.layout1.addWidget(self.Practice_child2but3, 0, 2)
        self.layout1.addWidget(self.Practice_child2but4, 0, 3)
        self.layout1.addWidget(self.Practice_child2but5, 1, 1)
        self.layout1.addWidget(self.Practice_child2but6, 1, 2)
        self.layout1.addWidget(self.Practice_child2but7, 1, 3)
        self.layout1.addWidget(self.Practice_child2but8, 2, 1)
        self.layout1.addWidget(self.Practice_child2but9, 2, 2)
        self.layout1.addWidget(self.Practice_child2but10, 2, 3)
        self.Practice_child2but1.show()
        self.Practice_child2but2.show()
        self.Practice_child2but3.show()
        self.Practice_child2but4.show()
        self.Practice_child2but5.show()
        self.Practice_child2but6.show()
        self.Practice_child2but7.show()
        self.Practice_child2but8.show()
        self.Practice_child2but9.show()
        self.Practice_child2but10.show()
        
        
        


    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date())
    
    
    def select(self, sign):
        self.sign2 = sign
        self.layout1.removeWidget(self.Practice_child2but1)
        self.layout1.removeWidget(self.Practice_child2but2)
        self.layout1.removeWidget(self.Practice_child2but3)
        self.layout1.removeWidget(self.Practice_child2but4)
        self.layout1.removeWidget(self.Practice_child2but5)
        self.layout1.removeWidget(self.Practice_child2but6)
        self.layout1.removeWidget(self.Practice_child2but7)
        self.layout1.removeWidget(self.Practice_child2but8)
        self.layout1.removeWidget(self.Practice_child2but9)
        self.layout1.removeWidget(self.Practice_child2but10)
        self.Practice_child2but1.close()
        self.Practice_child2but2.close()
        self.Practice_child2but3.close()
        self.Practice_child2but4.close()
        self.Practice_child2but5.close()
        self.Practice_child2but6.close()
        self.Practice_child2but7.close()
        self.Practice_child2but8.close()
        self.Practice_child2but9.close()
        self.Practice_child2but10.close()
        self.date_but1.clicked.connect(self.return_fun1)
        self.date_but2.clicked.connect(self.select2)
        self.date_but3.clicked.connect(self.select3)
        self.date_but4.clicked.connect(self.deletefile)
        self.select1()
    
        
    
    
    def read_usrdata(self):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        if self.sign1 == 1 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "语文" ))
        elif self.sign1 == 1 and self.sign2 == 2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "数学" ))
        elif self.sign1 == 1 and self.sign2 == 3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "英语" ))
        elif self.sign1 == 1 and self.sign2 == 4:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "物理" ))
        elif self.sign1 == 1 and self.sign2 == 5:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "化学" ))
        elif self.sign1 == 1 and self.sign2 == 6:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "生物" ))
        elif self.sign1 == 1 and self.sign2 == 7:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "政治" ))
        elif self.sign1 == 1 and self.sign2 == 8:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "历史" ))
        elif self.sign1 == 1 and self.sign2 == 9:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初一", "地理" ))
        elif self.sign1 == 2 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "语文" ))
        elif self.sign1 == 2 and self.sign2 == 2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "数学" ))
        elif self.sign1 == 2 and self.sign2 == 3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "英语" ))
        elif self.sign1 == 2 and self.sign2 == 4:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "物理" ))
        elif self.sign1 == 2 and self.sign2 == 5:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "化学" ))
        elif self.sign1 == 2 and self.sign2 == 6:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "生物" ))
        elif self.sign1 == 2 and self.sign2 == 7:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "政治" ))
        elif self.sign1 == 2 and self.sign2 == 8:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "历史" ))
        elif self.sign1 == 2 and self.sign2 == 9:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初二", "地理" ))
        elif self.sign1 == 3 and self.sign2==1:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "语文" ))
        elif self.sign1 == 3 and self.sign2 == 2:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "数学" ))
        elif self.sign1 ==3 and self.sign2 == 3:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "英语" ))
        elif self.sign1 == 3 and self.sign2 == 4:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "物理" ))
        elif self.sign1 == 3 and self.sign2 == 5:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "化学" ))
        elif self.sign1 == 3 and self.sign2 == 6:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "生物" ))
        elif self.sign1 == 3 and self.sign2 == 7:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "政治" ))
        elif self.sign1 == 3 and self.sign2 == 8:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "历史" ))
        elif self.sign1 == 3 and self.sign2 == 9:
            c.execute("select * from Education1 where level3 = (?)and level4 = (?)", ("初三", "地理" ))
        self.filedata = c.fetchall()
        c.close()
        conn.close()
    
    
    def select1(self):
        self.read_usrdata()
        self.sign = 1
        self.lab = QLabel("已添加成功的文件如下:")
        self.lab.setStyleSheet("QLabel{color:rgb(0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.lab.setMaximumSize(300, 40)
        self.date_but1.clicked.connect(self.delete2)
        self.layout1.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout1.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.win_layout1.addWidget(self.date_but1, 1, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but2, 2, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but3, 3, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but4, 4, 0, 1, 2)
        self.win_layout2.addWidget(self.lab, 0, 0, 1, 1)
        self.win_layout2.addWidget(self.table, 1, 0, 19, 20)
        self.table.setRowCount(len(self.filedata))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['选择', "序号", "类别", "文件名"])
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 110)
        self.table.setColumnWidth(2, 110)
        self.table.setColumnWidth(3, 820)
        i=0
        self.filedata1= []
        for variate in self.filedata:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.filedata1.append([ck, variate[0], variate[1], variate[2], variate[3], variate[4], variate[5], variate[6]])
            itemContent = variate[0]
            self.table.setItem(i, 1, QTableWidgetItem(itemContent))
            itemContent = variate[1]
            self.table.setItem(i, 2, QTableWidgetItem(itemContent))
            itemContent = variate[5]+'.'+variate[6]
            self.table.setItem(i, 3, QTableWidgetItem(itemContent))
            i= i+1
        self.date_but1.show()
        self.date_but2.show()
        self.date_but3.show()
        self.date_but4.show()
        self.lab.show()
        self.table.show()
        self.Lchild_win1.show()
        self.Rchild_win1.show()
        QApplication.processEvents()
    
    
    
  
        
    def delete_but(self):
        self.win_layout2.removeWidget(self.lab)
        self.win_layout2.removeWidget(self.table)
        self.lab.deleteLater()
        self.table.close()
    
    def select2(self):
        path, _ = QFileDialog.getOpenFileName(self, '请选择文件', 
        'C:\\', 'ppt(*.ppt *.pptx);;word(*.docx *.doc);;flash(*.swf)')
        if not path:
            return
        if self.sign1 == 1 and self.sign2 == 1:
            read_file().select_file("初中", "初一", "语文", path)
        elif self.sign1 == 1 and self.sign2 == 2:
            read_file().select_file("初中", "初一", "数学", path)
        elif self.sign1 == 1 and self.sign2 == 3:
            read_file().select_file("初中", "初一", "英语", path)    
        elif self.sign1 == 1 and self.sign2 == 4:
            read_file().select_file("初中", "初一", "物理", path)    
        elif self.sign1 == 1 and self.sign2 == 5:
            read_file().select_file("初中", "初一", "化学", path)    
        elif self.sign1 == 1 and self.sign2 == 6:
            read_file().select_file("初中", "初一", "生物", path)    
        elif self.sign1 == 1 and self.sign2 == 7:
            read_file().select_file("初中", "初一", "政治", path)    
        elif self.sign1 == 1 and self.sign2 == 8:
            read_file().select_file("初中", "初一", "历史", path)    
        elif self.sign1 == 1 and self.sign2 == 9:
            read_file().select_file("初中", "初一", "地理", path)    
        elif self.sign1 == 2 and self.sign2 == 1:
            read_file().select_file("初中", "初二", "语文", path)
        elif self.sign1 == 2 and self.sign2 == 2:
            read_file().select_file("初中", "初二", "数学", path)
        elif self.sign1 == 2 and self.sign2 == 3:
            read_file().select_file("初中", "初二", "英语", path)    
        elif self.sign1 == 2 and self.sign2 == 4:
            read_file().select_file("初中", "初二", "物理", path)    
        elif self.sign1 == 2 and self.sign2 == 5:
            read_file().select_file("初中", "初二", "化学", path)    
        elif self.sign1 == 2 and self.sign2 == 6:
            read_file().select_file("初中", "初二", "生物", path)    
        elif self.sign1 == 2 and self.sign2 == 7:
            read_file().select_file("初中", "初二", "政治", path)    
        elif self.sign1 == 2 and self.sign2 == 8:
            read_file().select_file("初中", "初二", "历史", path)    
        elif self.sign1 == 2 and self.sign2 == 9:
            read_file().select_file("初中", "初二", "地理", path)   
        elif self.sign1 == 3 and self.sign2 == 1:
            read_file().select_file("初中", "初三", "语文", path)
        elif self.sign1 == 3 and self.sign2 == 2:
            read_file().select_file("初中", "初三", "数学", path)
        elif self.sign1 == 3 and self.sign2 == 3:
            read_file().select_file("初中", "初三", "英语", path)    
        elif self.sign1 == 3 and self.sign2 == 4:
            read_file().select_file("初中", "初三", "物理", path)    
        elif self.sign1 == 3 and self.sign2 == 5:
            read_file().select_file("初中", "初三", "化学", path)    
        elif self.sign1 == 3 and self.sign2 == 6:
            read_file().select_file("初中", "初三", "生物", path)    
        elif self.sign1 == 3 and self.sign2 == 7:
            read_file().select_file("初中", "初三", "政治", path)    
        elif self.sign1 == 3 and self.sign2 == 8:
            read_file().select_file("初中", "初三", "历史", path)    
        elif self.sign1 == 3 and self.sign2 == 9:
            read_file().select_file("初中", "初三", "地理", path) 
    
    def select3(self):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        if self.sign1 == 1 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "语文" ))
        elif self.sign1 == 1 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "数学" ))
        elif self.sign1 == 1 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "英语" ))
        elif self.sign1 == 1 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "物理" ))
        elif self.sign1 == 1 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "化学" ))
        elif self.sign1 == 1 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "生物" ))
        elif self.sign1 == 1 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "政治" ))
        elif self.sign1 == 1 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "历史" ))
        elif self.sign1 == 1 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初一", "地理" ))
        elif self.sign1 == 2 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "语文" ))
        elif self.sign1 == 2 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "数学" ))
        elif self.sign1 == 2 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "英语" ))
        elif self.sign1 == 2 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "物理" ))
        elif self.sign1 == 2 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "化学" ))
        elif self.sign1 == 2 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "生物" ))
        elif self.sign1 == 2 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "政治" ))
        elif self.sign1 == 2 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "历史" ))
        elif self.sign1 == 2 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初二", "地理" ))
        elif self.sign1 == 3 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "语文" ))
        elif self.sign1 == 3 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "数学" ))
        elif self.sign1 ==3 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "英语" ))
        elif self.sign1 == 3 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "物理" ))
        elif self.sign1 == 3 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "化学" ))
        elif self.sign1 == 3 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "生物" ))
        elif self.sign1 == 3 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "政治" ))
        elif self.sign1 == 3 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "历史" ))
        elif self.sign1 == 3 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("初三", "地理" ))
        self.data = c.fetchall()
        c.close()
        conn.close()
        self.add_data()
        
    
    
    def add_data(self):
        self.delete_but()
        self.sign = 2
        self.lab = QLabel("可添加的文件如下:")
        self.lab.setStyleSheet("QLabel{color:rgb(0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.lab.setMaximumSize(300, 40)
        self.date_but1.clicked.connect(self.return_fun1)
        self.table.setRowCount(len(self.data))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['选择', "序号", "类别", "文件名"])
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 110)
        self.table.setColumnWidth(2, 110)
        self.table.setColumnWidth(3, 820)
        QApplication.processEvents()
        i=0
        self.data1= []
        for variate in self.data:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.data1.append([ck, variate[0], variate[1], variate[2], variate[3], variate[4], variate[5], variate[6]])
            itemContent = variate[0]
            self.table.setItem(i, 1, QTableWidgetItem(itemContent))
            itemContent = variate[1]
            self.table.setItem(i, 2, QTableWidgetItem(itemContent))
            itemContent = variate[5]+'.'+variate[6]
            self.table.setItem(i, 3, QTableWidgetItem(itemContent))
            i= i+1
        self.win_layout2.addWidget(self.add1, 0, 16, 1, 1)
        self.win_layout2.addWidget(self.lab, 0, 0, 1, 1)
        self.win_layout2.addWidget(self.table, 1, 0, 19, 20)
        self.add1.show()
        self.table.show()
        QApplication.processEvents()
    
    def return_fun1(self):
        if self.sign == 1:
            self.win_layout1.removeWidget(self.date_but1)
            self.win_layout1.removeWidget(self.date_but2)
            self.win_layout1.removeWidget(self.date_but3)
            self.win_layout1.removeWidget(self.date_but4)
            self.win_layout2.removeWidget(self.lab)
            self.win_layout2.removeWidget(self.table)
            self.layout1.removeWidget(self.Lchild_win1)
            self.layout1.removeWidget(self.Rchild_win1)
            self.lab.deleteLater()
            self.table.close()
            self.date_but1.close()
            self.date_but2.close()
            self.date_but3.close()
            self.date_but4.close()
            self.Lchild_win1.close()
            self.Rchild_win1.close()
            self.devise_Ui()
     
        elif self.sign ==2 :
            self.win_layout2.removeWidget(self.lab)
            self.win_layout2.removeWidget(self.add1)
            self.win_layout2.removeWidget(self.table)
            self.lab.deleteLater()
            self.table.close()
            self.add1.close()
            self.select1()
     
    def deletefile(self):
        removeline = []
        a= 0
        for data in self.filedata1:
            if data[0].isChecked():
                a+=1
        if a==0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件！！！" )
            return
        rely = QMessageBox.question(self, "提示!", "删除会造成数据无法恢复！！！\n确定删除？？？" , QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        for data in self.filedata1:
            if data[0].isChecked():
                file = 'D:/项目数据库'+ '/' +data[2]+'/' +data[3]+'/' +data[4]+ '/' + data[5]
                filepath= file +'/'+data[6]+'.'+data[7]
                os.remove(filepath)
                sqlpath = "D:/项目数据库/数据库/Data.db"
                conn=sqlite3.connect(sqlpath)
                c=conn.cursor()
                c.execute("delete from Education1 where no = (?)", (data[1], ))
                c.close()
                conn.close()
                row = self.table.rowCount()
                for x in range(row, 0, -1):
                    if (data[6]+'.'+data[7] )== (self.table.item(x-1 ,1).text()):
                        self.table.removeRow(x-1)
                        removeline.append(data)
        for data in removeline:
            self.filedata1.remove(data)  
        QMessageBox.about(self, "提示!", "文件删除成功！！！" )
     
    def save_file(self):
        removeline = []
        a= 0
        for data in self.data1:
            if data[0].isChecked():
                a+=1
        if a==0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件！！！" )
            return
        for data in self.data1:
            if data[0].isChecked():
                file = 'D:/项目数据库'
                file = file + '/' +data[2]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                file = file +'/' +data[3]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                file = file +'/' +data[4]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                path = file + '/' + data[5]
                if(not(os.path.exists(path))):   #创建文件夹。
                    os.mkdir(path)
                filepath= path +'/'+data[6]+'.'+data[7]
                sqlpath = "D:/项目数据库/数据库/Data.db"
                conn=sqlite3.connect(sqlpath)
                c=conn.cursor()
                c.execute("select * from filedata where no = (?)", (data[1], ))
                for d in c.fetchall():
                    total=base64.b64decode(d[1])
                    f=open(filepath,'wb')
                    f.write(total)
                    f.close()
                c.close()
                conn.close()
                row = self.table.rowCount()
                for x in range(row, 0, -1):
                    if data[1]==(self.table.item(x-1 ,1).text()):
                        self.table.removeRow(x-1)
                        removeline.append(data)
        for data in removeline:
            self.data1.remove(data)  
            sqlpath = "D:/项目数据库/数据库/Data.db"
            conn=sqlite3.connect(sqlpath)
            conn.execute("INSERT INTO Education1(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",(data[1], data[2], data[3], data[4], data[5], data[6], data[7]))
            conn.commit()
            conn.close()
        QMessageBox.about(self, "提示!", "文件已经保存成功了！！！" )

    
    
#添加数据高中
class addition_Date_child3(QFrame):
    def __init__(self, sign):
        super(addition_Date_child3, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.Practice_child3but1 = QPushButton("返回")
        self.Practice_child3but2 = QPushButton("语文")
        self.Practice_child3but3 = QPushButton("数学")
        self.Practice_child3but4 = QPushButton("英语")
        self.Practice_child3but5 = QPushButton("物理")
        self.Practice_child3but6 = QPushButton("化学")
        self.Practice_child3but7 = QPushButton("生物")
        self.Practice_child3but8 = QPushButton("政治")
        self.Practice_child3but9 = QPushButton("历史")
        self.Practice_child3but10 = QPushButton("地理")
        self.date_but1 = QPushButton("返回")
        self.date_but2 = QPushButton("导入文件")
        self.date_but3 = QPushButton("添加文件")
        self.date_but4 = QPushButton("删除")
        self.add1 =  QPushButton("添加")
        self.table = QTableWidget()
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout1 = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout1)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        self.sign1 = sign
        self.sign2 = []
        self.button()
        self.devise_Ui()
        
    
    def button(self):
        self.Practice_child3but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
                                 
        self.Practice_child3but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.Practice_child3but1.clicked.connect(self.return_fun)
        self.Practice_child3but2.clicked.connect(lambda:self.select(1))
        self.Practice_child3but3.clicked.connect(lambda:self.select(2))
        self.Practice_child3but4.clicked.connect(lambda:self.select(3))
        self.Practice_child3but5.clicked.connect(lambda:self.select(4))
        self.Practice_child3but6.clicked.connect(lambda:self.select(5))
        self.Practice_child3but7.clicked.connect(lambda:self.select(6))
        self.Practice_child3but8.clicked.connect(lambda:self.select(7))
        self.Practice_child3but9.clicked.connect(lambda:self.select(8))
        self.Practice_child3but10.clicked.connect(lambda:self.select(9))
        
        self.date_but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.date_but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:34px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        
        self.add1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.add1.setMaximumSize(60, 30)
        self.add1.clicked.connect(self.save_file)
        self.table.setStyleSheet("QTableWidget{background-color:rgb(235,235,235);font:13pt '宋体';font-weight:Bold;};")
        
    
    
    
    def devise_Ui(self):
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        a = self.x *1.0/6
        b = self.y*1.0/5
        self.Practice_child3but1.setMaximumSize(a, b)
        self.Practice_child3but2.setMaximumSize(a, b)
        self.Practice_child3but3.setMaximumSize(a, b)
        self.Practice_child3but4.setMaximumSize(a, b)
        self.Practice_child3but5.setMaximumSize(a, b)
        self.Practice_child3but6.setMaximumSize(a, b)
        self.Practice_child3but7.setMaximumSize(a, b)
        self.Practice_child3but8.setMaximumSize(a, b)
        self.Practice_child3but9.setMaximumSize(a, b)
        self.Practice_child3but10.setMaximumSize(a, b)
        self.layout1.addWidget(self.Practice_child3but1, 0, 0)
        self.layout1.addWidget(self.Practice_child3but2, 0, 1)
        self.layout1.addWidget(self.Practice_child3but3, 0, 2)
        self.layout1.addWidget(self.Practice_child3but4, 0, 3)
        self.layout1.addWidget(self.Practice_child3but5, 1, 1)
        self.layout1.addWidget(self.Practice_child3but6, 1, 2)
        self.layout1.addWidget(self.Practice_child3but7, 1, 3)
        self.layout1.addWidget(self.Practice_child3but8, 2, 1)
        self.layout1.addWidget(self.Practice_child3but9, 2, 2)
        self.layout1.addWidget(self.Practice_child3but10, 2, 3)
        self.Practice_child3but1.show()
        self.Practice_child3but2.show()
        self.Practice_child3but3.show()
        self.Practice_child3but4.show()
        self.Practice_child3but5.show()
        self.Practice_child3but6.show()
        self.Practice_child3but7.show()
        self.Practice_child3but8.show()
        self.Practice_child3but9.show()
        self.Practice_child3but10.show()
        
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,addition_Date())

    def select(self, sign):
        self.sign2 = sign
        self.layout1.removeWidget(self.Practice_child3but1)
        self.layout1.removeWidget(self.Practice_child3but2)
        self.layout1.removeWidget(self.Practice_child3but3)
        self.layout1.removeWidget(self.Practice_child3but4)
        self.layout1.removeWidget(self.Practice_child3but5)
        self.layout1.removeWidget(self.Practice_child3but6)
        self.layout1.removeWidget(self.Practice_child3but7)
        self.layout1.removeWidget(self.Practice_child3but8)
        self.layout1.removeWidget(self.Practice_child3but9)
        self.layout1.removeWidget(self.Practice_child3but10)
        self.Practice_child3but1.close()
        self.Practice_child3but2.close()
        self.Practice_child3but3.close()
        self.Practice_child3but4.close()
        self.Practice_child3but5.close()
        self.Practice_child3but6.close()
        self.Practice_child3but7.close()
        self.Practice_child3but8.close()
        self.Practice_child3but9.close()
        self.Practice_child3but10.close()
        self.date_but1.clicked.connect(self.return_fun1)
        self.date_but2.clicked.connect(self.select2)
        self.date_but3.clicked.connect(self.select3)
        self.date_but4.clicked.connect(self.deletefile)
        self.select1()
    

        
    
    
    def read_usrdata(self):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        if self.sign1 == 1 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "语文" ))
        elif self.sign1 == 1 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "数学" ))
        elif self.sign1 == 1 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "英语" ))
        elif self.sign1 == 1 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "物理" ))
        elif self.sign1 == 1 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "化学" ))
        elif self.sign1 == 1 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "生物" ))
        elif self.sign1 == 1 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "政治" ))
        elif self.sign1 == 1 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "历史" ))
        elif self.sign1 == 1 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "地理" ))
        elif self.sign1 == 2 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "语文" ))
        elif self.sign1 == 2 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "数学" ))
        elif self.sign1 == 2 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "英语" ))
        elif self.sign1 == 2 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "物理" ))
        elif self.sign1 == 2 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "化学" ))
        elif self.sign1 == 2 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "生物" ))
        elif self.sign1 == 2 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "政治" ))
        elif self.sign1 == 2 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "历史" ))
        elif self.sign1 == 2 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "地理" ))
        elif self.sign1 == 3 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "语文" ))
        elif self.sign1 == 3 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "数学" ))
        elif self.sign1 ==3 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "英语" ))
        elif self.sign1 == 3 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "物理" ))
        elif self.sign1 == 3 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "化学" ))
        elif self.sign1 == 3 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "生物" ))
        elif self.sign1 == 3 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "政治" ))
        elif self.sign1 == 3 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "历史" ))
        elif self.sign1 == 3 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "地理" ))
        self.filedata = c.fetchall()
        c.close()
        conn.close()
    
    def select1(self):
        self.read_usrdata()
        self.sign = 1
        self.lab = QLabel("已添加成功的文件如下:")
        self.lab.setStyleSheet("QLabel{color:rgb(0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.lab.setMaximumSize(300, 40)
        self.date_but1.clicked.connect(self.delete2)
        self.layout1.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout1.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.win_layout1.addWidget(self.date_but1, 1, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but2, 2, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but3, 3, 0, 1, 2)
        self.win_layout1.addWidget(self.date_but4, 4, 0, 1, 2)
        self.win_layout2.addWidget(self.lab, 0, 0, 1, 1)
        self.win_layout2.addWidget(self.table, 1, 0, 19, 20)
        self.table.setRowCount(len(self.filedata))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['选择', "序号", "类别", "文件名"])
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 110)
        self.table.setColumnWidth(2, 110)
        self.table.setColumnWidth(3, 820)
        i=0
        self.filedata1= []
        for variate in self.filedata:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.filedata1.append([ck, variate[0], variate[1], variate[2], variate[3], variate[4], variate[5], variate[6]])
            itemContent = variate[0]
            self.table.setItem(i, 1, QTableWidgetItem(itemContent))
            itemContent = variate[1]
            self.table.setItem(i, 2, QTableWidgetItem(itemContent))
            itemContent = variate[5]+'.'+variate[6]
            self.table.setItem(i, 3, QTableWidgetItem(itemContent))
            i= i+1
        self.date_but1.show()
        self.date_but2.show()
        self.date_but3.show()
        self.date_but4.show()
        self.lab.show()
        self.table.show()
        self.Lchild_win1.show()
        self.Rchild_win1.show()
        QApplication.processEvents()
    
        
        
    
     
    def delete_but(self):
        self.win_layout2.removeWidget(self.lab)
        self.win_layout2.removeWidget(self.table)
        self.lab.deleteLater()
        self.table.close()
    
    def select2(self):
        path, _ = QFileDialog.getOpenFileName(self, '请选择文件', 
        'C:\\', 'ppt(*.ppt *.pptx);;word(*.docx *.doc);;flash(*.swf)')
        if not path:
            return
        if self.sign1 == 1 and self.sign2 == 1:
            read_file().select_file("高中", "高一", "语文", path)
        elif self.sign1 == 1 and self.sign2 == 2:
            read_file().select_file("高中", "高一", "数学", path)
        elif self.sign1 == 1 and self.sign2 == 3:
            read_file().select_file("高中", "高一", "英语", path)    
        elif self.sign1 == 1 and self.sign2 == 4:
            read_file().select_file("高中", "高一", "物理", path)    
        elif self.sign1 == 1 and self.sign2 == 5:
            read_file().select_file("高中", "高一", "化学", path)    
        elif self.sign1 == 1 and self.sign2 == 6:
            read_file().select_file("高中", "高一", "生物", path)    
        elif self.sign1 == 1 and self.sign2 == 7:
            read_file().select_file("高中", "高一", "政治", path)    
        elif self.sign1 == 1 and self.sign2 == 8:
            read_file().select_file("高中", "高一", "历史", path)    
        elif self.sign1 == 1 and self.sign2 == 9:
            read_file().select_file("高中", "高一", "地理", path)    
        elif self.sign1 == 2 and self.sign2 == 1:
            read_file().select_file("高中", "高二", "语文", path)
        elif self.sign1 == 2 and self.sign2 == 2:
            read_file().select_file("高中", "高二", "数学", path)
        elif self.sign1 == 2 and self.sign2 == 3:
            read_file().select_file("高中", "高二", "英语", path)    
        elif self.sign1 == 2 and self.sign2 == 4:
            read_file().select_file("高中", "高二", "物理", path)    
        elif self.sign1 == 2 and self.sign2 == 5:
            read_file().select_file("高中", "高二", "化学", path)    
        elif self.sign1 == 2 and self.sign2 == 6:
            read_file().select_file("高中", "高二", "生物", path)    
        elif self.sign1 == 2 and self.sign2 == 7:
            read_file().select_file("高中", "高二", "政治", path)    
        elif self.sign1 == 2 and self.sign2 == 8:
            read_file().select_file("高中", "高二", "历史", path)    
        elif self.sign1 == 2 and self.sign2 == 9:
            read_file().select_file("高中", "高二", "地理", path)   
        elif self.sign1 == 3 and self.sign2 == 1:
            read_file().select_file("高中", "高三", "语文", path)
        elif self.sign1 == 3 and self.sign2 == 2:
            read_file().select_file("高中", "高三", "数学", path)
        elif self.sign1 == 3 and self.sign2 == 3:
            read_file().select_file("高中", "高三", "英语", path)    
        elif self.sign1 == 3 and self.sign2 == 4:
            read_file().select_file("高中", "高三", "物理", path)    
        elif self.sign1 == 3 and self.sign2 == 5:
            read_file().select_file("高中", "高三", "化学", path)    
        elif self.sign1 == 3 and self.sign2 == 6:
            read_file().select_file("高中", "高三", "生物", path)    
        elif self.sign1 == 3 and self.sign2 == 7:
            read_file().select_file("高中", "高三", "政治", path)    
        elif self.sign1 == 3 and self.sign2 == 8:
            read_file().select_file("高中", "高三", "历史", path)    
        elif self.sign1 == 3 and self.sign2 == 9:
            read_file().select_file("高中", "高三", "地理", path) 
    
    def select3(self):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        if self.sign1 == 1 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "语文" ))
        elif self.sign1 == 1 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "数学" ))
        elif self.sign1 == 1 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "英语" ))
        elif self.sign1 == 1 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "物理" ))
        elif self.sign1 == 1 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "化学" ))
        elif self.sign1 == 1 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "生物" ))
        elif self.sign1 == 1 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "政治" ))
        elif self.sign1 == 1 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "历史" ))
        elif self.sign1 == 1 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高一", "地理" ))
        elif self.sign1 == 2 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "语文" ))
        elif self.sign1 == 2 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "数学" ))
        elif self.sign1 == 2 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "英语" ))
        elif self.sign1 == 2 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "物理" ))
        elif self.sign1 == 2 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "化学" ))
        elif self.sign1 == 2 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "生物" ))
        elif self.sign1 == 2 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "政治" ))
        elif self.sign1 == 2 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "历史" ))
        elif self.sign1 == 2 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高二", "地理" ))
        elif self.sign1 == 3 and self.sign2==1:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "语文" ))
        elif self.sign1 == 3 and self.sign2 == 2:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "数学" ))
        elif self.sign1 ==3 and self.sign2 == 3:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "英语" ))
        elif self.sign1 == 3 and self.sign2 == 4:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "物理" ))
        elif self.sign1 == 3 and self.sign2 == 5:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "化学" ))
        elif self.sign1 == 3 and self.sign2 == 6:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "生物" ))
        elif self.sign1 == 3 and self.sign2 == 7:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "政治" ))
        elif self.sign1 == 3 and self.sign2 == 8:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "历史" ))
        elif self.sign1 == 3 and self.sign2 == 9:
            c.execute("select * from Education where level3 = (?)and level4 = (?)", ("高三", "地理" ))
        self.data = c.fetchall()
        c.close()
        conn.close()
        self.add_data()
        
    
    def add_data(self):
        self.delete_but()
        self.sign = 2
        self.lab = QLabel("可添加的文件如下:")
        self.lab.setStyleSheet("QLabel{color:rgb(0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.lab.setMaximumSize(300, 40)
        self.date_but1.clicked.connect(self.return_fun1)
        self.table.setRowCount(len(self.data))
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(['选择', "序号", "类别", "文件名"])
        self.table.setColumnWidth(0, 50)
        self.table.setColumnWidth(1, 110)
        self.table.setColumnWidth(2, 110)
        self.table.setColumnWidth(3, 820)
        QApplication.processEvents()
        i=0
        self.data1= []
        for variate in self.data:
            ck = QCheckBox()
            h = QHBoxLayout()
            h.setAlignment(Qt.AlignCenter)
            h.addWidget(ck)
            w = QWidget()
            w.setLayout(h)
            self.table.setCellWidget(i, 0, w)
            self.data1.append([ck, variate[0], variate[1], variate[2], variate[3], variate[4], variate[5], variate[6]])
            itemContent = variate[0]
            self.table.setItem(i, 1, QTableWidgetItem(itemContent))
            itemContent = variate[1]
            self.table.setItem(i, 2, QTableWidgetItem(itemContent))
            itemContent = variate[5]+'.'+variate[6]
            self.table.setItem(i, 3, QTableWidgetItem(itemContent))
            i= i+1
        self.win_layout2.addWidget(self.add1, 0, 16, 1, 1)
        self.win_layout2.addWidget(self.lab, 0, 0, 1, 1)
        self.win_layout2.addWidget(self.table, 1, 0, 19, 20)
        self.add1.show()
        self.table.show()
    
    
    def return_fun1(self):
        if self.sign == 1:
            self.win_layout1.removeWidget(self.date_but1)
            self.win_layout1.removeWidget(self.date_but2)
            self.win_layout1.removeWidget(self.date_but3)
            self.win_layout1.removeWidget(self.date_but4)
            self.win_layout2.removeWidget(self.lab)
            self.win_layout2.removeWidget(self.table)
            self.layout1.removeWidget(self.Lchild_win1)
            self.layout1.removeWidget(self.Rchild_win1)
            self.lab.deleteLater()
            self.table.close()
            self.date_but1.close()
            self.date_but2.close()
            self.date_but3.close()
            self.date_but4.close()
            self.Lchild_win1.close()
            self.Rchild_win1.close()
            self.devise_Ui()
     
        elif self.sign ==2 :
            self.win_layout2.removeWidget(self.lab)
            self.win_layout2.removeWidget(self.add1)
            self.win_layout2.removeWidget(self.table)
            self.lab.deleteLater()
            self.table.close()
            self.add1.close()
            self.select1()

    
    def deletefile(self):
        removeline = []
        a= 0
        for data in self.filedata1:
            if data[0].isChecked():
                a+=1
        if a==0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件！！！" )
            return
        rely = QMessageBox.question(self, "提示!", "删除会造成数据无法恢复！！！\n确定删除？？？" , QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
        if rely ==  65536:
            return 
        for data in self.filedata1:
            if data[0].isChecked():
                file = 'D:/项目数据库'+ '/' +data[2]+'/' +data[3]+'/' +data[4]+ '/' + data[5]
                filepath= file +'/'+data[6]+'.'+data[7]
                os.remove(filepath)
                sqlpath = "D:/项目数据库/数据库/Data.db"
                conn=sqlite3.connect(sqlpath)
                c=conn.cursor()
                c.execute("delete from Education1 where no = (?)", (data[1], ))
                c.close()
                conn.close()
                row = self.table.rowCount()
                for x in range(row, 0, -1):
                    if (data[6]+'.'+data[7] )== (self.table.item(x-1 ,1).text()):
                        self.table.removeRow(x-1)
                        removeline.append(data)
        for data in removeline:
            self.filedata1.remove(data)  
        QMessageBox.about(self, "提示!", "文件删除成功！！！" )
    
    def save_file(self):
        removeline = []
        a= 0
        for data in self.data1:
            if data[0].isChecked():
                a+=1
        if a==0:
            QMessageBox.about(self, "提示!", "您没有选择任何文件！！！" )
            return
        for data in self.data1:
            if data[0].isChecked():
                file = 'D:/项目数据库'
                file = file + '/' +data[2]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                file = file +'/' +data[3]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                file = file +'/' +data[4]
                if(not(os.path.exists(file))):   #创建文件夹。
                    os.mkdir(file)
                path = file + '/' + data[5]
                if(not(os.path.exists(path))):   #创建文件夹。
                    os.mkdir(path)
                filepath= path +'/'+data[6]+'.'+data[7]
                sqlpath = "D:/项目数据库/数据库/Data.db"
                conn=sqlite3.connect(sqlpath)
                c=conn.cursor()
                c.execute("select * from filedata where no = (?)", (data[1], ))
                for d in c.fetchall():
                    total=base64.b64decode(d[1])
                    f=open(filepath,'wb')
                    f.write(total)
                    f.close()
                c.close()
                conn.close()
                row = self.table.rowCount()
                for x in range(row, 0, -1):
                    if data[1]==(self.table.item(x-1 ,1).text()):
                        self.table.removeRow(x-1)
                        removeline.append(data)
        for data in removeline:
            self.data1.remove(data)  
            sqlpath = "D:/项目数据库/数据库/Data.db"
            conn=sqlite3.connect(sqlpath)
            conn.execute("INSERT INTO Education1(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",(data[1], data[2], data[3], data[4], data[5], data[6], data[7]))
            conn.commit()
            conn.close() 
        QMessageBox.about(self, "提示!", "文件已经保存成功了！！！" )
        

#导入文件时文件转换。
class read_file(QWidget):
    def __init__(self):
        super(read_file, self).__init__()
        
    def select_file(self, file1, file2, file3, path):
        file_extension = os.path.splitext(path)[1]
        if (file_extension=='.ppt'or file_extension== '.pptx'):
            print(1)
            try:
                self.ppt_to_image(path)
            except:
                QMessageBox.question(self, "提示!", "添加过程中出现错误了，请重新添加", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
                return
            file = os.path.split(path)[1]
            a = len(os.path.splitext(path)[1])
            file4 = file[0:-a]
            filepath1 = 'D:/项目数据库/tupian.zip'
            with open(filepath1, "rb") as f:
                total=base64.b64encode(f.read())   #将文件转换为字节。
            f.close()
            self.save_data('课件', file1, file2, file3, file4, 'zip', total)
            filepath2 = "D:/项目数据库/课件/"+file1+'/'+file2+"/"+file3 + "/"+file4 +".zip"
            f=open(filepath2,'wb')
            f.write(total)
            f.close()
        elif  file_extension=='.doc' or file_extension=='.docx':
            print(2)
            try:
                self.word_to_image(path)
            except:
                QMessageBox.question(self, "提示!", "添加过程中出现错误了，请重新添加", QMessageBox.Yes|QMessageBox.No ,QMessageBox.Yes)
                return
            file = os.path.split(path)[1]
            a = len(os.path.splitext(path)[1])
            file4 = file[0:-a]
            filepath1 = 'D:/项目数据库/tupian.zip'
            with open(filepath1, "rb") as f:
                total=base64.b64encode(f.read())   #将文件转换为字节。
                f.close()
            self.save_data('练习', file1, file2, file3, file4, 'zip', total)
            filepath2 = "D:/项目数据库/练习/"+file1+'/'+file2+"/"+file3 + "/"+file4 +".zip"
            f=open(filepath2,'wb')
            f.write(total)
            f.close()
        elif file_extension=='.swf' :
            print(3)
            file = os.path.split(path)[1]
            a = len(os.path.splitext(path)[1])
            file4 = file[0:-a]
            with open(path, "rb") as f:
                total=base64.b64encode(f.read())   #将文件转换为字节。
            f.close()
            self.save_data('课件', file1, file2, file3, file4, 'swf', total)
            filepath1 = "D:/项目数据库/课件/"+file1+'/'+file2+"/"+file3 + "/"+file4 +".swf"
            f=open(filepath2,'wb')
            f.write(total)
            f.close()
        QMessageBox.about(self, "提示!", "文件已经保存成功了！！！" )
    
    def file_to_zip(self, path):   #将文件夹压缩为压缩包。
        filepath = path +'.zip'
        if os.path.exists(filepath):
            os.remove(filepath)
        z = zipfile.ZipFile(filepath,'w',zipfile.ZIP_DEFLATED)
        for dirpath, dirnames, filenames in os.walk(path):
            fpath = dirpath.replace(path,'')
            fpath = fpath and fpath + os.sep or ''
            for filename in filenames:
                z.write(os.path.join(dirpath, filename),fpath+filename)
        z.close()
       
  
    
    def ppt_to_image(self,  input_file):  #将ppt 转换为图片
        file = 'D:/项目数据库/tupian'
        fileNames = glob.glob(file + r'\*')   
        if fileNames:
            for fileName in fileNames:     #将pa 文件夹中的文件删除。
                os.remove( fileName)
        #powerpoint = comtypes.client.CreateObject("kwpp.Application") #使用wps的接口
        powerpoint = client.DispatchEx("kwpp.Application") #使用wps的接口
        powerpoint.Visible = 1 #0 不显示wps程序,1 显示程序  ->不显示程序无法将ppt转化为图片。
        powerpoint.DisplayAlerts = 0  #不警告
        ppt = powerpoint.Presentations.Open(input_file)
        ppt.SaveAs(file + '.jpg', 17)    # 另存为
        ppt.Close()   # 退出
        powerpoint.Quit()
        time.sleep(2)
        self.file_to_zip(file)
    
    def word_to_image(self, filepath):
        file = 'D:/项目数据库/tupian'
        word = client.DispatchEx("kwps.Application")
        word.Visible = 0 #0 不显示wps程序
        word.DisplayAlerts = 0  #不警告
        new_file = 'D:/项目数据库/wen/xinwen.pdf'
        file_extension = os.path.splitext(filepath)[1]
        filepath1 = 'D:/项目数据库/wen/xinwen'+file_extension
        doc = word.Documents.Open(filepath)
        doc.PageSetup.PageWidth = 26*28.35     # 纸张大小, A3=6, A4=7 
        doc.PageSetup.PageHeight = 22*28.35   #1cm = 28.35pt
#            doc.PageSetup.PaperSize = 7     # 纸张大小, A3=6, A4=7 
#            doc.PageSetup.Orientation = 1     # 页面方向, 竖直=0, 水平=1
        doc.SaveAs(filepath1)  # 文档保存
        doc.Close(-1)     # doc.Close(-1)保存后关闭，doc.Close()或doc.Close(0)直接关闭不保存
        doc = word.Documents.Open(filepath1)
        doc.SaveAs(new_file, FileFormat = 17)
        doc.Close()
        word.Quit()
        os.remove(filepath1)
        try:
            pdf = fitz.open(new_file)
            for pg in range(pdf.pageCount):
                page = pdf[pg]
                rotate = int(0)
                # 每个尺寸的缩放系数为2，这将为我们生成分辨率提高4倍的图像。
                zoom_x = 3
                zoom_y = 3
                trans = fitz.Matrix(zoom_x, zoom_y).preRotate(rotate)
                pm = page.getPixmap(matrix=trans, alpha=False)
                pm.writePNG(file +'/幻灯片' + '%s.jpg' % (pg+1))
            pdf.close()
        except:
            pdf.close()
            doc.close()
        os.remove(new_file)
        self.file_to_zip(file)     

    def save_data(self, file1,file2,file3,file4,title, filename1,total):
        sqlpath = "D:/项目数据库/数据库/Data.db"
        conn=sqlite3.connect(sqlpath)
        c = conn.cursor()
        c.execute("select * from Education")
        no = len(c.fetchall())
        conn.execute("INSERT INTO Education(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename1))
        conn.commit()
        conn.execute("INSERT INTO Education1(no,level1,level2,level3,level4,name,filename) VALUES(?,?,?,?,?,?,?)",("S"+str(no+1), file1,file2,file3,file4,title, filename1))
        conn.commit()
        conn.execute("insert into filedata(no,total) values(?,?)", ("S"+str(no+1), total))
        conn.commit()
        c.close()
        conn.close()






#管理员统计界面
class Controller_census(QFrame):
    def __init__(self):
        super(Controller_census, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.returnBtn = QPushButton("返回")
        self.sex = QLabel("男女比例:")  #后期改成QPushButton 按钮，点击出现图片，男女比例图，曲线图
        self.age = QLabel("年龄分布:")
        self.usrs = QLabel("总用户数:")
        self.censustab = QTableWidget()
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.sex1 = QLabel("15:2")
        self.age1 = QLabel("18-25")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.usrs.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.sex.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.sex1.setStyleSheet("QLabel{color:rgb(0,0,240,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.age.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.age1.setStyleSheet("QLabel{color:rgb(0,0,240,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.returnBtn.setMaximumSize(50, 30) 
        self.usrs.setMaximumSize(120, 30)
        self.sex.setMaximumSize(80, 30)
        self.sex1.setMaximumSize(80, 30)
        self.age.setMaximumSize(80, 30)
        self.age1.setMaximumSize(50, 30)
        self.returnBtn.clicked.connect(self.return_fun)
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.usrs, 0, 3, 1, 1)
        self.layout.addWidget(self.sex, 0, 8, 1, 1)
        self.layout.addWidget(self.sex1, 0, 9, 1, 1)
        self.layout.addWidget(self.age, 0, 15, 1, 1)
        self.layout.addWidget(self.age1, 0, 16, 1, 1)
        self.finddata()
        
    def finddata(self):
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select User.numble,usrname,time,logonday,stude1_day,stude2_day from User_data1,User where User.numble=User_data1.numble")
        data = c.fetchall()
        b = len(data)
        self.usrs1 = QLabel(str(b))  #读取用户人数
        self.censustab.setRowCount(b)
        self.censustab.setColumnCount(8)
        self.censustab.setHorizontalHeaderLabels(['用户', '用户名','加入天数', '学习天数', '学习时长', '学习课件', '学习练习', '日均时间'])
        i =0
        for variate in data:
            self.censustab.setItem(i, 0, QTableWidgetItem(variate[0]))
            new = datetime.datetime.now()
            abcd = '%Y-%m-%d %H:%M:%S'
            b1 = datetime.datetime.strptime(variate[2], abcd)
            b = (new -b1).days +1
            self.censustab.setItem(i, 1, QTableWidgetItem(variate[1]))
            self.censustab.setItem(i, 2, QTableWidgetItem(str(b)+"天"))
            self.censustab.setItem(i, 3, QTableWidgetItem(str(variate[3])+"天"))
            ab = variate[4]+variate[5]
            if (ab/3600)>1:
                ac = str(int(ab/3600)) +'时'+str(round((ab/3600 - int(ab/3600))*60, 2)) +"分"
            else:
                ac = str(round(ab/60, 2))  +"分"
            self.censustab.setItem(i,4, QTableWidgetItem(ac))
            b = variate[4]
            if (b/3600)>1:
                c = str(int(b/3600)) +'时'+str(round((b/3600 - int(b/3600))*60, 2)) +"分"
            else:
                c = str(round(b/60, 2))  +"分"
            self.censustab.setItem(i,5, QTableWidgetItem(c))
            d = variate[5]
            if (d/3600)>1:
                e = str(int(d/3600)) +'时'+str(round((d/3600 - int(d/3600))*60, 2)) +"分"
            else:
                e = str(round(d/60, 2))+"分"
            self.censustab.setItem(i,6, QTableWidgetItem(e))
            ad = ab/variate[3]
            if (ad/3600)>1:
                ae = str(int(ad/3600)) +'时'+str(round((ad/3600 - int(ad/3600))*60, 2)) +"分"
            else:
                ae = str(round(ad/60, 2)) +"分"
            self.censustab.setItem(i,7, QTableWidgetItem(ae))
            i = i+1
        self.censustab.setStyleSheet("QTableWidget{background-color:rgb(255,255,255);font:13pt '宋体';font-weight:Bold;};");
        self.censustab.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.censustab.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.usrs1.setStyleSheet("QLabel{color:rgb(0,0,240,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.usrs1.setMaximumSize(50, 30)
        self.layout.addWidget(self.usrs1, 0, 4, 1, 1)
        self.layout.addWidget(self.censustab, 1, 0, 15, 19)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0,Controller_function())
  
        
#用户注册界面
class Usr_logon(QFrame):
    def __init__(self):
        super(Usr_logon, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usr =  QLabel("号码:")
        self.usrname = QLabel("用户名：")
        self.pwd2 = QLabel("密码:")
        self.pwd3 = QLabel("确认密码:")
        self.usrLineEdit = QLineEdit()
        self.usrLineEdit2 = QLineEdit()
        self.pwdLineEdit2 = QLineEdit()
        self.pwdLineEdit3 = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn1 = QPushButton("注册")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        self.usr.setMaximumSize(50, 40)
        self.pwd2.setMaximumSize(50, 40)
        self.pwd3.setMaximumSize(80, 40)
        #设置QLabel 的字体颜色，大小，
        self.usr.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.usrname.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        self.pwd2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.pwd3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.usrLineEdit.setMaximumSize(420, 40)
        self.usrLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit3.setMaximumSize(420, 40)
        self.codeLineEdit1.setMaximumSize(310, 40)
        #self.usrLineEdit2.setText(a)
        self.usrLineEdit.setPlaceholderText("请输入手机号码")
        self.usrLineEdit2.setPlaceholderText("请输入您的昵称")
        self.pwdLineEdit2.setPlaceholderText("请输入密码")
        self.pwdLineEdit3.setPlaceholderText("请重新输入密码")
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        self.usrLineEdit2.setFont(QFont("宋体" , 12))  #设置QLineEditn 的字体及大小
        self.usrLineEdit.setFont(QFont("宋体" , 12)) 
        self.pwdLineEdit2.setFont(QFont("宋体" , 12))
        self.pwdLineEdit3.setFont(QFont("宋体" , 12))
        self.codeLineEdit1.setFont(QFont("宋体" , 12))
        self.pwdLineEdit2.setEchoMode(QLineEdit.Password)
        self.pwdLineEdit3.setEchoMode(QLineEdit.Password)
        self.okBtn1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn1.setMaximumSize(420, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.codebel.setMaximumSize(100, 40)
        self.usrLineEdit.returnPressed.connect(self.enterPress1)  #输入结束后按回车键跳到下一个控件
        self.usrLineEdit2.returnPressed.connect(self.enterPress4)
        self.pwdLineEdit2.returnPressed.connect(self.enterPress2)
        self.pwdLineEdit3.returnPressed.connect(self.enterPress3)
        self.codeLineEdit1.returnPressed.connect(self.accept) #验证码输入后回车直接验证是否可以登录
        self.okBtn1.clicked.connect(self.accept)
        self.returnBtn.clicked.connect(self.return_record)   #点击返回键返回登录界面
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0,1 , 1, 1)
        self.layout.addWidget(self.usr, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit, 1, 5, 1, 14)
        self.layout.addWidget(self.usrname, 2, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit2, 2, 5, 1, 14)
        self.layout.addWidget(self.pwd2, 3, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit2, 3, 5, 1, 14)
        self.layout.addWidget(self.pwd3, 4, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit3, 4, 5,1, 14 )
        self.layout.addWidget(self.codeLineEdit1, 5, 5, 1, 5)
        self.layout.addWidget(self.codebel, 5, 10, 1, 6)
        self.layout.addWidget(self.change_code, 5, 12, 1, 1)
        self.layout.addWidget(self.okBtn1, 6, 5, 1, 14)
        self.renovate_code()
    
    def renovate_code(self):
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小

    def return_record(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_record())
    
    def checking1(self):  #注册时输入的号码检验是否已经注册过的
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit.text() :
                return True
        c.close()
        conn.close()
        return False
    
    def checking2(self):  #注册时密码在数据库中保存过来
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        a = self.usrLineEdit.text()
        b = self.usrLineEdit2.text()
        c = self.pwdLineEdit2.text()
        conn.execute("INSERT INTO User VALUES(?,?,?)",(a, b, c))
        conn.commit()	
        conn.close()
        
    def enterPress1(self):  #注册-》用户框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码已注册！\n请您登录！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_record())
        else:
            self.usrLineEdit2.setFocus()
    
    def enterPress4(self):  #注册-》用户名框回车确定时判断文字框是否有输入
        if len(self.usrLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "用户名不能为空！" )
            self.usrLineEdit2.setFocus()
        else:
            self.pwdLineEdit2.setFocus()
    
    def enterPress2(self):  #注册-》密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        else:
            self.pwdLineEdit3.setFocus()
     
    def enterPress3(self):   #注册-》确认密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        else:
            self.codeLineEdit1.setFocus()
    

    def accept(self):#注册时将账号密码保存并登录。
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text()) !=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "您输入的号码已注册！\n请您登录！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_record())
        elif len(self.usrLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "用户名不能为空！" )
            self.usrLineEdit2.setFocus()
        elif len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        elif len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        elif self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            win.numble = self.usrLineEdit.text()
            self.checking2()
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_informent())
      



    
#用户登录界面
class Usr_record(QFrame):
    def __init__(self):
        super(Usr_record, self).__init__()
        #self.setStyleSheet("background-color:white;")
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usr1 = QLabel("号码：")
        self.pwd1 = QLabel("密码：")
        self.usrLineEdit = QLineEdit()
        self.pwdLineEdit = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn = QPushButton("登录")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.forgetbtn = QLabel( )
        self.logonbtn = QLabel()
        self.found_sql()
        self.devise_Ui()

    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        
        self.usr1.setMaximumSize(60, 60)
        #设置QLabel 的字体颜色，大小，
        self.usr1.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.pwd1.setMaximumSize(60, 60)
        #设置QLabel 的字体颜色，大小，
        self.pwd1.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.usrLineEdit.setPlaceholderText("请输入手机号码")
        self.usrLineEdit.setMaximumSize(420, 40)
        self.usrLineEdit.setFont(QFont("宋体" , 16))  #设置QLineEditn 的字体及大小
        self.pwdLineEdit.setMaximumSize(420, 40)
        self.pwdLineEdit.setPlaceholderText("请输入密码") 
        self.pwdLineEdit.setFont(QFont("宋体" , 16))
        self.pwdLineEdit.setEchoMode(QLineEdit.Password)
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        self.codeLineEdit1.setFont(QFont("宋体" , 16))
        self.codeLineEdit1.setMaximumSize(310, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.codebel.setMaximumSize(100, 40)
        self.okBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn.setMaximumSize(420, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.forgetbtn.setText("<A href='www.baidu.com'>忘记密码</a>")
        self.logonbtn.setText("<A href='www.baidu.com'>注册</a>")
        self.forgetbtn.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:20px;font-weight:normal;font-family:Arial;}")
        self.logonbtn.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:20px;font-weight:normal;font-family:Arial;}")
        self.forgetbtn.setMaximumSize(90, 50)
        self.logonbtn.setMaximumSize(50, 50)
        self.usrLineEdit.returnPressed.connect(self.enterPress1)  #输入结束后按回车键跳到下一个控件
        self.okBtn.clicked.connect(self.accept)  #登录按钮点击判断能否登录
        self.returnBtn.clicked.connect(self.change_usr_record2)        #点击返回键连接选择身份界面
        self.forgetbtn.linkActivated.connect(self.usr_forgetbtn1)   #连接用户忘记密码界面
        self.logonbtn.linkActivated.connect(self.usr_logonbtn1)   #连接用户注册界面
        self.pwdLineEdit.returnPressed.connect(self.enterPress2)    
        self.codeLineEdit1.returnPressed.connect(self.accept)   #管理员忘记密码登录
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0, 1, 1, 1)
        self.layout.addWidget(self.usr1, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit, 1, 4, 1, 14)
        self.layout.addWidget(self.pwd1, 2, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit, 2, 4, 1, 14)
        self.layout.addWidget(self.codeLineEdit1, 3, 4, 1, 5)
        self.layout.addWidget(self.codebel, 3, 9, 1, 6)
        self.layout.addWidget(self.change_code, 3, 11, 1, 1)
        self.layout.addWidget(self.okBtn, 4, 4, 1, 14)
        self.layout.addWidget(self.forgetbtn, 5, 4, 1, 2)
        self.layout.addWidget(self.logonbtn, 5, 10, 1, 2)
        self.renovate_code()
    
    def renovate_code(self):
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小
    
    def found_sql(self):   #创建保存用户信息的数据库
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        try:                                                 #号码          用户名                  密码
            c.execute('''CREATE TABLE User(numble text,usrname text,password text)''')	
        except:
            pass
        try:                                                       #号码            姓名       出生年月               性别               学校              年级                 头像                              图片格式
            c.execute('''CREATE TABLE User_data(numble text,name text,birthday text,sex text,school text, grade text,tupian LONGBLOB,file text)''')	
        except:
            pass
        try:                                                       #号码            注册时间       加入时间            学习课件时间                    学习练习时间             上次登陆时间
            c.execute('''CREATE TABLE User_data1(numble text,time text,logonday int,stude1_day double, stude2_day double,lasttime text)''')	
        except:
            pass
        c.close()
        conn.close()
    

    def usr_forgetbtn1(self):   #连接用户忘记密码界面
        win.splitter.widget(0).setParent(None)
        Usr_forget().renovate_code()
        win.splitter.insertWidget(0, Usr_forget())
    
    def usr_logonbtn1(self):    #连接用户注册界面
        win.splitter.widget(0).setParent(None)
        Usr_logon().renovate_code()
        win.splitter.insertWidget(0, Usr_logon())

    def change_usr_record2(self):    #点击返回键连接选择身份界面
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Choice_status())
    
    def checking1(self):  #登录时检验号码是否没有注册
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit.text() :
                c.close()
                conn.close()
                return False
        c.close()
        conn.close()
        return True
    def enterPress1(self):  #登录回车确定时判断文字框是否有输入
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text())!=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "该账号还未注册！\n请先注册！" )
        else:
            self.pwdLineEdit.setFocus()
    
    def enterPress2(self):  #登录回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit.setFocus()
        else:
            self.codeLineEdit1.setFocus()
    
    def accept(self):         #登录时判断密码是否正确
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        if len(self.usrLineEdit.text())!=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        if (self.checking1()):
            QMessageBox.about(self, "提示!", "该账号还未注册！\n请先注册！" )
        if len(self.pwdLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit.setFocus()
        if self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("select * from User")
            d =0
            for variate in c.fetchall():
                if variate[0]==self.usrLineEdit.text() and variate[2]== self.pwdLineEdit.text():
                    win.numble = self.usrLineEdit.text()
                    d = 1
                    break
            c.close()
            conn.close()
            if d == 1:
                #设置一个查询用户的年级
                win.splitter.widget(0).setParent(None)
                win.splitter.insertWidget(0, Usr_function())
                self.finddata()
            else:
                QMessageBox.about(self, "提示!", "账号或密码输入错误" )
            #连接主界面函数
            
    def finddata(self):
        time1 = datetime.datetime.now()
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User_data")
        for variate in c.fetchall():
            if variate[0]== win.numble :
                win.grade = variate[5]
                win.data.append(variate[0])
                win.data.append(variate[1])
                win.data.append(variate[2])
                win.data.append(variate[3])
                win.data.append(variate[4])
                win.data.append(variate[5])
                win.data.append(variate[6])
                win.data.append(variate[7])
                break
        c.close()
        conn.close()
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User_data1")
        for variate in c.fetchall():
            if variate[0]== win.numble :
                win.data1.append(variate[0])
                win.data1.append(variate[1])
                win.data1.append(variate[2])
                win.data1.append(variate[3])
                win.data1.append(variate[4])
                ab = '%Y-%m-%d %H:%M:%S'
                b = datetime.datetime.strptime(variate[5], ab)
                time2 = time1.strftime(ab)
                if b.year ==time1.year and b.month == time1.month and b.day == time1.day:
                    a = variate[2]
                else:
                    a = variate[2] +1
                    win.data1[2] = a
                c.execute("update User_data1 set logonday=(?),lasttime = (?) where numble = (?)", (a, time2, win.numble))
                conn.commit()	
                break
        c.close()
        conn.close()
        
    
#用户忘记密码
class Usr_forget(QFrame):
    def __init__(self):
        super(Usr_forget, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        
        self.usr2 =  QLabel("号码:")
        self.pwd2 = QLabel("密码:")
        self.pwd3 = QLabel("确认密码:")
        self.usrLineEdit = QLineEdit()
        self.pwdLineEdit2 = QLineEdit()
        self.pwdLineEdit3 = QLineEdit()
        self.codeLineEdit1 = QLineEdit()
        self.okBtn1 = QPushButton("确认")
        self.returnBtn = QPushButton("返回")
        self.codebel = QLabel()
        self.change_code = QLabel( )
        self.devise_Ui()
        
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        self.usr2.setMaximumSize(50, 40)
        
        self.pwd2.setMaximumSize(50, 40)
        self.pwd3.setMaximumSize(80, 40)
        #设置QLabel 的字体颜色，大小，
        self.usr2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")
        
        self.pwd2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.pwd3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:18px;font-weight:Bold;font-family:Arial;}")      
        self.usrLineEdit.setMaximumSize(420, 40)
        
        self.pwdLineEdit2.setMaximumSize(420, 40)
        self.pwdLineEdit3.setMaximumSize(420, 40)
        self.codeLineEdit1.setMaximumSize(310, 40)
        self.usrLineEdit.setPlaceholderText("请输入手机号码")
       
        self.pwdLineEdit2.setPlaceholderText("请输入新的密码")
        self.pwdLineEdit3.setPlaceholderText("请重新输入新的密码")
        self.codeLineEdit1.setPlaceholderText("请输入右侧的验证码")
        #设置QLineEditn 的字体及大小
        self.usrLineEdit.setFont(QFont("宋体" , 12)) 
        self.pwdLineEdit2.setFont(QFont("宋体" , 12))
        self.pwdLineEdit3.setFont(QFont("宋体" , 12))
        self.codeLineEdit1.setFont(QFont("宋体" , 12))
        self.pwdLineEdit2.setEchoMode(QLineEdit.Password)
        self.pwdLineEdit3.setEchoMode(QLineEdit.Password)
        self.okBtn1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.okBtn1.setMaximumSize(420, 40)
        self.change_code.setText("<A href='www.baidu.com'>看不清，换一个</a>")
        self.change_code.setStyleSheet("QLabel{color:rgb(0,0,255,255);font-size:12px;font-weight:normal;font-family:Arial;}")
        self.change_code.setMaximumSize(120, 40)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.codebel.setMaximumSize(100, 40)
        self.usrLineEdit.returnPressed.connect(self.enterPress1)  #用户输入框按回车判断
        self.pwdLineEdit2.returnPressed.connect(self.enterPress2) #密码输入框按回车判断
        self.pwdLineEdit3.returnPressed.connect(self.enterPress3) #确认密码输入框回车判断
        self.codeLineEdit1.returnPressed.connect(self.accept)   #验证码输入框回车判断是否可登录
        self.returnBtn.clicked.connect(self.return_record)
        self.change_code.linkActivated.connect(self.renovate_code)
        self.layout.addWidget(self.returnBtn, 0, 1, 1, 1)
        self.layout.addWidget(self.usr2, 1, 3, 1, 1)
        self.layout.addWidget(self.usrLineEdit, 1, 5, 1, 14)
        self.layout.addWidget(self.pwd2, 3, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit2, 3, 5, 1, 14)
        self.layout.addWidget(self.pwd3, 4, 3, 1, 1)
        self.layout.addWidget(self.pwdLineEdit3, 4, 5,1, 14 )
        self.layout.addWidget(self.codeLineEdit1, 5, 5, 1, 5)
        self.layout.addWidget(self.codebel, 5, 10, 1, 6)
        self.layout.addWidget(self.change_code, 5, 12, 1, 1)
        self.layout.addWidget(self.okBtn1, 6, 5, 1, 14)
        self.renovate_code()
   
    def return_record(self): 
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_record())
    def renovate_code(self):   
        list = ['0', '1', '2', '3', '4', '5', '6', '7', '8', '9',
                 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z',
                 'A', 'B', 'C', 'D', 'E', 'F', 'G', 'H', 'I', 'J', 'K', 'L', 'M', 'N', 'O', 'P', 'Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z']
        self.code = ''
        for num in range(1,5):
            self.code = self.code + list[random.randint(0, 61)]
        image = ImageCaptcha().generate_image(self.code)
        image.save("D:/项目数据库/wen/code.png")
        self.codebel.setPixmap(QPixmap("D:/项目数据库/wen/code.png"))
        self.codebel.setScaledContents (True) # 让图片自适应label大小

    
    def checking1(self):  #忘记密码时检验号码是否没有注册
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit.text() :
                return False
        c.close()
        conn.close()
        return True
    

    def checking4(self):  #忘记密码时将新的密码在数据库中修改过来
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User")
        for variate in c.fetchall():
            if variate[0]==self.usrLineEdit.text() :
                conn.execute("update User set password=(?) where numble=(?)",(self.pwdLineEdit2.text(),variate[0],))
                break
        conn.commit()	
        c.close()
        conn.close()
    
    

    
    def enterPress1(self):  #忘记密码时回车确定时判断文字框是否有输入
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text())!=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "该账号还未注册！\n请先注册！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_logon())
        else:
            self.pwdLineEdit2.setFocus()
    
    
    
    def enterPress2(self):  #忘记密码-》密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit2.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        else:
            self.pwdLineEdit3.setFocus()
    
    def enterPress3(self):  #忘记密码-》确认密码框回车确定时判断文字框是否有输入
        if len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        else:
            self.codeLineEdit1.setFocus()
    

    def accept(self):  #忘记密码时验证是否可以登录
        if len(self.usrLineEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "号码不能为空！" )
            self.usrLineEdit.setFocus()
        elif len(self.usrLineEdit.text())!=11:
            QMessageBox.about(self, "提示!", "您输入的号码是错误的！\n请重新输入" )
            self.usrLineEdit2.setFocus()
        elif (self.checking1()):
            QMessageBox.about(self, "提示!", "该账号还未注册！\n请先注册！" )
            time.sleep(2)
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_logon())
        elif len(self.pwdLineEdit2.text()) == 0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit2.setFocus()
        elif len(self.pwdLineEdit3.text())==0:
            QMessageBox.about(self, "提示!", "密码不能为空！" )
            self.pwdLineEdit3.setFocus()
        elif self.pwdLineEdit2.text() != self.pwdLineEdit3.text():
            QMessageBox.about(self, "提示!", "您输入的密码前后不相同！！" )
        elif self.code.lower() !=self.codeLineEdit1.text().lower():
            QMessageBox.about(self, "提示!", "验证码输入错误" )
            self.renovate_code()
            self.codeLineEdit1.setFocus()
        else:
            win.numble = self.usrLineEdit.text()
            self.checking4()
            #设置一个查询用户年级的函数
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_function())
            self.finddata()
            #连接主窗口界面。
            
    
    def finddata(self):
        time1 = datetime.datetime.now()
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User_data")
        for variate in c.fetchall():
            if variate[0]== win.numble :
                win.data.append(variate[0])
                win.data.append(variate[1])
                win.data.append(variate[2])
                win.data.append(variate[3])
                win.data.append(variate[4])
                win.data.append(variate[5])
                win.data.append(variate[6])
                win.data.append(variate[7])
                break
        c.execute("select * from User_data1")
        for variate in c.fetchall():
            if variate[0]== win.numble :
                win.data1.append(variate[0])
                win.data1.append(variate[1])
                win.data1.append(variate[2])
                win.data1.append(variate[3])
                win.data1.append(variate[4])
                ab = '%Y-%m-%d %H:%M:%S'
                b = datetime.datetime.strptime(variate[5], ab)
                time2 = time1.strftime(ab)
                if b.year ==time1.year and b.month == time1.month and b.day == time1.day:
                    a = variate[2]
                else:
                    a = variate[2] +1
                    win.data1[2] = a
                c.execute("update User_data1 set logonday = (?),lasttime = (?) where numble = (?)", (a, time2, win.numble))
                conn.commit()	
                break
        c.close()
        conn.close()
    
    
#用户信息填写
class Usr_informent(QFrame):
    def __init__(self):
        super(Usr_informent, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.sure = QPushButton("确认")
        self.chang_image  = QPushButton("换头像")
        self.name =  QLabel("姓名:")
        self.year  = QLabel("出生年月")
        self.yearcb = QComboBox()
        self.monthcb = QComboBox()
        self.sex =  QLabel("性别:")
        self.sexcb = QComboBox()
        self.school = QLabel("学校:")
        self.grade = QLabel("选择年级")
        self.gradecb = QComboBox()
        self.nameEdit = QLineEdit()
        self.tupian = QLabel()
        self.schoolEiit = QLineEdit()
        self.devise_Ui()
     
   
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (100, 0, 0, 0)
        yearnb = []
        for  i in range(1980, 2020):
            yearnb.append(str(i))
        monthmb = []
        for i in range(1, 13):
            monthmb.append(str(i))
        grade = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级', 
                     '初一', '初二', '初三', 
                     '高一', '高二', '高三']
        self.sex.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.chang_image.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.school.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.grade.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.name.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.year.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.sure.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        self.nameEdit.setPlaceholderText("请输入姓名")
        self.schoolEiit.setPlaceholderText("请输入学校名称")
        self.nameEdit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.schoolEiit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.name.setMaximumSize(50, 40)
        self.chang_image.setMaximumSize(90, 40)
        self.school.setMaximumSize(50, 40)
        self.year.setMaximumSize(95, 40)
        self.sex.setMaximumSize(50, 40)
        self.grade.setMaximumSize(95, 40)
        self.nameEdit.setMaximumSize(420, 40)
        self.schoolEiit.setMaximumSize(420, 40)
        self.sure.setMaximumSize(420, 40)
        self.sexcb.setMaximumSize(420, 40)
        self.gradecb.setMaximumSize(420, 40)
        self.yearcb.setMaximumSize(220, 40)
        self.monthcb.setMaximumSize(175, 40)
        self.tupian.setMaximumSize(250, 250)
        self.sexcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.yearcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.monthcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.gradecb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.sexcb.addItems(['男', '女'])
        self.yearcb.addItems(yearnb)
        self.monthcb.addItems(monthmb)
        self.gradecb.addItems(grade)
        self.layout.addWidget(self.tupian, 1, 1, 4, 4)
        self.layout.addWidget(self.chang_image, 4, 2, 1, 2)
        self.layout.addWidget(self.name, 1, 6, 1, 1)
        self.layout.addWidget(self.nameEdit, 1, 8, 1, 8)
        self.layout.addWidget(self.sex, 2, 6, 1, 1)
        self.layout.addWidget(self.sexcb, 2, 8, 1, 8)
        self.layout.addWidget(self.year, 3, 6, 1, 1)
        self.layout.addWidget(self.yearcb,3, 8, 1, 4 )
        self.layout.addWidget(self.monthcb, 3, 11, 1, 7)
        self.layout.addWidget(self.school, 4, 6, 1, 1)
        self.layout.addWidget(self.schoolEiit, 4, 8, 1, 8)
        self.layout.addWidget(self.grade, 5, 6, 1, 1)
        self.layout.addWidget(self.gradecb, 5, 8, 1, 8)
        self.layout.addWidget(self.sure, 6, 8, 1, 8)
        self.image()
        self.sure.clicked.connect(self.connect_fun)
        self.chang_image.clicked.connect(self.chang_fun)
    
    def image(self):
        self.image_path = "D:/项目数据库/头像/a7.jpeg"
        self.file = os.path.splitext(self.image_path)[1]
        self.tupian.setPixmap(QPixmap(self.image_path))
        self.tupian.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
    
    def chang_fun(self):
        path, _ = QFileDialog.getOpenFileName(self, '请选择文件', 
        'D:\\', 'image(*.jpg)')
        if path:
            self.image_path = path
            self.file = os.path.splitext(self.image_path)[1]
            self.tupian.setPixmap(QPixmap(self.image_path))
            self.tupian.setScaledContents (True) # 让图片自适应label大小
        else:
            self.image()
            
        
    
    def save_data(self):
        a =  self.nameEdit.text()
        b = self.yearcb.currentText() +'-' +self.monthcb.currentText()
        c = self.sexcb.currentText()
        d = self.schoolEiit.text()
        e = self.gradecb.currentText()
        with open(self.image_path, "rb") as f:
            total=base64.b64encode(f.read())   #将文件转换为字节。
        f.close()
        win.grade = e
        win.data = [win.numble, a, b, c, d, e, total, self.file]
        ab = '%Y-%m-%d %H:%M:%S'
        theTime = datetime.datetime.now().strftime(ab) 
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        conn.execute("INSERT INTO User_data VALUES(?,?,?,?,?,?,?,?)",(win.numble, a, b, c, d, e,total, self.file))
        conn.execute("INSERT INTO User_data1 VALUES(?,?,?,?,?,?)",(win.numble, theTime , 1,0.0,0.0, theTime))
        win.data1 = [win.numble, theTime , 1, 0.0,0.0, theTime ]
        conn.commit()	
        conn.close()
        sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        try:                                                  #开始时间  ， 文件类别， 文件名 ， 结束时间
            c.execute('''CREATE TABLE User_data(strat_time text,category text, filename text,last_time text)''')	
        except:
            pass
        c.close()
        conn.close()
    
    def connect_fun(self):
        if len(self.nameEdit.text()) == 0:
            QMessageBox.about(self, "提示!", "姓名框不能为空！！" )
            self.nameEdit.setFocus()
        if len(self.schoolEiit.text()) == 0:
            QMessageBox.about(self, "提示!", "学校框不能为空！！" )
            self.schoolEiit.setFocus()
        else:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0, Usr_function())
            self.save_data()
        



#用户功能界面
class Usr_function(QFrame):
    def __init__(self):
        super(Usr_function, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.mainbutton1 = QPushButton("课件")  #用户功能界面的控件
        self.mainbutton2 = QPushButton("练习")
        self.mainbutton3 = QPushButton("学习报告")
        self.mainbutton4 = QPushButton("我的")
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        b= self.screenRect.height()  *1.0/5  
        a = self.screenRect.width() *1.0/5 
       
        self.mainbutton1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.mainbutton2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.mainbutton3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.mainbutton4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:32px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                QPushButton:hover{background-color:rgb(50, 170, 200)}")  
        self.mainbutton1.clicked.connect(self.connect_fun1)
        self.mainbutton2.clicked.connect(self.connect_fun2)
        self.mainbutton3.clicked.connect(self.connect_fun3)
        self.mainbutton4.clicked.connect(self.connect_fun4)
        self.layout.addWidget(self.mainbutton1, 0, 0)#往网格的不同坐标添加不同的组件
        self.layout.addWidget(self.mainbutton2, 0, 1)
        self.layout.addWidget(self.mainbutton3, 1, 0)
        self.layout.addWidget(self.mainbutton4, 1, 1)
        self.mainbutton1.setMaximumSize(a, b)
        self.mainbutton2.setMaximumSize(a, b)
        self.mainbutton3.setMaximumSize(a, b)
        self.mainbutton4.setMaximumSize(a, b)
        
    def connect_fun1(self):
        win.splitter.widget(0).setParent(None)
        if (win.grade =='一年级' or win.grade =='二年级' or win.grade == "三年级"):
            win.splitter.insertWidget(0, win.usr_window1_child1)
        elif(win.grade =='四年级' or win.grade =='五年级' or win.grade == "六年级"):
            win.splitter.insertWidget(0, win.usr_window1_child1)
        elif(win.grade =='初一' or win.grade =='初二' or win.grade == "初三") :
            win.splitter.insertWidget(0, win.usr_window1_child2)
        elif(win.grade =='高一' or win.grade =='高二' or win.grade == "高三") :
            win.splitter.insertWidget(0, win.usr_window1_child3)
    
    def connect_fun2(self):
        win.splitter.widget(0).setParent(None)
        win.sign = 2
        if (win.grade =='一年级' or win.grade =='二年级' or win.grade == "三年级"):
            win.splitter.insertWidget(0, win.usr_window2_child1)
        elif(win.grade =='四年级' or win.grade =='五年级' or win.grade == "六年级"):
            win.splitter.insertWidget(0, win.usr_window2_child1)
        elif(win.grade =='初一' or win.grade =='初二' or win.grade == "初三") :
            win.splitter.insertWidget(0, win.usr_window2_child2)
        elif(win.grade =='高一' or win.grade =='高二' or win.grade == "高三") :
            win.splitter.insertWidget(0, win.usr_window2_child3)
    
    def connect_fun3(self):
        win.splitter.widget(0).setParent(None)
        Usr_report().information()
        win.splitter.insertWidget(0, Usr_report())
    
    def connect_fun4(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_myself())
        
    
    
#用户课件小学界面
class Usr_window1_child1(QFrame):
    def __init__(self):
        super(Usr_window1_child1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.child1_win1but1 = QPushButton("返回")#课件-小学的控件
        self.child1_win1but2 = QPushButton("语文")
        self.child1_win1but3 = QPushButton("数学")
        self.child1_win1but4 = QPushButton("英语")
        self.window1tree = QTreeView()
        self.window1tree1 = QTextEdit()
        self.win_lab = QLabel()
        self.pa = 'D://项目数据库/tupian'
        self.a =0
        self.sign =0
        self.path = []
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        
        self.child1_win1but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child1_win1but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                        
        self.child1_win1but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child1_win1but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child1_win1but1.clicked.connect(self.return_fun)
        self.child1_win1but2.clicked.connect(lambda:self.select1(1))
        self.child1_win1but3.clicked.connect(lambda:self.select1(2))
        self.child1_win1but4.clicked.connect(lambda:self.select1(3))
        self.win_layout1.addWidget(self.child1_win1but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.child1_win1but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.child1_win1but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.child1_win1but4, 4, 0,1,2)
        self.win_layout2.addWidget(self.window1tree1, 0, 0, 20, 20)
        
    def return_fun(self):
        win.sign = 0
        if self.sign == 0 or self.sign == 1:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Usr_function())
        elif self.sign == 2:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        elif self.sign == 3:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
    
    def select1(self, sign):
        if sign == 1:
            self.TreeView("D:/项目数据库/课件/小学/" +win.grade+"/语文")
        elif sign ==2:
            self.TreeView("D:/项目数据库/课件/小学/" +win.grade+"/数学")
        elif sign == 3:
            self.TreeView("D:/项目数据库/课件/小学/" +win.grade+"/英语")
    
    def delLwin(self):
        if self.sign == 0:
            self.win_layout2.removeWidget(self.window1tree1)
            self.window1tree1.close()
        elif self.sign ==1:
            self.win_layout2.removeWidget(self.window1tree)
            self.window1tree.close()
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.aswiget)
            self.aswiget.close()
        elif self.sign == 3:
            self.win_layout2.removeWidget(self.win_lab)
            self.win_lab.close()
            self.a = 0
    
    
    def TreeView(self, path):   #目录树
        if(os.path.exists(path)):
            self.path = path
            self.delLwin()
            self.sign = 1
            self.model = QFileSystemModel()
            self.model.setRootPath(path)
            #model.setFilter(QDir.NoDotAndDotDot | QDir.Dirs)# 只显示文件夹，不显示文件中的其他小文件。
            self.window1tree.setModel(self.model)
            self.window1tree.setAnimated(False)
            self.window1tree.setIndentation(20)
            self.window1tree.setSortingEnabled(True)
            self.window1tree.setRootIndex(self.model.index(path)) #只显示设置的那个文件路径。
            self.window1tree.doubleClicked.connect(self.file_name) #双击文件打开
            self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
            self.window1tree.show()
            QApplication.processEvents()
        else:
            QMessageBox.about(self, "提示!","文件夹没有任何有效文件\n请您先进行《保存数据》功能保存数据！！" )
    
    def file_name(self,Qmodelidx):   #判断文件
        win.time2 =datetime.datetime.now()
        win.sign = 1
        path = self.model.filePath(Qmodelidx) #文件的地址。
        win.filename = os.path.split(path)[1]
        pa = path[-4:]
        if pa == '.swf':    #如果是flash 文件
            self.read_flash(path)
        elif pa == '.zip':
            self.read_zip(path)
    
    
    
    def read_flash(self, path):  #读取flash 文件
        self.delLwin()
        self.sign = 2
        self.aswiget = QAxContainer.QAxWidget()
        self.win_layout2.addWidget(self.aswiget, 0, 0, 20, 20)
        self.aswiget.adjustSize()
        self.aswiget.setFocusPolicy(QtCore.Qt.StrongFocus)
        self.aswiget.setControl("{D27CDB6E-AE6D-11cf-96B8-444553540000}")  #flash的com接口
        self.aswiget.dynamicCall("LoadMovie(long,string)", 'something',path)
        self.aswiget.show()
    
    def read_zip(self, path):   #读取zip文件
        self.delLwin()
        self.sign = 3
        self.a= 1  #ppt页数变化,
        self.win_layout2.addWidget(self.win_lab, 0, 0, 20, 20)
        self.MaximumButton = QPushButton(self.win_lab)
        self.win_lab.setMouseTracking(True) # 设置widget鼠标跟踪
        self.MaximumButton.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(100, 100, 100)} ")
        QApplication.processEvents()
        self.width = self.win_lab.width()
        self.height = self.win_lab.height()
        self.MaximumButton.resize(20, 20)
        self.MaximumButton.move(self.width-22, self.height-22)
        self.MaximumButton.clicked.connect(self.select_class)
        
        if os.path.isdir(self.pa):  #建立新的文件夹
            fileNames = glob.glob(self.pa + r'/*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(self.pa)
        try:
            with zipfile.ZipFile(path, 'r') as zf:    
                for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
                    right_fn =  self.pa + '/'+fn
                    with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                        with zf.open(fn, 'r') as origin_file:  # 打开原文件
                            shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
            pa1 = self.pa + '/幻灯片1.jpg'
            pa2 = self.pa +'/幻灯片1.jpeg'
            img = Image.open(pa1)
            out = img.resize((self.width, self.height),Image.ANTIALIAS) #将图片改变分辨率为窗口大小
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
        except:
            QMessageBox.about(self, "提示!", "此文件是破损的，已删除！！" )
        self.win_lab.show()
        
    
    def select_class(self):
        self.max =  max_widget(1, self.a)
        self.max.show()
    
    
    
    def close_max_widget(self, a):  #最大化变回原来的
        self.a = a
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width, self.height),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.win_lab.setPixmap (pixmap) # 在label上显示图片
#        self.win_lab.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
        
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
           

     
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if self.a != 0:
            if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
                self.cut_images()
            if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
                self.add_images()
            if  event.pos().y() > 620:  
                self.MaximumButton.show()
            else:
                self.MaximumButton.hide()
                QApplication.processEvents()

 
#用户课件初中界面
class Usr_window1_child2(QFrame):
    def __init__(self):
        super(Usr_window1_child2, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.child2_win1but1 = QPushButton("返回")#课件-初中的控件
        self.child2_win1but2 = QPushButton("语文")
        self.child2_win1but3 = QPushButton("数学")
        self.child2_win1but4 = QPushButton("英语")
        self.child2_win1but5 = QPushButton("物理")
        self.child2_win1but6 = QPushButton("化学")
        self.child2_win1but7 = QPushButton("生物")
        self.child2_win1but8 = QPushButton("政治")
        self.child2_win1but9 = QPushButton("历史")
        self.child2_win1but10 = QPushButton("地理")
        self.window1tree = QTreeView()
        self.window1tree1 = QTextEdit()
        self.win_lab = QLabel()
        self.pa = 'D://项目数据库/tupian'
        self.a =0
        self.sign =0
        self.path = []
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        
        self.child2_win1but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                        
        self.child2_win1but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}") 
        self.child2_win1but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but1.clicked.connect(self.return_fun)
        self.child2_win1but2.clicked.connect(lambda:self.select1(1))
        self.child2_win1but3.clicked.connect(lambda:self.select1(2))
        self.child2_win1but4.clicked.connect(lambda:self.select1(3))
        self.child2_win1but5.clicked.connect(lambda:self.select1(4))
        self.child2_win1but6.clicked.connect(lambda:self.select1(5))
        self.child2_win1but7.clicked.connect(lambda:self.select1(6))
        self.child2_win1but8.clicked.connect(lambda:self.select1(7))
        self.child2_win1but9.clicked.connect(lambda:self.select1(8))
        self.child2_win1but10.clicked.connect(lambda:self.select1(9))
        self.win_layout1.addWidget(self.child2_win1but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree1, 0, 0, 20, 20)
        
        


    def return_fun(self):
        win.sign = 0
        if self.sign == 0 or self.sign == 1:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Usr_function())
        elif self.sign == 2:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        elif self.sign == 3:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        
        
        
    def select1(self, sign):
        if sign == 1:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/语文")
        elif sign ==2:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/数学")
        elif sign == 3:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/英语")
        elif sign == 4:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/物理")
        elif sign == 5:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/化学")
        elif sign == 6:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/生物")
        elif sign == 7:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/政治")
        elif sign == 8:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/历史")
        elif sign == 9:
            self.TreeView("D:/项目数据库/课件/初中/" +win.grade+"/地理")
        
        
        
    def delLwin(self):
        if self.sign == 0:
            self.win_layout2.removeWidget(self.window1tree1)
            self.window1tree1.close()
        elif self.sign ==1:
            self.win_layout2.removeWidget(self.window1tree)
            self.window1tree.close()
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.aswiget)
            self.aswiget.close()
        elif self.sign == 3:
            self.win_layout2.removeWidget(self.win_lab)
            self.win_lab.close()
            self.a = 0
    
    
    
    def TreeView(self, path):   #目录树
        if(os.path.exists(path)):
            self.path = path
            self.delLwin()
            self.sign = 1
            self.model = QFileSystemModel()
            self.model.setRootPath(path)
            #model.setFilter(QDir.NoDotAndDotDot | QDir.Dirs)# 只显示文件夹，不显示文件中的其他小文件。
            self.window1tree.setModel(self.model)
            self.window1tree.setAnimated(False)
            self.window1tree.setIndentation(20)
            self.window1tree.setSortingEnabled(True)
            self.window1tree.setRootIndex(self.model.index(path)) #只显示设置的那个文件路径。
            self.window1tree.doubleClicked.connect(self.file_name) #双击文件打开
            self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
            self.window1tree.show()
            QApplication.processEvents()
        else:
            QMessageBox.about(self, "提示!","文件夹没有任何有效文件\n请您先进行《保存数据》功能保存数据！！" )
    
    def file_name(self,Qmodelidx):   #判断文件
        win.time2 =datetime.datetime.now()
        win.sign = 1
        path = self.model.filePath(Qmodelidx) #文件的地址。
        win.filename = os.path.split(path)[1]
        pa = path[-4:]
        if pa == '.swf':    #如果是flash 文件
            self.read_flash(path)
        elif pa == '.zip':
            self.read_zip(path)
    
    
    
    def read_flash(self, path):  #读取flash 文件
        self.delLwin()
        self.sign = 2
        self.aswiget = QAxContainer.QAxWidget()
        self.win_layout2.addWidget(self.aswiget, 0, 0, 20, 20)
        self.aswiget.adjustSize()
        self.aswiget.setFocusPolicy(QtCore.Qt.StrongFocus)
        self.aswiget.setControl("{D27CDB6E-AE6D-11cf-96B8-444553540000}")  #flash的com接口
        self.aswiget.dynamicCall("LoadMovie(long,string)", 'something',path)
        self.aswiget.show()
    
    def read_zip(self, path):   #读取zip文件
        self.delLwin()
        self.sign = 3
        self.a= 1  #ppt页数变化,
        self.win_layout2.addWidget(self.win_lab, 0, 0, 20, 20)
        self.MaximumButton = QPushButton(self.win_lab)
        self.win_lab.setMouseTracking(True) # 设置widget鼠标跟踪
        self.MaximumButton.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(100, 100, 100)} ")
        QApplication.processEvents()
        self.width = self.win_lab.width()
        self.height = self.win_lab.height()
        self.MaximumButton.resize(20, 20)
        self.MaximumButton.move(self.width-22, self.height-22)
        self.MaximumButton.clicked.connect(self.select_class)
        
        if os.path.isdir(self.pa):  #建立新的文件夹
            fileNames = glob.glob(self.pa + r'/*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(self.pa)
        try:
            with zipfile.ZipFile(path, 'r') as zf:    
                for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
                    right_fn =  self.pa + '/'+fn
                    with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                        with zf.open(fn, 'r') as origin_file:  # 打开原文件
                            shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
            pa1 = self.pa + '/幻灯片1.jpg'
            pa2 = self.pa +'/幻灯片1.jpeg'
            img = Image.open(pa1)
            out = img.resize((self.width, self.height),Image.ANTIALIAS) #将图片改变分辨率为窗口大小
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
        except:
            QMessageBox.about(self, "提示!", "此文件是破损的，已删除！！" )
        self.win_lab.show()
        
    
    def select_class(self):
        self.max =  max_widget(2, self.a)
        self.max.show()
    
    
    
    def close_max_widget(self, a):  #最大化变回原来的
        self.a = a
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width, self.height),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.win_lab.setPixmap (pixmap) # 在label上显示图片
#        self.win_lab.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
        
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
           

     
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if self.a != 0:
            if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
                self.cut_images()
            if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
                self.add_images()
            if  event.pos().y() > 620:  
                self.MaximumButton.show()
            else:
                self.MaximumButton.hide()
                QApplication.processEvents()

#用户课件高中界面
class Usr_window1_child3(QFrame):
    def __init__(self):
        super(Usr_window1_child3, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.child2_win1but1 = QPushButton("返回")#课件-初中的控件
        self.child2_win1but2 = QPushButton("语文")
        self.child2_win1but3 = QPushButton("数学")
        self.child2_win1but4 = QPushButton("英语")
        self.child2_win1but5 = QPushButton("物理")
        self.child2_win1but6 = QPushButton("化学")
        self.child2_win1but7 = QPushButton("生物")
        self.child2_win1but8 = QPushButton("政治")
        self.child2_win1but9 = QPushButton("历史")
        self.child2_win1but10 = QPushButton("地理")
        self.window1tree = QTreeView()
        self.window1tree1 = QTextEdit()
        self.win_lab = QLabel()
        self.pa = 'D://项目数据库/tupian'
        self.a =0
        self.sign =0
        self.path = []
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        
        self.child2_win1but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                        
        self.child2_win1but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}") 
        self.child2_win1but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win1but1.clicked.connect(self.return_fun)
        self.child2_win1but2.clicked.connect(lambda:self.select1(1))
        self.child2_win1but3.clicked.connect(lambda:self.select1(2))
        self.child2_win1but4.clicked.connect(lambda:self.select1(3))
        self.child2_win1but5.clicked.connect(lambda:self.select1(4))
        self.child2_win1but6.clicked.connect(lambda:self.select1(5))
        self.child2_win1but7.clicked.connect(lambda:self.select1(6))
        self.child2_win1but8.clicked.connect(lambda:self.select1(7))
        self.child2_win1but9.clicked.connect(lambda:self.select1(8))
        self.child2_win1but10.clicked.connect(lambda:self.select1(9))
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.win_layout1.addWidget(self.child2_win1but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.child2_win1but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree1, 0, 0, 20, 20)
        


    def return_fun(self):
        win.sign = 0
        if self.sign == 0 or self.sign == 1:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Usr_function())
        elif self.sign == 2:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
            
        elif self.sign == 3:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude1_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "课件", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        
        
    
    def select1(self, sign):
        if sign == 1:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/语文")
        elif sign ==2:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/数学")
        elif sign == 3:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/英语")
        elif sign == 4:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/物理")
        elif sign == 5:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/化学")
        elif sign == 6:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/生物")
        elif sign == 7:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/政治")
        elif sign == 8:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/历史")
        elif sign == 9:
            self.TreeView("D:/项目数据库/课件/高中/" +win.grade+"/地理")
        
        
        
    def delLwin(self):
        if self.sign == 0:
            self.win_layout2.removeWidget(self.window1tree1)
            self.window1tree1.close()
        elif self.sign ==1:
            self.win_layout2.removeWidget(self.window1tree)
            self.window1tree.close()
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.aswiget)
            self.aswiget.close()
        elif self.sign == 3:
            self.win_layout2.removeWidget(self.win_lab)
            self.win_lab.close()
            self.a = 0
    
    
    
    def TreeView(self, path):   #目录树
        if(os.path.exists(path)):
            self.path = path
            self.delLwin()
            self.sign = 1
            self.model = QFileSystemModel()
            self.model.setRootPath(path)
            #model.setFilter(QDir.NoDotAndDotDot | QDir.Dirs)# 只显示文件夹，不显示文件中的其他小文件。
            self.window1tree.setModel(self.model)
            self.window1tree.setAnimated(False)
            self.window1tree.setIndentation(20)
            self.window1tree.setSortingEnabled(True)
            self.window1tree.setRootIndex(self.model.index(path)) #只显示设置的那个文件路径。
            self.window1tree.doubleClicked.connect(self.file_name) #双击文件打开
            self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
            self.window1tree.show()
            QApplication.processEvents()
        else:
            QMessageBox.about(self, "提示!","文件夹没有任何有效文件\n请您先进行《保存数据》功能保存数据！！" )
    
    def file_name(self,Qmodelidx):   #判断文件
        win.time2 =datetime.datetime.now()
        win.sign = 1
        path = self.model.filePath(Qmodelidx) #文件的地址。
        win.filename = os.path.split(path)[1]
        pa = path[-4:]
        if pa == '.swf':    #如果是flash 文件
            self.read_flash(path)
        elif pa == '.zip':
            self.read_zip(path)
    
    
    
    def read_flash(self, path):  #读取flash 文件
        self.delLwin()
        self.sign = 2
        self.aswiget = QAxContainer.QAxWidget()
        self.win_layout2.addWidget(self.aswiget, 0, 0, 20, 20)
        self.aswiget.adjustSize()
        self.aswiget.setFocusPolicy(QtCore.Qt.StrongFocus)
        self.aswiget.setControl("{D27CDB6E-AE6D-11cf-96B8-444553540000}")  #flash的com接口
        self.aswiget.dynamicCall("LoadMovie(long,string)", 'something',path)
        self.aswiget.show()
    
    def read_zip(self, path):   #读取zip文件
        self.delLwin()
        self.sign = 3
        self.a= 1  #ppt页数变化,
        self.win_layout2.addWidget(self.win_lab, 0, 0, 20, 20)
        self.MaximumButton = QPushButton(self.win_lab)
        self.win_lab.setMouseTracking(True) # 设置widget鼠标跟踪
        self.MaximumButton.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(100, 100, 100)} ")
        QApplication.processEvents()
        self.width = self.win_lab.width()
        self.height = self.win_lab.height()
        self.MaximumButton.resize(20, 20)
        self.MaximumButton.move(self.width-22, self.height-22)
        self.MaximumButton.clicked.connect(self.select_class)
        
        if os.path.isdir(self.pa):  #建立新的文件夹
            fileNames = glob.glob(self.pa + r'/*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(self.pa)
        try:
            with zipfile.ZipFile(path, 'r') as zf:    
                for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
                    right_fn =  self.pa + '/'+fn
                    with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                        with zf.open(fn, 'r') as origin_file:  # 打开原文件
                            shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
            pa1 = self.pa + '/幻灯片1.jpg'
            pa2 = self.pa +'/幻灯片1.jpeg'
            img = Image.open(pa1)
            out = img.resize((self.width, self.height),Image.ANTIALIAS) #将图片改变分辨率为窗口大小
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            self.win_lab.show()
        except:
            QMessageBox.about(self, "提示!", "此文件是破损的，已删除！！" )
        
        
    
    def select_class(self):
        self.max =  max_widget(3, self.a)
        self.max.show()
    
    
    
    def close_max_widget(self, a):  #最大化变回原来的
        self.a = a
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width, self.height),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.win_lab.setPixmap (pixmap) # 在label上显示图片
#        self.win_lab.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
        
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
           

     
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if self.a != 0:
            if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
                self.cut_images()
            if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
                self.add_images()
            if  event.pos().y() > 620:  
                self.MaximumButton.show()
            else:
                self.MaximumButton.hide()
                QApplication.processEvents()

    
#用户练习小学界面
class Usr_window2_child1(QFrame):
    def __init__(self):
        super(Usr_window2_child1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.child1_win2but1 = QPushButton("返回")#课件-小学的控件
        self.child1_win2but2 = QPushButton("语文")
        self.child1_win2but3 = QPushButton("数学")
        self.child1_win2but4 = QPushButton("英语")
        self.window1tree = QTreeView()
        self.window1tree1 = QTextEdit()
        self.win_lab = QLabel()
        self.pa = 'D://项目数据库/tupian'
        self.a =0
        self.sign =0
        self.path = []
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        
        self.child1_win2but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child1_win2but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                        
        self.child1_win2but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child1_win2but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child1_win2but1.clicked.connect(self.return_fun)
        self.child1_win2but2.clicked.connect(lambda:self.select1(1))
        self.child1_win2but3.clicked.connect(lambda:self.select1(2))
        self.child1_win2but4.clicked.connect(lambda:self.select1(3))
        self.win_layout1.addWidget(self.child1_win2but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.child1_win2but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.child1_win2but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.child1_win2but4, 4, 0,1,2)
        self.win_layout2.addWidget(self.window1tree1, 0, 0, 20, 20)
        
        

    def return_fun(self):
        win.sign = 0
        if self.sign == 0 or self.sign == 1:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Usr_function())
        elif self.sign == 2:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude2_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "练习", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        
        
    def select1(self, sign):
        if sign == 1:
            self.TreeView("D:/项目数据库/练习/小学/" +win.grade+"/语文")
        elif sign ==2:
            self.TreeView("D:/项目数据库/练习/小学/" +win.grade+"/数学")
        elif sign == 3:
            self.TreeView("D:/项目数据库/练习/小学/" +win.grade+"/英语")

        
        
        
    def delLwin(self):
        if self.sign == 0:
            self.win_layout2.removeWidget(self.window1tree1)
            self.window1tree1.close()
        elif self.sign ==1:
            self.win_layout2.removeWidget(self.window1tree)
            self.window1tree.close()
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.win_lab)
            self.win_lab.close()
            self.a = 0
    
    
    
    def TreeView(self, path):   #目录树
        if(os.path.exists(path)):
            self.path = path
            self.delLwin()
            self.sign = 1
            self.model = QFileSystemModel()
            self.model.setRootPath(path)
            #model.setFilter(QDir.NoDotAndDotDot | QDir.Dirs)# 只显示文件夹，不显示文件中的其他小文件。
            self.window1tree.setModel(self.model)
            self.window1tree.setAnimated(False)
            self.window1tree.setIndentation(20)
            self.window1tree.setSortingEnabled(True)
            self.window1tree.setRootIndex(self.model.index(path)) #只显示设置的那个文件路径。
            self.window1tree.doubleClicked.connect(self.file_name) #双击文件打开
            self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
            self.window1tree.show()
            QApplication.processEvents()
        else:
            QMessageBox.about(self, "提示!","文件夹没有任何有效文件\n请您先进行《保存数据》功能保存数据！！" )
    
    def file_name(self,Qmodelidx):   #判断文件
        win.time2 =datetime.datetime.now()
        win.sign = 2
        path = self.model.filePath(Qmodelidx) #文件的地址。
        win.filename = os.path.split(path)[1]
        pa = path[-4:]
        if pa == '.zip':
            self.read_zip(path)
    
    def read_zip(self, path):   #读取zip文件
        self.delLwin()
        self.sign = 2
        self.a= 1  #ppt页数变化,
        self.win_layout2.addWidget(self.win_lab, 0, 0, 20, 20)
        self.MaximumButton = QPushButton(self.win_lab)
        self.win_lab.setMouseTracking(True) # 设置widget鼠标跟踪
        self.MaximumButton.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(100, 100, 100)} ")
        QApplication.processEvents()
        self.width = self.win_lab.width()
        self.height = self.win_lab.height()
        self.MaximumButton.resize(20, 20)
        self.MaximumButton.move(self.width-22, self.height-22)
        self.MaximumButton.clicked.connect(self.select_class)
        
        if os.path.isdir(self.pa):  #建立新的文件夹
            fileNames = glob.glob(self.pa + r'/*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(self.pa)
        try:
            with zipfile.ZipFile(path, 'r') as zf:    
                for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
                    right_fn =  self.pa + '/'+fn
                    with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                        with zf.open(fn, 'r') as origin_file:  # 打开原文件
                            shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
            pa1 = self.pa + '/幻灯片1.jpg'
            pa2 = self.pa +'/幻灯片1.jpeg'
            img = Image.open(pa1)
            out = img.resize((self.width, self.height),Image.ANTIALIAS) #将图片改变分辨率为窗口大小
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            self.win_lab.show()
        except:
            QMessageBox.about(self, "提示!", "此文件是破损的，已删除！！" )
 
    
    def select_class(self):
        self.max =  max_widget(4, self.a)
        self.max.show()
    
    
    
    def close_max_widget(self, a):  #最大化变回原来的
        self.a = a
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width, self.height),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.win_lab.setPixmap (pixmap) # 在label上显示图片
#        self.win_lab.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
        
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
           

     
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if self.a != 0:
            if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
                self.cut_images()
            if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
                self.add_images()
            if  event.pos().y() > 620:  
                self.MaximumButton.show()
            else:
                self.MaximumButton.hide()
                QApplication.processEvents()

    
    
#用户练习初中界面
class Usr_window2_child2(QFrame):
    def __init__(self):
        super(Usr_window2_child2, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.child2_win2but1 = QPushButton("返回")#课件-初中的控件
        self.child2_win2but2 = QPushButton("语文")
        self.child2_win2but3 = QPushButton("数学")
        self.child2_win2but4 = QPushButton("英语")
        self.child2_win2but5 = QPushButton("物理")
        self.child2_win2but6 = QPushButton("化学")
        self.child2_win2but7 = QPushButton("生物")
        self.child2_win2but8 = QPushButton("政治")
        self.child2_win2but9 = QPushButton("历史")
        self.child2_win2but10 = QPushButton("地理")
        self.window1tree = QTreeView()
        self.window1tree1 = QTextEdit()
        self.win_lab = QLabel()
        self.pa = 'D://项目数据库/tupian'
        self.a =0
        self.sign =0
        self.path = []
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        
        self.child2_win2but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                        
        self.child2_win2but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}") 
        self.child2_win2but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but1.clicked.connect(self.return_fun)
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.win_layout1.addWidget(self.child2_win2but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree1, 0, 0, 20, 20)
        
        
 

 
    def return_fun(self):
        win.sign = 0
        if self.sign == 0 or self.sign == 1:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Usr_function())
        elif self.sign == 2:
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude2_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "练习", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
        
    
    def select1(self, sign):
        if sign == 1:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/语文")
        elif sign ==2:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/数学")
        elif sign == 3:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/英语")
        elif sign == 4:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/物理")
        elif sign == 5:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/化学")
        elif sign == 6:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/生物")
        elif sign == 7:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/政治")
        elif sign == 8:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/历史")
        elif sign == 9:
            self.TreeView("D:/项目数据库/练习/初中/" +win.grade+"/地理")
        
    
    
    def delLwin(self):
        if self.sign == 0:
            self.win_layout2.removeWidget(self.window1tree1)
            self.window1tree1.close()
        elif self.sign ==1:
            self.win_layout2.removeWidget(self.window1tree)
            self.window1tree.close()
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.win_lab)
            self.win_lab.close()
            self.a = 0
    
    
    
    def TreeView(self, path):   #目录树
        if(os.path.exists(path)):
            self.path = path
            self.delLwin()
            self.sign = 1
            self.model = QFileSystemModel()
            self.model.setRootPath(path)
            #model.setFilter(QDir.NoDotAndDotDot | QDir.Dirs)# 只显示文件夹，不显示文件中的其他小文件。
            self.window1tree.setModel(self.model)
            self.window1tree.setAnimated(False)
            self.window1tree.setIndentation(20)
            self.window1tree.setSortingEnabled(True)
            self.window1tree.setRootIndex(self.model.index(path)) #只显示设置的那个文件路径。
            self.window1tree.doubleClicked.connect(self.file_name) #双击文件打开
            self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
            self.window1tree.show()
            QApplication.processEvents()
        else:
            QMessageBox.about(self, "提示!","文件夹没有任何有效文件\n请您先进行《保存数据》功能保存数据！！" )
    
    def file_name(self,Qmodelidx):   #判断文件
        win.time2 =datetime.datetime.now()
        win.sign = 2
        path = self.model.filePath(Qmodelidx) #文件的地址。
        win.filename = os.path.split(path)[1]
        path = self.model.filePath(Qmodelidx) #文件的地址。
        pa = path[-4:]
        if pa == '.zip':
            self.read_zip(path)
    
    def read_zip(self, path):   #读取zip文件
        self.delLwin()
        self.sign = 2
        self.a= 1  #ppt页数变化,
        self.win_layout2.addWidget(self.win_lab, 0, 0, 20, 20)
        self.MaximumButton = QPushButton(self.win_lab)
        self.win_lab.setMouseTracking(True) # 设置widget鼠标跟踪
        self.MaximumButton.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(100, 100, 100)} ")
        QApplication.processEvents()
        self.width = self.win_lab.width()
        self.height = self.win_lab.height()
        self.MaximumButton.resize(20, 20)
        self.MaximumButton.move(self.width-22, self.height-22)
        self.MaximumButton.clicked.connect(self.select_class)
        
        if os.path.isdir(self.pa):  #建立新的文件夹
            fileNames = glob.glob(self.pa + r'/*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(self.pa)
        try:
            with zipfile.ZipFile(path, 'r') as zf:    
                for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
                    right_fn =  self.pa + '/'+fn
                    with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                        with zf.open(fn, 'r') as origin_file:  # 打开原文件
                            shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
            pa1 = self.pa + '/幻灯片1.jpg'
            pa2 = self.pa +'/幻灯片1.jpeg'
            img = Image.open(pa1)
            out = img.resize((self.width, self.height),Image.ANTIALIAS) #将图片改变分辨率为窗口大小
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            self.win_lab.show()
        except:
            QMessageBox.about(self, "提示!", "此文件是破损的，已删除！！" )
 
    
    def select_class(self):
        self.max =  max_widget(5, self.a)
        self.max.show()
    
    
    
    def close_max_widget(self, a):  #最大化变回原来的
        self.a = a
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width, self.height),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.win_lab.setPixmap (pixmap) # 在label上显示图片
#        self.win_lab.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
        
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
           

     
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if self.a != 0:
            if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
                self.cut_images()
            if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
                self.add_images()
            if  event.pos().y() > 620:  
                self.MaximumButton.show()
            else:
                self.MaximumButton.hide()
                QApplication.processEvents()

    
    
    
    
#用户练习高中界面
class Usr_window2_child3(QFrame):
    def __init__(self):
        super(Usr_window2_child3, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.child2_win2but1 = QPushButton("返回")#课件-初中的控件
        self.child2_win2but2 = QPushButton("语文")
        self.child2_win2but3 = QPushButton("数学")
        self.child2_win2but4 = QPushButton("英语")
        self.child2_win2but5 = QPushButton("物理")
        self.child2_win2but6 = QPushButton("化学")
        self.child2_win2but7 = QPushButton("生物")
        self.child2_win2but8 = QPushButton("政治")
        self.child2_win2but9 = QPushButton("历史")
        self.child2_win2but10 = QPushButton("地理")
        self.window1tree = QTreeView()
        self.window1tree1 = QTextEdit()
        self.win_lab = QLabel()
        self.pa = 'D://项目数据库/tupian'
        self.a =0
        self.sign =0
        self.path = []
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        
        self.Lchild_win1 = QWidget()  #左侧控件布局
        self.win_layout1 = QGridLayout()  # 创建左侧部件的网格布局层
        self.Lchild_win1.setLayout(self.win_layout1)   # 设置左侧部件布局为网格    
        self.Rchild_win1 = QWidget()     #右侧控件布局
        self.win_layout2 = QGridLayout()   # 创建右侧部件的网格布局层
        self.Rchild_win1.setLayout(self.win_layout2)   # 设置右侧部件布局为网格
        self.layout.addWidget(self.Lchild_win1,0,0,20,2) # 左侧部件在第0行第0列，占20行2列
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        
        self.child2_win2but1.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but2.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")                        
        self.child2_win2but3.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but4.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but5.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but6.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but7.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}") 
        self.child2_win2but8.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but9.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but10.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}\
                                 QPushButton{background-color:rgb(170,200, 50)}\
                                 QPushButton:hover{background-color:rgb(50, 170, 200)}")
        self.child2_win2but1.clicked.connect(self.return_fun)
        self.layout.addWidget(self.Rchild_win1,0,2,20,20) # 右侧部件在第1行第3列，占20行20列   
        self.win_layout1.addWidget(self.child2_win2but1, 1, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but2, 2, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but3, 3, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but4, 4, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but5, 5, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but6, 6, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but7, 7, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but8, 8, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but9, 9, 0,1,2)
        self.win_layout1.addWidget(self.child2_win2but10, 10, 0,1,2)
        self.win_layout2.addWidget(self.window1tree1, 0, 0, 20, 20)
        
    def return_fun(self):
        win.sign = 0
        if self.sign == 0 or self.sign == 1:
            win.splitter.widget(0).setParent(None)
            win.splitter.insertWidget(0,Usr_function())
        elif self.sign == 2:
            
            self.TreeView(self.path)
            a = datetime.datetime.now()
            b = a - win.time2
            win.data1[3] = win.data1[3] + b.seconds
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data1 set stude2_day =(?) where numble=(?)",(win.data1[3],win.numble))
            conn.commit()
            conn.close()
            da = '%Y-%m-%d %H:%M:%S'
            a1 = datetime.datetime.strftime(win.time2, da)
            b1= datetime.datetime.strftime(a, da)
            sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("insert into User_data values(?,?,?,?)", (a1, "练习", win.filename, b1))
            conn.commit()
            c.close()
            conn.close()
    
    def select1(self, sign):
        if sign == 1:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/语文")
        elif sign ==2:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/数学")
        elif sign == 3:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/英语")
        elif sign == 4:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/物理")
        elif sign == 5:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/化学")
        elif sign == 6:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/生物")
        elif sign == 7:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/政治")
        elif sign == 8:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/历史")
        elif sign == 9:
            self.TreeView("D:/项目数据库/练习/高中/" +win.grade+"/地理")
        
    
    
    def delLwin(self):
        if self.sign == 0:
            self.win_layout2.removeWidget(self.window1tree1)
            self.window1tree1.close()
        elif self.sign ==1:
            self.win_layout2.removeWidget(self.window1tree)
            self.window1tree.close()
        elif self.sign == 2:
            self.win_layout2.removeWidget(self.win_lab)
            self.win_lab.close()
            self.a = 0
    
    
    
    def TreeView(self, path):   #目录树
        if(os.path.exists(path)):
            self.path = path
            self.delLwin()
            self.sign = 1
            self.model = QFileSystemModel()
            self.model.setRootPath(path)
            #model.setFilter(QDir.NoDotAndDotDot | QDir.Dirs)# 只显示文件夹，不显示文件中的其他小文件。
            self.window1tree.setModel(self.model)
            self.window1tree.setAnimated(False)
            self.window1tree.setIndentation(20)
            self.window1tree.setSortingEnabled(True)
            self.window1tree.setRootIndex(self.model.index(path)) #只显示设置的那个文件路径。
            self.window1tree.doubleClicked.connect(self.file_name) #双击文件打开
            self.win_layout2.addWidget(self.window1tree, 0, 0, 20, 20)
            self.window1tree.show()
            QApplication.processEvents()
        else:
            QMessageBox.about(self, "提示!","文件夹没有任何有效文件\n请您先进行《保存数据》功能保存数据！！" )
    
    def file_name(self,Qmodelidx):   #判断文件
        win.time2 =datetime.datetime.now()
        win.sign = 2
        path = self.model.filePath(Qmodelidx) #文件的地址。
        win.filename = os.path.split(path)[1]
        pa = path[-4:]
        if pa == '.zip':
            self.read_zip(path)
    
    def read_zip(self, path):   #读取zip文件
        self.delLwin()
        self.sign = 2
        self.a= 1  #ppt页数变化,
        self.win_layout2.addWidget(self.win_lab, 0, 0, 20, 20)
        self.MaximumButton = QPushButton(self.win_lab)
        self.win_lab.setMouseTracking(True) # 设置widget鼠标跟踪
        self.MaximumButton.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(100, 100, 100)} ")
        QApplication.processEvents()
        self.width = self.win_lab.width()
        self.height = self.win_lab.height()
        self.MaximumButton.resize(20, 20)
        self.MaximumButton.move(self.width-22, self.height-22)
        self.MaximumButton.clicked.connect(self.select_class)
        
        if os.path.isdir(self.pa):  #建立新的文件夹
            fileNames = glob.glob(self.pa + r'/*') 
            if fileNames:
                for fileName in fileNames:     #将pa 文件夹中的文件删除。
                   os.remove( fileName)
        else:
            os.mkdir(self.pa)
        try:
            with zipfile.ZipFile(path, 'r') as zf:    
                for fn in zf.namelist():  #循环压缩包中的文件并保存进新文件夹。
                    right_fn =  self.pa + '/'+fn
                    with open(right_fn, 'wb') as output_file:  # 创建并打开新文件
                        with zf.open(fn, 'r') as origin_file:  # 打开原文件
                            shutil.copyfileobj(origin_file, output_file)  # 将原文件内容复制到新文件
            pa1 = self.pa + '/幻灯片1.jpg'
            pa2 = self.pa +'/幻灯片1.jpeg'
            img = Image.open(pa1)
            out = img.resize((self.width, self.height),Image.ANTIALIAS) #将图片改变分辨率为窗口大小
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            self.win_lab.show()
        except:
            QMessageBox.about(self, "提示!", "此文件是破损的，已删除！！" )
 
    
    def select_class(self):
        self.max =  max_widget(6, self.a)
        self.max.show()
    
    
    
    def close_max_widget(self, a):  #最大化变回原来的
        self.a = a
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width, self.height),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.win_lab.setPixmap (pixmap) # 在label上显示图片
#        self.win_lab.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
        
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
            
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width, self.height),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.win_lab.setPixmap (pixmap) # 在label上显示图片
            #self.win_lab.setScaledContents (True) # 让图片自适应label大小
           

     
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if self.a != 0:
            if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
                self.cut_images()
            if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
                self.add_images()
            if  event.pos().y() > 620:  
                self.MaximumButton.show()
            else:
                self.MaximumButton.hide()
                QApplication.processEvents()

    
    
    
    
    
    

#播放图片最大化的设置
class max_widget(QWidget):
    def __init__(self, sign, a):
        super(max_widget, self).__init__()
        self.pa = 'D://项目数据库/tupian'
        self.sign = sign
        self.a = a
        self.setWindowFlags(Qt.FramelessWindowHint)  #无边框 
        self.desktop = QApplication.desktop()
        #获取显示器分辨率大小
        self.screenRect = self.desktop.screenGeometry()
        self.height1 = self.screenRect.height()
        self.width1 = self.screenRect.width()
        self.resize(self.width1, self.height1)
        self.setMouseTracking(True) # 设置widget鼠标跟踪
        self.lab2 = QtWidgets.QLabel(self)
        self.lab2.resize(self.width1, self.height1)
        self.MaximumButton1 = QPushButton(self)
        self.MaximumButton1.resize(20, 20)
        self.MaximumButton1.setStyleSheet("QPushButton{background-color:rgb(255,255, 255)}\
                            QPushButton:hover{background-color:rgb(50, 10, 50)} ")
        self.MaximumButton1.move(self.width1-24, self.height1-24)
        self.MaximumButton1.clicked.connect(self.closewin)
        self.lab2.setMouseTracking(True) # 设置widget鼠标跟踪
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
        out = img.resize((self.width1, self.height1),Image.ANTIALIAS) 
        out.save(pa2, 'jpeg')
        pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
        self.lab2.setPixmap (pixmap) # 在label上显示图片
        #self.lab2.setScaledContents (True) # 让图片自适应label大小

    def closewin(self):
        if self.sign == 1:
            win.usr_window1_child1.close_max_widget(self.a)
        elif self.sign ==2:
            win.usr_window1_child2.close_max_widget(self.a)
        elif self.sign == 3:
            win.usr_window1_child3.close_max_widget(self.a)
        elif self.sign == 4:
            win.usr_window2_child1.close_max_widget(self.a)
        elif self.sign == 5:
            win.usr_window2_child2.close_max_widget(self.a)
        elif self.sign == 6:
            win.usr_window2_child3.close_max_widget(self.a)
        self.close() 
   
    def mousePressEvent(self, event):   # 重写鼠标点击的事件
        self.desktop = QApplication.desktop() #获取屏幕分辨率
        self.screenRect = self.desktop.screenGeometry()
        self.y = self.screenRect.height()
        self.x = self.screenRect.width()
        if (event.button() == Qt.LeftButton) and (event.pos().x()<self.x/2):
            self.cut_images()
        if (event.button() == Qt.LeftButton) and (event.pos().x()>self.x/2):
            self.add_images()
    
    def add_images(self):  #下一页ppt
        self.a = self.a+1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a -1
            QMessageBox.about(self, "提示!", "这是最后一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width1, self.height1),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.lab2.setPixmap(pixmap)
        
    def cut_images(self):  #上一页ppt
        self.a = self.a-1
        pa1 = self.pa +"\幻灯片" + str(self.a) +".jpg"
        pa2 = self.pa +"\幻灯片" + str(self.a) +".jpeg"
        if not os.path.exists(pa1):
            self.a = self.a+1
            QMessageBox.about(self, "提示!", "这是第一页" )
        else:
            img = Image.open(pa1)  #将图片改变分辨率为self.lab窗口大小
            out = img.resize((self.width1, self.height1),Image.ANTIALIAS) 
            out.save(pa2, 'jpeg')
            pixmap = QPixmap (pa2) # 按指定路径找到图片，注意路径必须用双引号包围，不能用单引号
            self.lab2.setPixmap(pixmap)
    
    
    
    
    

 
#用户我的界面
class Usr_myself(QFrame):   #增加一个编辑资料的按钮
    def __init__(self):
        super(Usr_myself, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.returnBtn = QPushButton("返回")
        self.ExditBtn = QPushButton("编辑")
        self.chang_image = QPushButton("换头像")
        self.name = QLabel("姓名:")
        self.sex = QLabel("性别:")
        self.number = QLabel("手机号:")
        self.year  = QLabel("出生年月:")
        self.school = QLabel("学校:")
        self.grade = QLabel("年级:")
        self.amend = QPushButton("修改密码")
        self.withdraw = QPushButton('退出')
        self.tupian = QLabel()
        self.devise_Ui()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (100, 0, 0, 0)        
        self.name1 = QLabel(win.data[1])    #读取数据库中的信息，将信息输出label中
        self.sex1 = QLabel(win.data[3])
        self.number1 = QLabel(win.data[0])
        self.year1  = QLabel(win.data[2][0:4]+"年 " +win.data[2][5: ]+' 月')
        self.school1 = QLabel(win.data[4])
        self.grade1 = QLabel(win.data[5])
        self.returnBtn.setMaximumSize(60, 40)
        self.ExditBtn.setMaximumSize(60, 40)
        self.name.setMaximumSize(70, 40)
        self.sex.setMaximumSize(70, 40)
        self.number.setMaximumSize(70, 40)
        self.school.setMaximumSize(70, 40)
        self.year.setMaximumSize(100, 40)
        self.grade.setMaximumSize(70, 40)
        self.name1.setMaximumSize(350, 40)
        self.sex1.setMaximumSize(350, 40)
        self.number1.setMaximumSize(350, 40)
        self.school1.setMaximumSize(350, 40)
        self.grade1.setMaximumSize(350, 40)
        self.year1.setMaximumSize(350, 40)
        self.amend.setMaximumSize(500, 40)
        self.withdraw.setMaximumSize(500, 40)
        self.chang_image.setMaximumSize(90, 40)
        self.tupian.setMaximumSize(250, 250)
        self.chang_image.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.name.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.year.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.sex.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.number.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.school.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.grade.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.amend.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.withdraw.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(255,0,0,255);}")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.ExditBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.name1.setStyleSheet("QLabel{color:rgb(255,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.sex1.setStyleSheet("QLabel{color:rgb(255,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.year1.setStyleSheet("QLabel{color:rgb(255,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.number1.setStyleSheet("QLabel{color:rgb(255,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.school1.setStyleSheet("QLabel{color:rgb(255,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.grade1.setStyleSheet("QLabel{color:rgb(255,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.withdraw.clicked.connect(self.return_win)
        self.returnBtn.clicked.connect(self.return_fun)
        self.ExditBtn.clicked.connect(self.edit_fun)
        self.chang_image.clicked.connect(self.chang_fun)
        self.amend.clicked.connect(self.amend_fun)
        self.layout.addWidget(self.tupian, 1, 1, 4, 4)
        self.layout.addWidget(self.chang_image, 5, 2, 1, 2)
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.ExditBtn, 0, 10, 1, 1)
        self.layout.addWidget(self.name, 1, 6, 1, 1)
        self.layout.addWidget(self.name1, 1, 8, 1, 10)
        self.layout.addWidget(self.year, 2, 6, 1, 1)
        self.layout.addWidget(self.year1, 2, 8, 1, 10)
        self.layout.addWidget(self.sex, 3, 6, 1, 1)
        self.layout.addWidget(self.sex1, 3, 8, 1, 10)
        self.layout.addWidget(self.number, 4, 6, 1, 1)
        self.layout.addWidget(self.number1, 4, 8, 1, 10)
        self.layout.addWidget(self.school, 5, 6, 1, 1)
        self.layout.addWidget(self.school1, 5, 8, 1, 10)
        self.layout.addWidget(self.grade, 6, 6, 1, 1)
        self.layout.addWidget(self.grade1, 6, 8, 1, 10)
        self.layout.addWidget(self.amend, 7, 6, 1, 10)
        self.layout.addWidget(self.withdraw, 8, 6, 1, 10)
        self.image()
    
    def amend_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_amend())    
    
    def image(self):
        self.image_path = "D:/项目数据库/头像/image" +win.data[7]
        total=base64.b64decode(win.data[6])
        f=open(self.image_path,'wb')
        f.write(total)
        f.close()
        self.tupian.setPixmap(QPixmap(self.image_path))
        self.tupian.setScaledContents (True) # 让图片自适应label大小
        QApplication.processEvents()
    
    
    def chang_fun(self):
        path, _ = QFileDialog.getOpenFileName(self, '请选择文件', 
        'D:\\', 'image(*.jpg)')
        if path:
            self.file = os.path.splitext(path)[1]
            self.tupian.setPixmap(QPixmap(path))
            self.tupian.setScaledContents (True) # 让图片自适应label大小
            with open(path, "rb") as f:
                total=base64.b64encode(f.read())   #将文件转换为字节。
            f.close()
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            conn.execute("update User_data set tupian = (?),file = (?) where numble = (?)", (total, self.file, win.numble))
            conn.commit()
            conn.close()
            win.data[6] = total
            win.data[7] = self.file
        else:
            self.image()
            
        
    
    def return_win(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Choice_status())
        
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_function())    
        
    def edit_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_informent1())   

#用户修改密码
class Usr_amend(QFrame):
    def __init__(self):
        super(Usr_amend, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.usrlab = QLabel("账号:")
        
        self.amendlab1 = QLabel("原密码:")
        self.amendlab2 = QLabel("新密码:")
        self.amendlab3 = QLabel("确认密码:")
        self.amendedit1 = QLineEdit()
        self.amendedit2 = QLineEdit()
        self.amendedit3 = QLineEdit()
        self.sure = QPushButton("确认修改")
        self.returnBtn = QPushButton("返回")
        self.devise_Ui()
        
    def devise_Ui(self):
        self.usrlab1 = QLabel(win.numble)
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (350, 0, 0, 0)
        self.usrlab.setMaximumSize(80, 40)
        self.amendlab1.setMaximumSize(80, 40)
        self.amendlab2.setMaximumSize(80, 40)
        self.amendlab3.setMaximumSize(100, 40)
        #设置QLabel 的字体颜色，大小，
        self.usrlab.setStyleSheet("QLabel{color:rgb(100,100,100,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.usrlab1.setStyleSheet("QLabel{color:rgb(100,100,100,255);font-size:20px;font-weight:Bold;font-family:Arial;}")
        self.amendlab1.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")      
        self.amendlab2.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")      
        self.amendlab3.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:20px;font-weight:Bold;font-family:Arial;}")      
        self.usrlab1.setMaximumSize(420, 40)
        self.amendedit1.setMaximumSize(420, 40)
        self.amendedit2.setMaximumSize(420, 40)
        self.amendedit3.setMaximumSize(420, 40)
        self.sure.setMaximumSize(420, 40)
        self.amendedit1.setPlaceholderText("请输入原密码")
        self.amendedit2.setPlaceholderText("请输入新密码")
        self.amendedit3.setPlaceholderText("请重新输入密码")
        self.amendedit1.setFont(QFont("宋体" , 16))  #设置QLineEditn 的字体及大小
        self.amendedit2.setFont(QFont("宋体" , 16))
        self.amendedit3.setFont(QFont("宋体" , 16))
        self.amendedit1.setEchoMode(QLineEdit.Password)
        self.amendedit2.setEchoMode(QLineEdit.Password)
        self.amendedit3.setEchoMode(QLineEdit.Password)
        self.sure.setStyleSheet("QPushButton{ font-family:'宋体';font-size:28px;color:rgb(0,0,0,255);}")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:24px;color:rgb(0,0,0,255);}")
        self.returnBtn.setMaximumSize(60, 40)
        self.returnBtn.clicked.connect(self.return_fun)
        self.amendedit1.returnPressed.connect(self.enterPress1)
        self.amendedit2.returnPressed.connect(self.enterPress2)
        self.sure.clicked.connect(self.accept)
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.usrlab, 1, 3, 1, 1)
        self.layout.addWidget(self.usrlab1, 1, 5, 1, 5)
        self.layout.addWidget(self.amendlab1, 2, 3, 1, 1)
        self.layout.addWidget(self.amendedit1, 2, 5, 1, 5)
        self.layout.addWidget(self.amendlab2, 3, 3, 1, 1)
        self.layout.addWidget(self.amendedit2, 3, 5, 1, 5)
        self.layout.addWidget(self.amendlab3, 4, 3, 1, 1)
        self.layout.addWidget(self.amendedit3, 4, 5, 1, 5)
        self.layout.addWidget(self.sure, 5, 5, 1, 5)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_myself())
        
    def enterPress1(self):
        if len(self.amendedit1.text()) == 0:
            QMessageBox.about(self, "提示!", "原密码没有填写" )
            self.amendedit1.setFocus()
        else:
            self.amendedit2.setFocus()
            
    def enterPress2(self):
        if len(self.amendedit2.text()) == 0:
            QMessageBox.about(self, "提示!", "新密码框不能为空！" )
            self.amendedit2.setFocus()
        else:
            self.amendedit3.setFocus()
            
    def accept(self):
        if len(self.amendedit1.text()) == 0:
            QMessageBox.about(self, "提示!", "原密码没有填写" )
            self.amendedit1.setFocus()
        elif len(self.amendedit2.text()) == 0:
            QMessageBox.about(self, "提示!", "新密码框不能为空！" )
            self.amendedit2.setFocus()
        elif len(self.amendedit3.text()) == 0:
            QMessageBox.about(self, "提示!", "新密码框不能为空！" )
            self.amendedit3.setFocus()
        elif self.amendedit3.text() != self.amendedit2.text():
            QMessageBox.about(self, "提示!", "前后密码输入不一样！" )
            self.amendedit3.setFocus()
        else:
            sqlpath ='D:/项目数据库/数据库/Information.db'
            conn=sqlite3.connect(sqlpath)
            c=conn.cursor()
            c.execute("select * from User")
            sign = 0
            for variate in c.fetchall():
                if variate[0]==win.numble and variate[2]== self.amendedit1.text():
                    conn.execute("update User set password=(?) where numble=(?)",(self.amendedit2.text(),variate[0],))
                    conn.commit()
                    sign = 1
                    break
            c.close()
            conn.close()
            if sign ==0:
                QMessageBox.about(self, "提示!", "原密码输入错误！！" )
                self.amendedit1.setFocus()
            else:
                QMessageBox.about(self, "提示!", "修改成功！！" )
                time.sleep(1)
                win.splitter.widget(0).setParent(None)
                win.splitter.insertWidget(0, Usr_myself())
            
   

#用户我的编辑信息
class Usr_informent1(QFrame):  
    def __init__(self):
        super(Usr_informent1, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.sure = QPushButton("确认")
        self.returnBtn = QPushButton("返回")
        self.name =  QLabel("姓名:")
        self.year  = QLabel("出生年月")
        self.yearcb = QComboBox()
        self.monthcb = QComboBox()
        self.sex =  QLabel("性别:")
        self.sexcb = QComboBox()
        self.school = QLabel("学校:")
        self.grade = QLabel("选择年级")
        self.gradecb = QComboBox()
        self.nameEdit = QLineEdit()
        self.schoolEiit = QLineEdit()
        self.devise_Ui()
     
   
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
        self.layout.setContentsMargins (300, 0, 0, 0)
        yearnb = []
        for  i in range(1980, 2020):
            yearnb.append(str(i))
        monthmb = []
        for i in range(1, 13):
            monthmb.append(str(i))
        grade = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级', 
                     '初一', '初二', '初三', 
                     '高一', '高二', '高三']
        self.sex.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        self.school.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.grade.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.name.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.year.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:Arial;}")
        self.sure.setStyleSheet("QPushButton{ font-family:'宋体';font-size:26px;color:rgb(0,0,0,255);}")
        
        self.nameEdit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.schoolEiit.setFont(QFont("宋体" , 14))  #设置QLineEditn 的字体及大小
        self.name.setMaximumSize(50, 40)
        self.school.setMaximumSize(50, 40)
        self.returnBtn.setMaximumSize(60, 40)
        self.year.setMaximumSize(95, 40)
        self.sex.setMaximumSize(50, 40)
        self.grade.setMaximumSize(95, 40)
        self.nameEdit.setMaximumSize(420, 40)
        self.schoolEiit.setMaximumSize(420, 40)
        self.sure.setMaximumSize(420, 40)
        self.sexcb.setMaximumSize(420, 40)
        self.gradecb.setMaximumSize(420, 40)
        self.yearcb.setMaximumSize(220, 40)
        self.monthcb.setMaximumSize(175, 40)
        self.sexcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.yearcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.monthcb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.gradecb.setStyleSheet("QComboBox{font-family:'宋体';font-size: 22px;}")
        self.sexcb.addItems(['男', '女'])
        self.sexcb.setCurrentText(win.data[3])  #设置文本的默认选项
        self.yearcb.addItems(yearnb)
        self.yearcb.setCurrentText(win.data[2][0:4])  #设置文本的默认选项
        self.monthcb.addItems(monthmb)
        self.monthcb.setCurrentText(win.data[2][5:7])  #设置文本的默认选项
        self.gradecb.addItems(grade)
        self.gradecb.setCurrentText(win.data[5])  #设置文本的默认选项
        self.nameEdit.setText(win.data[1])
        self.schoolEiit.setText(win.data[4])
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.name, 1, 3, 1, 1)
        self.layout.addWidget(self.nameEdit, 1, 4, 1, 18)
        self.layout.addWidget(self.sex, 2, 3, 1, 1)
        self.layout.addWidget(self.sexcb, 2, 4, 1, 18)
        self.layout.addWidget(self.year, 3, 3, 1, 1)
        self.layout.addWidget(self.yearcb,3, 4, 1, 8 )
        self.layout.addWidget(self.monthcb, 3, 9, 1, 7)
        self.layout.addWidget(self.school, 4, 3, 1, 1)
        self.layout.addWidget(self.schoolEiit, 4, 4, 1, 18)
        self.layout.addWidget(self.grade, 5, 3, 1, 1)
        self.layout.addWidget(self.gradecb, 5, 4, 1, 18)
        self.layout.addWidget(self.sure, 6, 4, 1, 18)
        self.sure.clicked.connect(self.connect_fun)
        self.returnBtn.clicked.connect(self.return_fun)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_myself())
    
    def save_data(self):
        a =  self.nameEdit.text()
        b = self.yearcb.currentText() +'-' +self.monthcb.currentText()
        c = self.sexcb.currentText()
        d = self.schoolEiit.text()
        e = self.gradecb.currentText()
        win.grade = e
        win.data[1] = a
        win.data[2] = b
        win.data[3] = c
        win.data[4] = d
        win.data[5] = e
        sqlpath ='D:/项目数据库/数据库/Information.db'
        conn=sqlite3.connect(sqlpath)
        conn.execute("update User_data set name =(?),birthday=(?),sex=(?),school=(?),grade=(?) where numble=(?)",(a, b, c, d, e, win.numble))
        conn.commit()	
        conn.close()
    
    def connect_fun(self):
        win.splitter.widget(0).setParent(None)
        self.save_data()
        Usr_myself().devise_Ui()
        win.splitter.insertWidget(0, Usr_myself())
        
        





#用户学习报告的界面
class Usr_report(QFrame):
    def __init__(self):
        super(Usr_report, self).__init__()
        self.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.setFrameShadow(QtWidgets.QFrame.Raised)
        self.returnBtn = QPushButton("返回")
        self.day = QLabel("学习天数:")
        self.learntime = QLabel("学习总时长:")
        self.learncou =  QLabel("学习课件:")
        self.learnexe = QLabel("学习练习:")
        self.avglearn = QLabel("日均学习:")
        self.table  = QTableWidget()
        self.devise_Ui()
        self.information()
    
    def devise_Ui(self):
        self.horizontalLayout = QtWidgets.QHBoxLayout(self)
        self.layout = QGridLayout()
        self.win = QWidget()
        self.win.setLayout(self.layout)    #设置顶级布局管理器
        self.horizontalLayout.addWidget(self.win)
        self.win.setMouseTracking(True) # 设置widget鼠标跟踪
#        self.layout.setContentsMargins (300, 0, 0, 0)
        self.returnBtn.setStyleSheet("QPushButton{ font-family:'宋体';font-size:22px;color:rgb(0,0,0,255);}")
        self.day .setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:'宋体';}")
        self.learntime.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:'宋体';}")
        self.learncou.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:'宋体';}")
        self.learnexe.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:'宋体';}")
        self.avglearn.setStyleSheet("QLabel{color:rgb(0,0,0,255);font-size:22px;font-weight:Bold;font-family:'宋体';}")
        self.returnBtn.setMaximumSize(60, 40)
        self.day.setMaximumSize(120, 40)
        self.learntime.setMaximumSize(130, 40)
        self.learncou.setMaximumSize(120, 40)
        self.learnexe.setMaximumSize(120, 40)
        self.avglearn.setMaximumSize(120, 40)
        self.returnBtn.clicked.connect(self.return_fun)
        self.layout.addWidget(self.returnBtn, 0, 0, 1, 1)
        self.layout.addWidget(self.day, 0, 5, 1, 1)
        self.layout.addWidget(self.learntime, 0, 8, 1, 1)
        self.layout.addWidget(self.learncou, 1, 2, 1, 1)
        self.layout.addWidget(self.learnexe, 1, 5, 1, 1)
        self.layout.addWidget(self.avglearn, 1, 8, 1, 1)
        
    
    def information(self):
        new = datetime.datetime.now()
        abcd = '%Y-%m-%d %H:%M:%S'
        a1 = datetime.datetime.strptime(win.data1[1], abcd)
        a = (new -a1).days +1
        self.join =  QLabel("已加入"+ str(a)+ "天")
        self.day1 = QLabel(str(win.data1[2])+"天")
        ab = win.data1[3]+win.data1[4]
        if (ab/3600)>1:
            ac = str(int(ab/3600)) +'时'+str(round((ab/3600 - int(ab/3600))*60, 2)) +"分"
        else:
            ac = str(round(ab/60, 2))  +"分"
        self.learntime1 = QLabel(ac)
        b = win.data1[3]
        if (b/3600)>1:
            c = str(int(b/3600)) +'时'+str(round((b/3600 - int(b/3600))*60, 2)) +"分"
        else:
            c = str(round(b/60, 2))  +"分"
        self.learncou1 =  QLabel(c)
        d = win.data1[4]
        if (d/3600)>1:
            e = str(int(d/3600)) +'时'+str(round((d/3600 - int(d/3600))*60, 2)) +"分"
        else:
            e = str(round(d/60, 2))+"分"
        self.learnexe1 = QLabel(e)
        ad = ab/win.data1[2]
        if (ad/3600)>1:
            ae = str(int(ad/3600)) +'时'+str(round((ad/3600 - int(ad/3600))*60, 2)) +"分"
        else:
            ae = str(round(ad/60, 2)) +"分"
        self.avglearn1 = QLabel(ae)
        self.join.setStyleSheet("QLabel{color:rgb(0,200,0,255);font-size:20px;font-weight:Bold;font-family:'宋体';}")
        self.day1 .setStyleSheet("QLabel{color:rgb(0,255,0,255);font-size:20px;font-weight:Bold;font-family:'宋体';}")
        self.learntime1.setStyleSheet("QLabel{color:rgb(0,255,0,255);font-size:20px;font-weight:Bold;font-family:'宋体';}")
        self.learncou1.setStyleSheet("QLabel{color:rgb(0,255,0,255);font-size:20px;font-weight:Bold;font-family:'宋体';}")
        self.learnexe1.setStyleSheet("QLabel{color:rgb(0,255,0,255);font-size:20px;font-weight:Bold;font-family:'宋体';}")
        self.avglearn1.setStyleSheet("QLabel{color:rgb(0,255,0,255);font-size:20px;font-weight:Bold;font-family:'宋体';}")
        self.table.setStyleSheet("QTableWidget{background-color:rgb(255,255,255);font:13pt '宋体';font-weight:Bold;};");
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        sqlpath = "D:/项目数据库/数据库/SQ"+str(win.numble)+"L.db"
        conn=sqlite3.connect(sqlpath)
        c=conn.cursor()
        c.execute("select * from User_data")
        data = c.fetchall()
        b = len(data)
        self.table.setRowCount(b)
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels(['开始学习时间', '文件类型', '文件名', '结束时间', '学习时长'])
        i= 0
        for variate in data:
            self.table.setItem(i, 0, QTableWidgetItem(variate[0]))
            self.table.setItem(i, 1, QTableWidgetItem(variate[1]))
            self.table.setItem(i, 2, QTableWidgetItem(variate[2]))
            self.table.setItem(i, 3, QTableWidgetItem(variate[3]))
            min = (datetime.datetime.strptime(variate[3], abcd)- datetime.datetime.strptime(variate[0], abcd)).seconds
            if (min/3600)>1:
                ac = str(int(min/3600)) +'时'+str(round((min/3600 - int(min/3600))*60, 2)) +"分"
            else:
                ac = str(round(min/60, 2))  +"分"
            self.table.setItem(i, 4, QTableWidgetItem(ac))
            i+=1
            
        self.join.setMaximumSize(300, 40)
        self.day1.setMaximumSize(150, 40)
        self.learntime1.setMaximumSize(150, 40)
        self.learncou1.setMaximumSize(150, 40)
        self.learnexe1.setMaximumSize(150, 40)
        self.avglearn1.setMaximumSize(150, 40)
        self.layout.addWidget(self.day1, 0, 6, 1, 1)
        self.layout.addWidget(self.join, 0, 2, 1, 1)
        self.layout.addWidget(self.learntime1, 0, 9, 1, 1)
        self.layout.addWidget(self.learncou1, 1, 3, 1, 1)
        self.layout.addWidget(self.learnexe1, 1, 6, 1, 1)
        self.layout.addWidget(self.avglearn1, 1, 9, 1, 1)
        self.layout.addWidget(self.table, 3, 0, 1, 10)
    
    def return_fun(self):
        win.splitter.widget(0).setParent(None)
        win.splitter.insertWidget(0, Usr_function())
        
        

#主函数
if __name__ == "__main__":
    app = QApplication(sys.argv)
    win = QUnFrameWindow()
    win.show()
    sys.exit(app.exec_())

  
